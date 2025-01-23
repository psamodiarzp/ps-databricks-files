# Databricks notebook source
!/databricks/python3/bin/python -m pip install --upgrade pip
!/databricks/python3/bin/python -m pip install -U kaleido

# COMMAND ----------

!pip install python-docx
!pip install plotly
!pip install -U kaleido

# COMMAND ----------

from datetime import datetime, timedelta
from datetime import date
import calendar
import pandas as pd
import os
import io
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.offline as pyo
import docx
import plotly.io as pio
from docx import Document
from docx.shared import Inches
from PIL import Image
import math
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# COMMAND ----------

def percentage_conversion(value):
    return f'{value/100:.2%}'
    
def convert_week_to_date_month(value):
    date_values = []
    for i in range(len(value)):
        week_int = int(value[i])
        date = datetime.date(datetime(current_year, 1, 1) + relativedelta(weeks=week_int,weekday=0))
        formated_date = date.strftime('%d %B')
        date_values.append(formated_date)
    return date_values

def week_to_date(week_number):
    date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number,days=1))
    return date.strftime('%d %B')
    
def week_to_date_running_avg(week_number):
    start_date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number,days=1))
    end_date = start_date + timedelta(days=-28)
    return f"{end_date.strftime('%d %B')} - {start_date.strftime('%d %B')}"

current_dt = date.today().strftime('%Y-%m-%d')
today_date = date.today()
num_weeks = 4 #Input number of weeks we will be comparing the data for
waterfall_num_weeks = 2 #Input number of weeks we will be comparing the data for in the waterfall charts
num_days = num_weeks * 7
dt = datetime.strptime(current_dt, '%Y-%m-%d')
current_monday = (dt - timedelta(days = dt.weekday()))
week = dt.isocalendar()[1] #current_monday is also the ending date
starting_dt = (current_monday - timedelta(days=num_days))
ending_dt = (current_monday - timedelta(days=1)).strftime('%Y-%m-%d')
waterfall_starting_dt = (current_monday - timedelta(days=waterfall_num_weeks*7))
current_date = datetime.now().date()
previous_week_monday = current_date - timedelta(days=current_date.weekday(), weeks=1)
previous_week_month = previous_week_monday.strftime("%B")
previous_week_day = previous_week_monday.strftime("%d")
current_year = datetime.now().year
prev_week_monday = today_date - timedelta(days=today_date.weekday(), weeks=1)
this_week_sunday = prev_week_monday + timedelta(days=6)

# Get the month names
month_name_today = calendar.month_name[today_date.month]
month_name_prev_week_monday = calendar.month_name[prev_week_monday.month]
month_name_this_week_sunday = calendar.month_name[this_week_sunday.month]


print('This report was generated on '+ str(current_dt) + ', in the week of '+ str(current_monday.strftime('%Y-%m-%d'))+ ' that is the week '+str(week)+' of the current year')
print('This report is comparing data between weeks of '+ str(starting_dt.strftime('%Y-%m-%d')) + ' and '+ str(ending_dt))
print('The waterfall charts will be comparing data between weeks of '+ str(waterfall_starting_dt.strftime('%Y-%m-%d')) + ' and '+ str(ending_dt))


# COMMAND ----------


base_db = sqlContext.sql("""
    SELECT * ,
    weekofyear(payments.created_date) as payments_created_week
    FROM aggregate_pa.magic_payment_fact as payments
    WHERE payments.created_date >= '{0}' 
    and payments.created_date < '{1}'
    and payments.segment is not null
    """.format(starting_dt.strftime('%Y-%m-%d'),ending_dt))
base_df = base_db.toPandas()
base_df.head()

# COMMAND ----------

#wtu_group_by

wtu_grouped_df = base_df[base_df['payment_success']>0].groupby(['segment','payments_created_week']).agg({
    'merchant_id':'nunique',
}).reset_index()

# COMMAND ----------

#wtu_pivot

wtu_df = pd.pivot_table(wtu_grouped_df, values='merchant_id', columns='payments_created_week', index='segment').reset_index()
wtu_df = wtu_df.dropna()
integer_columns = wtu_df.columns[1:]
wtu_df[integer_columns] = wtu_df[integer_columns].astype(int)

#converting week number to date
for weeks in range(week - num_weeks, week):
    val = week_to_date(weeks)
    wtu_df = wtu_df.rename(columns={weeks:val})

# total
totals = wtu_df.sum().to_frame().T
totals['segment'] = 'Total'

wtu_df = pd.concat([wtu_df, totals], ignore_index=True)
wtu_df

# COMMAND ----------

 #mtu_group_by

mtu_grouped_df = base_df[base_df['running_30day_payment_success']==1].groupby(['payments_created_week']).agg({
    'merchant_id':'nunique',
})
mtu_grouped_df

# COMMAND ----------

#mtu_pivot_table

mtu_df = pd.pivot_table(mtu_grouped_df, values='merchant_id', index='payments_created_week')
mtu_df = mtu_df.rename(columns={'merchant_id': 'MTU Count'})

mtu_df.index = mtu_df.index.map(week_to_date_running_avg)
mtu_df = mtu_df.reset_index()
mtu_df = mtu_df.rename(columns={'payments_created_week':'Week Range'})
mtu_df

# COMMAND ----------

#std e-commerce dataframe


std_ecommerce_db = sqlContext.sql("""
Select 
weekofyear(payments_fact_druid.payments_created_date) as payments_created_week,
ROUND((( COUNT(DISTINCT CASE WHEN payments_fact_druid.payments_authorized_at IS NOT NULL then payments_fact_druid.payments_id END ) )*1.000000/( COUNT(DISTINCT payments_fact_druid.payments_id ) ))*100 ,4 ) as Std_checkout_Ecommerce_Sr
FROM warehouse.payments  AS payments_fact_druid
where payments_fact_druid.payments_created_date between '{0}' and '{1}'
and (CASE
          WHEN payments_fact_druid.payment_analytics_library = 1 THEN 'CHECKOUTJS'
          WHEN payments_fact_druid.payment_analytics_library = 2 THEN 'RAZORPAYJS'
          WHEN payments_fact_druid.payment_analytics_library = 3 THEN 'S2S'
          WHEN payments_fact_druid.payment_analytics_library = 4 THEN 'CUSTOM'
          WHEN payments_fact_druid.payment_analytics_library = 5 THEN 'DIRECT'
          WHEN payments_fact_druid.payment_analytics_library = 6 THEN 'PUSH'
          WHEN payments_fact_druid.payment_analytics_library = 7 THEN 'LEGACYJS'
          WHEN payments_fact_druid.payment_analytics_library = 8 THEN 'HOSTED' /* Hosted and embedded checkout identifers are added at the end of Nov, 2020 */
          WHEN payments_fact_druid.payment_analytics_library = 9 THEN 'EMBEDDED'
          ELSE 'UNKNOWN'
          END ) IN ('CHECKOUTJS', 'HOSTED') 
AND (payments_fact_druid.merchants_category2 = 'ecommerce') 
group by 1
ORDER BY 1
    """.format(starting_dt.strftime('%Y-%m-%d'),ending_dt))
std_ecommerce_df = std_ecommerce_db.toPandas().set_index('payments_created_week')
std_ecommerce_df.index = std_ecommerce_df.index.map(week_to_date)
std_ecommerce_df['Std_checkout_Ecommerce_Sr'] = std_ecommerce_df['Std_checkout_Ecommerce_Sr'].apply(percentage_conversion)
std_ecommerce_df.head()

# COMMAND ----------

#orders grouped

orders_grouped = base_df.groupby(['payments_created_week']).agg({
    'payment_attempt': 'sum',
    'payment_success': 'sum',
})
#payment_success for cod
orders_grouped['COD Payment Success'] = (base_df.loc[base_df['cod_or_noncod'] == 'cod'].groupby('payments_created_week')['payment_success'].sum()).astype(int)

#payment success for non cod
orders_grouped['Non COD Payment Success'] = (base_df.loc[base_df['cod_or_noncod'] == 'non cod'].groupby('payments_created_week')['payment_success'].sum())

#COD%
orders_grouped['COD%'] = ((orders_grouped['COD Payment Success'] / orders_grouped['payment_success'])*100).apply(percentage_conversion)

#NON_COD%
orders_grouped['Non COD%'] = ((orders_grouped['Non COD Payment Success'] / (orders_grouped['COD Payment Success'] + orders_grouped['Non COD Payment Success']))*100).apply(percentage_conversion)

# #Overall SR%
orders_grouped['Overall SR%'] = ((orders_grouped['payment_success'] / orders_grouped['payment_attempt'])*100).apply(percentage_conversion)

#Non COD SR%
orders_grouped['Non COD SR%'] = ((orders_grouped['Non COD Payment Success'] / (base_df.loc[base_df['cod_or_noncod'] == 'non cod'].groupby('payments_created_week')['payment_attempt'].sum()))*100).apply(percentage_conversion)

orders_grouped = orders_grouped.rename(columns={'payment_attempt': 'Payment Attempt','payment_success':'Payment Success'})
orders_grouped.index = orders_grouped.index.map(week_to_date)
orders_grouped

# COMMAND ----------

#Orders

orders_df = pd.merge(orders_grouped, std_ecommerce_df, on='payments_created_week')
orders_df = orders_df.rename(columns={'Std_checkout_Ecommerce_Sr':'SR% for ECommerce[Std Checkout]'})
orders_df = orders_df.transpose().reset_index()
orders_df =  orders_df.rename(columns={'index':'Payment Week'})
orders_df


# COMMAND ----------

#Merchant Sql
merchant_sql = sqlContext.sql("""
    SELECT merchants.id as merchant_id, merchants.website 
    FROM realtime_hudi_api.merchants as merchants
    """)
merchant_sql = merchant_sql.toPandas()
merchant_sql.head()


# COMMAND ----------

#merge merchants table and payments table

merchant_and_payment_table_merged = pd.merge(merchant_sql, base_df, on='merchant_id')
merchant_and_payment_table_merged.head()


# COMMAND ----------

#group by gmv on week 
gmv_grouped_data = merchant_and_payment_table_merged.groupby(['payments_created_week','merchant_id','website' ])['gmv_authorized'].sum().reset_index()
gmv_pivot = pd.pivot_table(gmv_grouped_data, values='gmv_authorized', columns='payments_created_week', index=['merchant_id','website']).reset_index()
gmv_pivot = gmv_pivot.fillna(0)

#converting week number to date
for weeks in range(week - num_weeks, week):
    val = week_to_date(weeks)
    gmv_pivot = gmv_pivot.rename(columns={weeks:val})


last_two_weeks = gmv_pivot.columns[-2:]
gmv_pivot['week2 - week1'] = gmv_pivot[last_two_weeks[1]] - gmv_pivot[last_two_weeks[0]]
gmv_pivot['abs( week2 - week1 )'] = (abs(gmv_pivot[last_two_weeks[1]] - gmv_pivot[last_two_weeks[0]])).round(2)
gmv_pivot = gmv_pivot.sort_values(by='abs( week2 - week1 )',ascending=False)
gmv_pivot.head(10)


# COMMAND ----------

# GMV

gmv_pivot = gmv_pivot.sort_values(by='abs( week2 - week1 )',ascending=False)
filtered_df = gmv_pivot[gmv_pivot['abs( week2 - week1 )'] > 500000]
other_gmv_sum = gmv_pivot[gmv_pivot['abs( week2 - week1 )'] <= 500000]['week2 - week1'].sum()
previous_week_sum = gmv_pivot.iloc[:,-4].sum()
subtotal_sum = gmv_pivot.iloc[:,-3].sum()

#creating row for others, previous week and subtotal 
filtered_websites = filtered_df[['website', 'week2 - week1']].rename(columns={'week2 - week1':'GMV'}).round(2)
other_row = pd.DataFrame({'website': ['Others'], 'GMV': [other_gmv_sum]}).round(2)
previous_week = pd.DataFrame({'website':[f"Week of {previous_week_day} {previous_week_month}"],'GMV':[previous_week_sum]}).round(2)
subtotal = pd.DataFrame({'website':['subtotal'],'GMV':[subtotal_sum]}).round(2)

#concating previous week, subtotal and other
result_df = pd.concat([filtered_websites, other_row, previous_week, subtotal]).reset_index()

last_second_row = result_df.iloc[-2]
result_df = pd.concat([last_second_row.to_frame().T, result_df.drop(result_df.index[-2])])
result_df


# COMMAND ----------


previous_week_total = gmv_pivot.iloc[:,-4].sum()
subtotal = gmv_pivot.iloc[:,-3].sum()

merchant_gmv_values = result_df['GMV'].values
merchant_website = result_df['website'].values

base_index = result_df[result_df['website'] == f"Week of {previous_week_day} {previous_week_month}"].index[0]
base = previous_week_total

measures =  ['absolute'] + ['relative'] * (len(merchant_gmv_values) - 2) + ['total']

trace = go.Waterfall(
    x=merchant_website,
    y=merchant_gmv_values,
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.7,
    text=result_df['GMV'].apply(lambda x: round(x / 100000, 2)).apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    textposition='outside',
)

layout = go.Layout(
    title='GMV Merchant Influence(In Lakhs)',
    title_font=dict(size=24),
    xaxis=dict(title='Website'),
    yaxis=dict(title='GMV'),
    width=750,
    height=1050,
)

gmv_fig = go.Figure(data=trace, layout=layout)
gmv_fig = gmv_fig.update_xaxes(tickfont=dict(size=30,family='Arial, bold'))
gmv_fig.show()


# COMMAND ----------

#orders(MID wise)

orders_grouped_data = merchant_and_payment_table_merged.groupby(['payments_created_week','merchant_id','website' ])['order_count_payment_success'].sum().reset_index()

orders_pivot = pd.pivot_table(orders_grouped_data, values='order_count_payment_success', columns='payments_created_week', index=['merchant_id','website']).reset_index()
orders_pivot = orders_pivot.fillna(0)

#converting week number to date
for weeks in range(week - num_weeks, week):
    val = week_to_date(weeks)
    orders_pivot = orders_pivot.rename(columns={weeks:val})


last_two_weeks = orders_pivot.columns[-2:]
orders_pivot['week2 - week1'] = orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]]
orders_pivot['abs( week2 - week1 )'] = (abs(orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]])).round(2)
orders_pivot = orders_pivot.sort_values(by="abs( week2 - week1 )",ascending=False)
orders_pivot.head(10)


# COMMAND ----------


filtered_df = orders_pivot[orders_pivot['abs( week2 - week1 )'] > 500]
other_orders_sum = orders_pivot[orders_pivot['abs( week2 - week1 )'] <= 500]['week2 - week1'].sum()
previous_week_sum = orders_pivot.iloc[:,-4].sum()
subtotal_sum = orders_pivot.iloc[:,-3].sum()

filtered_websites = filtered_df[['website', 'week2 - week1']].rename(columns={'week2 - week1':'Orders'}).round(2)
other_row = pd.DataFrame({'website': ['Others'], 'Orders': [other_orders_sum]}).round(2)
previous_week = pd.DataFrame({'website':[f"Week of {previous_week_day} {previous_week_month}"],'Orders':[previous_week_sum]}).round(2)
subtotal = pd.DataFrame({'website':['subtotal'],'Orders':[subtotal_sum]}).round(2)

result_df = pd.concat([filtered_websites, other_row, previous_week, subtotal]).reset_index()
last_second_row = result_df.iloc[-2]
result_df = pd.concat([last_second_row.to_frame().T, result_df.drop(result_df.index[-2])])
result_df


# COMMAND ----------

previous_week_total = gmv_pivot.iloc[:,-4].sum()
subtotal = gmv_pivot.iloc[:,-3].sum()

merchant_orders_values = result_df['Orders'].values
merchant_website = result_df['website'].values

base_index = result_df[result_df['website'] == f"Week of {previous_week_day} {previous_week_month}"].index[0]
base = previous_week_total

measures =  ['absolute'] + ['relative'] * (len(merchant_orders_values) - 2) + ['total']

trace = go.Waterfall(
    x=merchant_website,
    y=merchant_orders_values,
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.7,
    text=result_df['Orders'].apply(lambda x: round(x , 2)).apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    textposition='outside',
)

layout = go.Layout(
    title='Orders Merchant Influence',
    title_font=dict(size=24),
    xaxis=dict(title='Website'),
    yaxis=dict(title='Orders'),
    width=750,
    height=1050,
)

orders_fig = go.Figure(data=trace, layout=layout)
orders_fig = orders_fig.update_xaxes(tickfont=dict(size=30,family='Arial, bold'))
orders_fig.show()


# COMMAND ----------

cod_non_cod_grouped_data = merchant_and_payment_table_merged.groupby(['merchant_id','website','payments_created_week','cod_or_noncod'])['order_count_payment_success'].sum()

cod_non_cod_grouped_data.index = cod_non_cod_grouped_data.index.set_levels(cod_non_cod_grouped_data.index.levels[2].map(week_to_date), level='payments_created_week')
cod_non_cod_grouped_data = cod_non_cod_grouped_data.reset_index()

orders_sum = merchant_and_payment_table_merged.groupby(['merchant_id','website','payments_created_week','cod_or_noncod'])['order_count_payment_success'].sum().reset_index()
orders_sum_pivot = pd.pivot_table(orders_sum,index=['merchant_id','website'],values=['order_count_payment_success'],aggfunc='sum')

cod_non_cod_pivot = pd.pivot_table(cod_non_cod_grouped_data,index=['merchant_id','website'],columns=['payments_created_week','cod_or_noncod'],values=['order_count_payment_success'])
cod_non_cod_pivot = cod_non_cod_pivot.fillna(0)
cod_non_cod_pivot = cod_non_cod_pivot.drop(cod_non_cod_pivot.columns[0:4],axis =1)


for weeks in range(week - num_weeks + 2, week):
    cod = []
    non_cod = []
    val = week_to_date(weeks)
    index = weeks-(week - num_weeks + 2)
    for row_count in range(len(cod_non_cod_pivot.iloc[:,index])):
        value_cod = cod_non_cod_pivot.iloc[row_count].values[index*2]
        value_non_cod = cod_non_cod_pivot.iloc[row_count].values[(index*2)+1]
        total = value_cod + value_non_cod
        cod_result = value_cod/total*100
        non_cod_result = value_non_cod/total*100
        cod.append(percentage_conversion(cod_result))
        non_cod.append(percentage_conversion(non_cod_result))
    cod_non_cod_pivot[('order_count_payment_success',val,'cod%')] = cod
    cod_non_cod_pivot[('order_count_payment_success',val,'non cod%')] = non_cod
    
cod_non_cod_pivot = cod_non_cod_pivot.drop(cod_non_cod_pivot.columns[0:4],axis=1)
cod_non_cod_pivot = cod_non_cod_pivot.fillna(0).round(2)
cod_non_cod_pivot['orders'] = orders_sum_pivot
cod_non_cod_pivot= cod_non_cod_pivot.sort_values(by='orders',ascending=False)
cod_non_cod_pivot = cod_non_cod_pivot.drop(cod_non_cod_pivot.columns[4:5],axis=1)

cod_non_cod_pivot = cod_non_cod_pivot.head(10).reset_index()

cod_non_cod_pivot

# COMMAND ----------

def save_plotly_figure_as_image(fig, document,width_inches, height_inches, dpi):
    print(document)
    image_stream = io.BytesIO()
    fig.write_image(image_stream, format='png',width=width_inches * dpi, height=height_inches * dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def create_document_from_multi_level_list(data):
    doc = Document()
    for heading in data:
        section_heading = heading[0]
        section_heading_content = heading[1:]
        heading_formating = doc.add_heading(section_heading, level=1)
        heading_run = heading_formating.runs[0]
        heading_run.font.size = Pt(16)  
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_formating.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        review_date_paragraph = f"Review Date - {month_name_today} {today_date.day}, {today_date.year} ({month_name_prev_week_monday} {prev_week_monday.day} - {month_name_this_week_sunday} {this_week_sunday.day}, {this_week_sunday.year})"
        heading_data = doc.add_paragraph(review_date_paragraph)
        heading_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for section in section_heading_content:
            section_title = section[0]
            section_content = section[1:]
            doc.add_heading(section_title, level=1)
            
            if not isinstance(section_content, list):
                section_content = [section_content]

            for item in section_content:
                if isinstance(item, list):
                    subheading = item[0]
                    subcontent = item[1]
                    headings=doc.add_heading(subheading, level=2)
                    headings.paragraph_format.space_before = Pt(8)
                    headings.paragraph_format.space_before = Pt(10)
                    if isinstance(subcontent, go.Figure):
                        save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=320)
                
                    elif isinstance(subcontent.columns, pd.MultiIndex):
                        paragraph = doc.add_paragraph()
                        headings.paragraph_format.space_before = Pt(8)
                        headings.paragraph_format.space_after = Pt(10)
                        table = doc.add_table(subcontent.shape[0]+2, subcontent.shape[1])
                        table.style = 'LightGrid-Accent1'
                        last_paragraph = table.rows[-1].cells[0].paragraphs[-1]
                        last_paragraph.paragraph_format.space_after = Pt(10)  

                        
                        for i,col_tuple in enumerate(subcontent.columns):
                            for j, col in enumerate(col_tuple):
                                if(len(col_tuple)==3):
                                        
                                    if(j%2 == 0 ):
                                        table.cell(1,i).text = str(col)
                                    else:
                                        continue
                        for i,col_tuple in enumerate(subcontent.columns):
                            for j, col in enumerate(col_tuple):
                                if(j%2 != 0):
                                    table.cell(0,i).text = str(col)
                                else:
                                    continue
                                
                        for i, row_tuple in enumerate(subcontent.itertuples(index=False)):
                            for j, value in enumerate(row_tuple):
                                table.cell(i + 2, j).text = str(value)
                        
                        table_spacing = doc.add_paragraph()
                        table_spacing.paragraph_format.space_before = Pt(8)
                        table_spacing.paragraph_format.space_after = Pt(8)

                    elif isinstance(subcontent, pd.DataFrame):
                        doc.add_paragraph()
                        table = doc.add_table(subcontent.shape[0]+1, subcontent.shape[1])
                        table.style = 'LightGrid-Accent1'
                        for i, column in enumerate(subcontent.columns):
                            table.cell(0, i).text = str(column)
                            for j, value in enumerate(subcontent[column]):
                                table.cell(j+1, i).text = str(value)
                    
                else:
                    raise ValueError("Invalid content format. Expected list.")
    return doc


# COMMAND ----------


data_list = [
    ['Magic Checkout Weekly Health Metrics',
        ['Acquisition',
            ['MTU',mtu_df],
            ['WTU',wtu_df],
            ['Orders', orders_df],
            ['GMV Waterfall Chart',gmv_fig],
            ['Orders Waterfall Chart',orders_fig],
            ['COD - Non-COD tracking Merchant wise (Top 10 Merchants by Orders Volume)',cod_non_cod_pivot]
            
            
        ],
        
    ]
]
document = create_document_from_multi_level_list(data_list)
document.save('hm_data_list.docx')

# COMMAND ----------

dbutils.fs.cp("file:/databricks/driver/hm_data_list.docx", "dbfs:/FileStore/shared_transfer/Pallavi_Samodia/hm_data_list.docx")
# https://razorpay-dev.cloud.databricks.com/files/shared_transfer/Pallavi_Samodia/hm_data_list.docx

# COMMAND ----------


