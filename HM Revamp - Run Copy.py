# Databricks notebook source
from datetime import datetime, timedelta

from datetime import date

# COMMAND ----------

from datetime import datetime


# COMMAND ----------


import pandas as pd
import numpy as np
import plotly.graph_objects as go


# COMMAND ----------

import os
import io
import calendar
import matplotlib.pyplot as plt
# import seaborn as sns
import plotly.graph_objects as go
import plotly.offline as pyo

import plotly.io as pio


from PIL import Image
import math


# COMMAND ----------


!/databricks/python3/bin/python -m pip install --upgrade pip
!/databricks/python3/bin/python -m pip install -U kaleido

# COMMAND ----------

#import warnings
#from pandas.util._exceptions import find_stack_level


# COMMAND ----------

!pip install python-docx


# COMMAND ----------

#!/databricks/python3/bin/python -m pip install --upgrade pip

# COMMAND ----------

from docx import Document
from docx.shared import Cm, Pt
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# COMMAND ----------

# MAGIC %md
# MAGIC                                               Helper Functions

# COMMAND ----------

def percentage_conversion(value):
    return f'{value:.2%}'
    
def convert_week_to_date_month(value):
    current_year = pd.Timestamp.now().year
    date_values = []
    for i in range(len(value)):
        week_int = int(value[i])
        date = datetime.date(datetime(current_year, 1, 1) + relativedelta(weeks=week_int,weekday=0))
        formated_date = date.strftime('%d %b')
        date_values.append(formated_date)
    return date_values

def week_to_date(week_number):
    current_year = pd.Timestamp.now().year
    # if week_number 
    date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number,days=-7))
    return date.strftime('%d %b')

def month_number_to_name(month_number):
    # if 1 <= month_number <= 12:
    #     return calendar.month_name[month_number]
    # else:
    #     return None 
    check = month_number
    if isinstance(check, int):
        if 1 <= month_number <= 12:
            return calendar.month_name[month_number]
        else:
            return month_number
    else:
        return month_number
    

def week_to_date_running_avg(week_number):
    current_year = pd.Timestamp.now().year
    start_date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number,days=-6))
    end_date = start_date + timedelta(days=-28)
    return f"{end_date.strftime('%d %b')} - {start_date.strftime('%d %b')}"

def format_indian_number(number):
    """
    Format a number in the Indian numbering system without abbreviations and with two decimal places.

    Args:
        number (int or float): The number to format.

    Returns:
        str: The formatted number as a string.
    """
    # Round the number to two decimal places
    formatted_number = "{:,.0f}".format(number)

    # Convert the number to a string with commas as thousands separators
    formatted_number = formatted_number[:-3] + formatted_number[-3:].replace(",", ",")

    return formatted_number

def formatINR(number):
    s, *d = str(number).partition(".")
    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
    return r
  #  return "".join([r] + d)
    
def convert_to_lakhs(number):
    """
    Format a number in the Indian numbering system.

    Args:
        number (int or float): The number to format.

    Returns:
        str: The formatted number as a string.
    """
    # Convert the number to a string with commas as thousands separators
    formatted_number = "{:,.0f}".format(number)

    # Replace commas with the Indian numbering system separators
    formatted_number = formatted_number.replace(",", ",")

    # Replace thousands with lakhs and millions with crores
    formatted_number = formatted_number.replace(",K", "K").replace(",M", "Cr")

    return formatted_number

def week_to_dates(week_number):
    try:
        week_number = int(week_number)  # Ensure week_number is an integer
        current_year = pd.Timestamp.now().year
        # Calculate the first day of the year
        first_day_of_year = datetime(current_year, 1, 1)
        # Calculate the date corresponding to the given week number
        date = first_day_of_year + timedelta(weeks=week_number - 1) - timedelta(days=first_day_of_year.weekday())
        return date.strftime('%d %b')
    except ValueError:
        return week_number 
    
def month_number_to_names(month_number):
    try:
        month_number = int(month_number)
        month_name = datetime(1900, month_number, 1).strftime('%B')
        return month_name
    except ValueError:
        return month_number
    

# COMMAND ----------



def get_starting_ending_dates(current_dt, period, term):
    # Parse the current_dt string to a datetime object
    current_date = datetime.strptime(current_dt, "%Y-%m-%d")
    # In waterfall charts only last two periods are compared
    waterfall_term = 2 
    
    if period == "month":
        # Calculate the starting date by subtracting term months
        check = current_date.month - term
        check2 = current_date.month - waterfall_term
        if check < 0:
            starting_date = current_date.replace(day=1, month=current_date.month - term + 12, year=current_date.year - 1)
            if check2 < 0:
                waterfall_starting_dt = current_date.replace(day=1, month=current_date.month - waterfall_term + 12, year=current_date.year - 1)
            else: 
                waterfall_starting_dt = current_date.replace(day=1, month=current_date.month - waterfall_term)    
        elif check == 0:
            starting_date = current_date.replace(day=1, month=current_date.month - term + 1, year=current_date.year)
            if check2 < 0:
                waterfall_starting_dt = current_date.replace(day=1, month=current_date.month - waterfall_term + 12, year=current_date.year - 1)
            else: 
                waterfall_starting_dt = current_date.replace(day=1, month=current_date.month - waterfall_term)
        else:
            starting_date = current_date.replace(day=1, month=current_date.month - term)
            waterfall_starting_dt = current_date.replace(day=1, month=current_date.month - waterfall_term)    
        
        
        
        # Calculate the ending date as the last day of the month before the starting date
        ending_date = current_date.replace(day=1, month=current_date.month)
    elif period == "week":
        dt = datetime.strptime(current_dt, '%Y-%m-%d')
        current_monday = dt - timedelta(days = dt.weekday())

        # Calculate the starting date by subtracting term weeks (7 days)
        starting_date = (current_monday - timedelta(days=7 * term))
        waterfall_starting_dt = current_monday - timedelta(days=waterfall_term*7)
        ending_date = current_monday
        
    else:
        return None  # Invalid period

    # Format the starting and ending dates in "YYYY-MM-DD" format
    starting_date_str = starting_date.strftime("%Y-%m-%d")
    waterfall_starting_dt = waterfall_starting_dt.strftime("%Y-%m-%d")
    ending_date_str = ending_date.strftime("%Y-%m-%d")

    return starting_date_str, waterfall_starting_dt, ending_date_str


# COMMAND ----------

## TO CHECK THE Starting, waterfall and ending date
# current_dt = date.today().strftime('%Y-%m-%d')
current_dt = date.today().strftime('%Y-%m-%d')
# current_dt = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')
starting , waterfall, ending = get_starting_ending_dates(current_dt, 'week', 4)
print(starting)
print(waterfall)
print(ending)

# COMMAND ----------

def assign_bucket(value):
        if value >= 1.1:
            return 'A. >10%'
        elif value >= 1.05:
            return 'B. 5%-10%'
        elif value >= 1.01:
            return 'C. 1%-5%'
        elif value >= 1:
            return 'D. 0%-1%'
        else:
            return 'E. <0%'

# COMMAND ----------

def create_mhi_cr_bucket(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        
        select DISTINCT {2}(observation_date) as producer_created_term,
            merchant_id,
            modified_pre_magic_cr_percentage,
            magic_cr_percentage, 
            round(magic_cr_percentage*1.00/(modified_pre_magic_cr_percentage*1.00),2) as delta

            from aggregate_ba.merchant_happiness_index_metrics
            where observation_date > '{0}'
            and observation_date <= '{1}'
            and ( pre_magic_cr_percentage is NOT NULL
                OR pre_magic_rto_percentage is NOT NULL
                OR modified_pre_magic_cr_percentage is NOT NULL 
                OR magic_cr_percentage is NOT NULL
                OR modified_pre_magic_rto_percentage is NOT NULL
                OR magic_rto_percentage is NOT NULL
                OR p50_completion_time_seconds is NOT NULL
                OR magic_prefill_percentage is NOT NULL
                OR p50_page_load_time_seconds is NOT NULL
                )

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df[['modified_pre_magic_cr_percentage', 'magic_cr_percentage','delta']] = cx_base_df[['modified_pre_magic_cr_percentage', 'magic_cr_percentage','delta']].astype(float)
    # cx_base_df = cx_base_df.astype('int')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    

    cx_base_df['Bucket'] = cx_base_df['delta'].apply(assign_bucket)
     
    cx_base_df.drop('modified_pre_magic_cr_percentage', axis=1, inplace=True)
    cx_base_df.drop('magic_cr_percentage', axis=1, inplace=True)
    cx_base_df.drop('delta', axis=1, inplace=True)

    grouped_counts = cx_base_df.groupby(['producer_created_term', 'Bucket']).size().reset_index(name='Merchant_Count')

    # Pivot the grouped DataFrame based on 'week', with 'Bucket' as columns and 'Merchant_Count' as values

    pivot_df = grouped_counts.pivot(index='Bucket', columns='producer_created_term', values='Merchant_Count').fillna(0)

    # Calculate the total number of merchants per bucket

    bucket_totals = cx_base_df.groupby('Bucket').size().reset_index(name='Total_Merchants')
    
    # Merge bucket totals with the pivot table
    result_df = pd.merge(bucket_totals, pivot_df, on='Bucket')
    
    for week in result_df.columns[2:]:
        total_merchants = result_df[week].sum()
        result_df[week + '_Proportion'] = (result_df[week] / total_merchants).apply(lambda x: f"{x:.2%}")
        result_df.drop(week, axis=1, inplace=True)

    result_df.drop('Total_Merchants', axis=1, inplace=True)

    # Display the resulting DataFrame
    print(result_df)

# COMMAND ----------

current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months
# current_dt = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')
period = 'week'
term = 3
starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)
# starting_dt = '2023-10-01'
# ending_dt = '2023-10-03'
print(starting_dt)
print(ending_dt)
base_df = create_mhi_cr_bucket(period, starting_dt, ending_dt)
print(base_df)


# COMMAND ----------

# MAGIC %md  
# MAGIC                                                     Metric's Function 

# COMMAND ----------

def create_acquisition_base_df(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
# spark.sql("")
    base_db = sqlContext.sql("""

    WITH magic_merchant_list AS (
        SELECT 
        mid,
        max(segment) as segment
        FROM batch_sheets.magic_merchant_list
        group by 1
    )
    SELECT 
    {2}(payments.created_date) as payments_created_term,
    payments.merchant_id,
    magic_merchant_list.segment as segment,
    (
        CASE
        WHEN lower(payments.method) = 'cod' THEN 'cod'
        ELSE 'non cod'
        END
    ) AS cod_or_noncod,
    COUNT(DISTINCT payments.id) AS payment_attempt,
    COUNT(
        DISTINCT CASE
        WHEN payments.authorized_at IS NOT NULL
        OR lower(payments.method) = 'cod' THEN payments.id
        ELSE null
        END
    ) AS payment_success,

    cast(round(COALESCE(
        SUM(CASE WHEN payments.authorized_at IS NOT NULL
        OR lower(payments.method) = 'cod'
    THEN payments.base_amount*1.00/100.0  ELSE NULL END), 0
    ),0) as INTEGER) AS gmv_authorized,

    COUNT(DISTINCT CASE WHEN payments.method = 'cod' OR payments.authorized_at IS NOT NULL THEN payments.merchant_id ELSE NULL END) AS mtus_cod_non_cod
    
    FROM realtime_hudi_api.payments AS payments
    
    LEFT JOIN realtime_hudi_api.payment_analytics as pa
    on payments.id = pa.payment_id

    inner JOIN realtime_hudi_api.order_meta  AS order_meta ON payments.order_id =order_meta.order_id
    LEFT JOIN magic_merchant_list AS magic_merchant_list ON payments.merchant_id = magic_merchant_list.mid
    where payments.created_date >= '{0}'
    and payments.created_date < '{1}'
    and order_meta.created_date >= '{0}'
    and order_meta.created_date < '{1}'
    and pa.checkout_id in (
        SELECT distinct checkout_id
        from aggregate_pa.magic_checkout_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}') 
        and (library <> 'magic-x' or library is null)
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
    )
    and payments.merchant_id not in  (
       'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')
    and type='one_click_checkout'
    GROUP BY 1,2,3,4
    """.format(starting_dt,ending_dt,sql_period)
    )

    #%sql $base_db
    print('Base df created but not converted to pandas yet')
    base_df = base_db.toPandas()
    # df = spark.sql(base_db)

    # Convert to Pandas DataFrame (for smaller datasets)
    # base_df = df.toPandas()
    print('Base df created and converted to pandas')
    return base_df

     

# COMMAND ----------

def create_cod_funnel(base_df,period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql("""
    
    SELECT 
    
    {2}(cc.producer_created_date) as producer_created_term,
    
    count(DISTINCT case when cc.event_name = 'behav:checkoutcodoptionselected' then cc.checkout_id end) as cod_select,

    count(DISTINCT case when cc.event_name = 'submit' and get_json_object(cc.properties,'$.data.data.method') = 'cod' then cc.checkout_id end) as cod_submit

 
    from aggregate_pa.cx_1cc_events_dump_v1 as cc 
    where producer_created_date >=  date('{0}') 
    and  producer_created_date < date('{1}')
    and (get_json_object(properties,'$.options._.integration_type') <> 'x' or get_json_object(properties,'$.options._.integration_type') is NULL)
    GROUP BY 1
    """.format(starting_dt,ending_dt,sql_period))
    
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.astype('int')

    payment_table = base_df.groupby(['payments_created_term']).agg({
            'payment_attempt': 'sum',
            'payment_success': 'sum',
            
        }).reset_index()

    merged_df = cx_base_df.merge(payment_table, how='inner', left_on='producer_created_term', right_on='payments_created_term')

    merged_df['COD Payment Success'] = (base_df.loc[base_df['cod_or_noncod'] == 'cod'].groupby('payments_created_term')['payment_success'].sum())

    if period=='week':
        merged_df['producer_created_term'] = merged_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        merged_df['producer_created_term']= merged_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    merged_df = merged_df.drop(columns=['payments_created_term'])
    return merged_df

     

# COMMAND ----------

def create_mtu(period, starting_dt, ending_dt, base_df):
    '''

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'

    wtu_db = sqlContext.sql("""
    Select 
    {2}(payments.created_date) as payments_created_term,
    magic_merchant_list.segment,
        COUNT(DISTINCT CASE WHEN payments.method = 'cod' OR payments.authorized_at IS NOT NULL THEN payments.merchant_id ELSE NULL END) AS mtus_cod_non_cod
    FROM realtime_hudi_api.payments  AS payments
    LEFT JOIN batch_sheets.magic_merchant_list AS magic_merchant_list ON payments.merchant_id = magic_merchant_list.mid
    inner JOIN realtime_hudi_api.order_meta  AS order_meta ON payments.order_id =order_meta.order_id
    where payments.created_date >= '{0}'
    and payments.created_date < '{1}'
    and order_meta.created_date >= '{0}'
    and order_meta.created_date < '{1}'
    and magic_merchant_list.segment is not null
    group by 1,2
    ORDER BY 1,2
        """.format(starting_dt,ending_dt, period)) 
    wtu_temp_df = wtu_db.toPandas()
    '''
    #base_df = base_df.rename(columns={'magic_merchant_list.segment':'segment'})
    #wtu_temp_df = base_df.groupby(by=['segment','payments_created_term']).agg({'mtus_cod_non_cod':'sum'}).reset_index()

    
    
    wtu_temp_df = base_df[(base_df['mtus_cod_non_cod']==1)].groupby(by=['segment','payments_created_term']).agg({'merchant_id':'nunique'}).reset_index()

    if period=='week':
        wtu_temp_df['payments_created_term'] = wtu_temp_df['payments_created_term'].apply(week_to_date)
    elif period=='month':
        wtu_temp_df['payments_created_term']= wtu_temp_df['payments_created_term'].apply(month_number_to_name)
    else:
        None

    

    #Pivot the Table
    wtu_df = wtu_temp_df.pivot(index='segment', columns='payments_created_term', values='merchant_id').reset_index()
    


    

    #Adding the totals 
    sum_row = wtu_df.sum().to_frame().T
    sum_row['segment'] = 'Total'

    return pd.concat([wtu_df, sum_row])
#wtu_df_trial = create_mtu('month', starting_dt, ending_dt)

# COMMAND ----------

def create_payment_table(base_df, period, term):
    #orders grouped

    orders_grouped = base_df.groupby(['payments_created_term']).agg({
        'gmv_authorized': 'sum',
        'payment_attempt': 'sum',
        'payment_success': 'sum',
        
    })
    #payment_success for cod
    orders_grouped['COD Payment Success'] = (base_df.loc[base_df['cod_or_noncod'] == 'cod'].groupby('payments_created_term')['payment_success'].sum())

    #payment success for non cod
    orders_grouped['Non COD Payment Success'] = (base_df.loc[base_df['cod_or_noncod'] == 'non cod'].groupby('payments_created_term')['payment_success'].sum())

    #COD%
    orders_grouped['COD%'] = ((orders_grouped['COD Payment Success'] / orders_grouped['payment_success'])).apply(percentage_conversion)

    #NON_COD%
    orders_grouped['Non COD%'] = ((orders_grouped['Non COD Payment Success'] / (orders_grouped['COD Payment Success'] + orders_grouped['Non COD Payment Success']))).apply(percentage_conversion)

    #Non COD SR%
    orders_grouped['Non COD SR%'] = ((orders_grouped['Non COD Payment Success'] / (base_df.loc[base_df['cod_or_noncod'] == 'non cod'].groupby('payments_created_term')['payment_attempt'].sum()))).apply(percentage_conversion)

    # #Overall SR%
    orders_grouped['Overall SR%'] = ((orders_grouped['payment_success'] / orders_grouped['payment_attempt'])).apply(percentage_conversion)


    orders_grouped = orders_grouped.rename(columns={'payment_attempt': 'Payment Attempt','payment_success':'Payment Success','gmv_authorized':'GMV Authorized'})

    if period=='week':
        orders_grouped.index = orders_grouped.index.map(week_to_date)
    elif period=='month':
        orders_grouped.index = orders_grouped.index.map(month_number_to_name)
    else:
        None
    orders_grouped = orders_grouped.transpose()
    orders_grouped = orders_grouped.fillna(0)
    orders_grouped.iloc[:5] = orders_grouped.iloc[:5].astype(int)
    orders_grouped.iloc[:5] = orders_grouped.iloc[:5].applymap(formatINR)
    return orders_grouped.reset_index()


# COMMAND ----------

def create_merchant_df():
    merchant_sql = sqlContext.sql("""
        SELECT merchants.id as merchant_id, max(merchants.website) as website
        FROM realtime_hudi_api.merchants as merchants
        GROUP BY 1
        """)
    merchant_sql = merchant_sql.toPandas()
    return merchant_sql

# COMMAND ----------

def create_orders_waterfall(base_df, merchant_df, waterfall_starting_dt, period):
    
    

    #Merging the two tables
    merchant_and_payment_table_merged = pd.merge(base_df, merchant_df, how='left', on='merchant_id')
    #merchant_and_payment_table_merged.head()

    #converting week/month number to date/month name
    
    waterfall_starting_dt =  datetime.strptime(waterfall_starting_dt, "%Y-%m-%d")
    if period=='week':
        waterfall_starting_dt = week_to_date(waterfall_starting_dt.isocalendar()[1])
    elif period=='month':
        waterfall_starting_dt =  month_number_to_name(waterfall_starting_dt.month)
    else:
        None


    orders_grouped_data = merchant_and_payment_table_merged.groupby(['payments_created_term','merchant_id','website' ])['payment_success'].sum().reset_index()

    print(orders_grouped_data)

    orders_pivot = pd.pivot_table(orders_grouped_data, values='payment_success', columns='payments_created_term', index=['merchant_id','website']).reset_index()

    #orders_grouped_data = orders_grouped_data['payments_created_term'].map(month_number_to_name)
    
    orders_pivot = orders_pivot.fillna(0)
    
    print(orders_pivot)

    last_two_weeks = orders_pivot.columns[-2:]
    orders_pivot['current_period - prev_period'] = orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]]
    orders_pivot['abs( current_period - prev_period )'] = (abs(orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]])).round(2)
    orders_pivot = orders_pivot.sort_values(by="abs( current_period - prev_period )",ascending=False)

    filtered_df = orders_pivot.head(10)
    # Calculate the sum of all rows
    total_sum = orders_pivot['current_period - prev_period'].sum()

    # Calculate the sum of the top 10 rows
    top_sum = orders_pivot['current_period - prev_period'].head(10).sum()

    # Calculate the sum of the rows that are NOT in the top 15 rows
    other_orders_sum = total_sum - top_sum
    previous_week_sum = orders_pivot.iloc[:,-4].sum()
    subtotal_sum = orders_pivot.iloc[:,-3].sum()

    filtered_websites = filtered_df[['website', 'current_period - prev_period']].rename(columns={'current_period - prev_period':'Orders'})

    other_row = pd.DataFrame({'website': ['Others'], 'Orders': [other_orders_sum]})

    previous_week = pd.DataFrame({'website':[f"{waterfall_starting_dt}"],'Orders':[previous_week_sum]})

    subtotal = pd.DataFrame({'website':['subtotal'],'Orders':[subtotal_sum]})

    result_df = pd.concat([filtered_websites, other_row, previous_week, subtotal]).reset_index()

    last_second_row = result_df.iloc[-2]

    result_df = pd.concat([last_second_row.to_frame().T, result_df.drop(result_df.index[-2])])

    previous_week_total = orders_pivot.iloc[:,-4].sum()

    subtotal = orders_pivot.iloc[:,-3].sum()

    merchant_orders_values = result_df['Orders'].values
    
    merchant_website = result_df['website'].astype('str').values

    #base_index = result_df[result_df['website'] == f"{waterfall_starting_dt}"].index[0]
    #base = previous_week_total

    measures =  ["absolute"] + ["relative"] * (len(merchant_orders_values) - 2) + ["total"]

    trace= go.Waterfall(
    x=(merchant_website),
    y=(merchant_orders_values),
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.85,
    #text=list(merchant_orders_values),
    text=result_df['Orders'].apply(lambda x: f'{x:,}').apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    textposition='outside',
    base=50000

    )


    layout = go.Layout(
    title='Orders Merchant Influence',
    title_font=dict(size=1),
    xaxis=dict(title='Website'),
    yaxis=dict(title='Orders'),
    plot_bgcolor='rgba(0,0,0,0)',
    width=1050,
    height=500,
    #grid = True
    )

    orders_fig = go.Figure(data=trace, layout=layout)
    orders_fig = orders_fig.update_xaxes(tickfont=dict(size=30,family='Arial, bold'))
    return orders_fig



# COMMAND ----------

def create_gmv_waterfall(base_df, merchant_df, waterfall_starting_dt, period):
    


    #Merging the two tables
    merchant_and_payment_table_merged = pd.merge(base_df, merchant_df, how='left', on='merchant_id')
    #merchant_and_payment_table_merged.head()

    #converting week/month number to date/month name
    
    waterfall_starting_dt =  datetime.strptime(waterfall_starting_dt, "%Y-%m-%d")
    if period=='week':
        waterfall_starting_dt = week_to_date(waterfall_starting_dt.isocalendar()[1])
    elif period=='month':
        waterfall_starting_dt =  month_number_to_name(waterfall_starting_dt.month)
    else:
        None
    
    print(waterfall_starting_dt) 
    print('waterfall date ^')

    orders_grouped_data = merchant_and_payment_table_merged.groupby(['payments_created_term','merchant_id','website'])['gmv_authorized'].sum().reset_index()
    print(orders_grouped_data['gmv_authorized']  )

    orders_pivot = pd.pivot_table(orders_grouped_data, values='gmv_authorized', columns='payments_created_term', index=['merchant_id','website']).reset_index()
    #orders_grouped_data = orders_grouped_data['payments_created_term'].map(month_number_to_name)
    orders_pivot = orders_pivot.fillna(0)

    last_two_weeks = orders_pivot.columns[-2:]
    orders_pivot['current_period - prev_period'] = (orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]]).round(2)
    orders_pivot['abs( current_period - prev_period )'] = (abs(orders_pivot[last_two_weeks[1]] - orders_pivot[last_two_weeks[0]])).round(2)
    orders_pivot = orders_pivot.sort_values(by="abs( current_period - prev_period )",ascending=False)

    filtered_df = orders_pivot.head(10)
    # Calculate the sum of all rows
    total_sum = orders_pivot['current_period - prev_period'].sum()

    # Calculate the sum of the top 10 rows
    top_sum = orders_pivot['current_period - prev_period'].head(10).sum()

    # Calculate the sum of the rows that are NOT in the top 15 rows
    other_orders_sum = total_sum - top_sum
    previous_week_sum = orders_pivot.iloc[:,-4].sum()
    subtotal_sum = orders_pivot.iloc[:,-3].sum()

    filtered_websites = filtered_df[['website', 'current_period - prev_period']].rename(columns={'current_period - prev_period':'Orders'})
    other_row = pd.DataFrame({'website': ['Others'], 'Orders': [other_orders_sum]})

    previous_week = pd.DataFrame({'website':[f"{waterfall_starting_dt}"],'Orders':[previous_week_sum]})
    
    subtotal = pd.DataFrame({'website':['subtotal'],'Orders':[subtotal_sum]})

    result_df = pd.concat([filtered_websites, other_row, previous_week, subtotal]).reset_index()
    last_second_row = result_df.iloc[-2]
    result_df = pd.concat([last_second_row.to_frame().T, result_df.drop(result_df.index[-2])])

    previous_week_total = orders_pivot.iloc[:,-4].sum()
    subtotal = orders_pivot.iloc[:,-3].sum()

    merchant_orders_values = result_df['Orders'].values
    merchant_website = result_df['website'].astype('str').values

    #base_index = result_df[result_df['website'] == f"{waterfall_starting_dt}"].index[0]
    #base = previous_week_total

    measures =  ["absolute"] + ["relative"] * (len(merchant_orders_values) - 2) + ["total"]

    trace= go.Waterfall(
    x=(merchant_website),
    y=(merchant_orders_values),
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.85,
    #text=list(merchant_orders_values),
    text=result_df['Orders'].apply(lambda x: f'{x:,}').apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    #
    textposition='outside',
    base=50000

    )


    layout = go.Layout(
    title='GMV: Merchant Influence',
    title_font=dict(size=1),
    xaxis=dict(title='Website'),
    yaxis=dict(title='GMV Authorized'),
    plot_bgcolor='rgba(0,0,0,0)',
    width=1050,
    height=500,
    #grid = True
    )

    orders_fig = go.Figure(data=trace, layout=layout)
    orders_fig = orders_fig.update_xaxes(tickfont=dict(size=30,family='Arial, bold'))
    return orders_fig



# COMMAND ----------

def get_contact_prefill_rate(period, starting_dt, ending_dt):
#contact prefill rate
    if period=='week':
            sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    contact_prefill_db = sqlContext.sql(
    """
    WITH non_prefill AS(
        SELECT
        
        DISTINCT 
        checkout_id,
        producer_created_date,
        
    max(cast(CASE
        WHEN (
            CAST(
            coalesce(get_json_object(properties, '$.data.prefill_contact_number'), get_json_object(properties, '$.data.meta.prefill_contact_number'))AS string 
            ) IS NULL
            OR 
          CAST(
            coalesce(get_json_object(properties, '$.data.prefill_contact_number'), get_json_object(properties, '$.data.meta.prefill_contact_number') ) AS string
            ) = ''
        ) THEN 0
        ELSE 1
        END as double)) AS prefill_contact_number,

    max(
      get_json_object(properties, '$.data.meta.v2_result')
    ) AS v2_result,
  
    max(get_json_object(properties,'$.data.meta.v2_eligible')) as v2_eligible,  
  
  max (get_json_object(properties,'$.options._.integration_type')) as checkout_integration_type,
  
    max(case when lower(cast(get_json_object(context,'$.user_agent_parsed.os.family') as string)) ='ios' or lower(browser_name) like '%safari%' then 'not-applicable' else 'applicable' end) as prefill_applicable
    
  FROM
        aggregate_pa.cx_1cc_events_dump_v1
    WHERE
    producer_created_date >= date('{0}')
    and producer_created_date <  date('{1}')
    and (library <> 'magic-x' or library is null)
    and (is_magic_x <> 1 or is_magic_x is NULL)
    group by 1,2
    )

    SELECT

    {2}(producer_created_date) as producer_created_term,


    cast(round(COUNT(DISTINCT case when prefill_contact_number=1 then  non_prefill.checkout_id else null end)*1.0/count(DISTINCT non_prefill.checkout_id),4) * 10000 as INTEGER) as prefill_rate

    FROM
    non_prefill
    where producer_created_date >= date('{0}') and producer_created_date  < date('{1}')
    and prefill_applicable = 'applicable'
    and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
    group by 1
    order by 1 desc 
    
    """.format(starting_dt,ending_dt, sql_period))

    contact_prefill_df = contact_prefill_db.toPandas()

    contact_prefill_df['prefill_rate'] = contact_prefill_df['prefill_rate'] / 10000

    contact_prefill_df['prefill_rate'] = contact_prefill_df['prefill_rate'].apply(percentage_conversion)
    if period=='week':
        contact_prefill_df['producer_created_term'] = contact_prefill_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        contact_prefill_df['producer_created_term']= contact_prefill_df['producer_created_term'].apply(month_number_to_name)
    else:
        None
    return contact_prefill_df

# COMMAND ----------

def get_contact_prefill_rate_new(period, starting_dt, ending_dt):
#contact prefill rate
    if period=='week':
            sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    contact_prefill_db = sqlContext.sql(
    """
    WITH non_prefill AS(
    SELECT
        -- count(distinct checkout_id)
        merchant_id,
        checkout_id,
        producer_created_date,
        prefill_contact_number,
        case when lower(os_brand_family)='ios' then 'iOS' else 'non-iOS' end as os_family,
        case when lower(browser_name) like '%safari%' then 'safari' else 'non-safari' end as browser,
        case when lower(os_brand_family)='ios' or lower(browser_name) like '%safari%' then 'not-applicable' else 'applicable' end as prefill_applicable
        FROM
            aggregate_pa.magic_checkout_fact
        WHERE
            --event_name = 'render:1cc_summary_screen_loaded_completed'
            --AND producer_created_date = date('2023-02-01') 
        producer_created_date >= date('{0}')
        and producer_created_date <  date('{1}')
        and summary_screen_loaded = 1
        )

    SELECT

    {2}(producer_created_date) as producer_created_term,


    cast(round(COUNT(DISTINCT case when prefill_contact_number=1 then  non_prefill.checkout_id else null end)*1.0/count(DISTINCT non_prefill.checkout_id),4) * 10000 as INTEGER) as prefill_rate

    FROM
    non_prefill
    where producer_created_date >= date('{0}') and producer_created_date  < date('{1}')
    and prefill_applicable = 'applicable'
    group by 1
    order by 1 desc 
    
    """.format(starting_dt,ending_dt, sql_period))

    contact_prefill_df = contact_prefill_db.toPandas()

    contact_prefill_df['prefill_rate'] = contact_prefill_df['prefill_rate'] / 10000

    contact_prefill_df['prefill_rate'] = contact_prefill_df['prefill_rate'].apply(percentage_conversion)
    if period=='week':
        contact_prefill_df['producer_created_term'] = contact_prefill_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        contact_prefill_df['producer_created_term']= contact_prefill_df['producer_created_term'].apply(month_number_to_name)
    else:
        None
    return contact_prefill_df

# COMMAND ----------

def get_saved_address_prefill_rate(period,starting_dt,ending_dt):

    if period=='week':
            sql_period = 'weekofyear'
    else:
        sql_period = 'month'

    browser_db = sqlContext.sql(
    """
   
    WITH summary AS
    (
    SELECT  DISTINCT checkout_id,
      
    {2}(producer_created_date) as producer_created_term,
      
    max (get_json_object(properties,'$.options._.integration_type')) as checkout_integration_type,

    max(coalesce( try_cast(get_json_object(properties,'$.data.meta.initial_loggedIn') as boolean),try_cast(get_json_object(properties,'$.data.meta.is_user_pre_logged_In') as boolean ) ) ) as logged_status,
      
    
    max( try_cast(
          get_json_object(
            properties,
            '$.data.meta.hasSavedAddress'
          ) AS boolean
        ) ) AS hasSavedAddress,
        
        max( try_cast(
          get_json_object(
            properties,
            '$.data.meta.initial_hasSavedAddress'
          ) AS boolean
        )
      ) AS initial_hasSavedAddress
      
      
    from aggregate_pa.cx_1cc_events_dump_v1
    where 
      producer_created_date  >= date('{0}') 
                and producer_created_date  < date('{1}') 
                and (library <> 'magic-x' or library is null)
                and (is_magic_x <> 1 or is_magic_x is NULL)
    and checkout_id not in 
            (
                SELECT DISTINCT checkout_id
                from aggregate_pa.cx_1cc_events_dump_v1
                where event_name = 'render:1cc_saved_shipping_address_screen_loaded'
                and producer_created_date  >= date('{0}') 
                and producer_created_date  < date('{1}')     
            
    )
    group by 1,2
    
    ),
    
    saved_add as
    (
    select
      DISTINCT checkout_id,
    {2}(producer_created_date) as producer_created_term

    from aggregate_pa.cx_1cc_events_dump_v1
    where event_name = 'render:1cc_saved_shipping_address_screen_loaded'
    and producer_created_date  >= date('{0}') 
                and producer_created_date  < date('{1}') 
                and (library <> 'magic-x' or library is null)
    group by 1,2
    ),
   
   
   total as 
    (
    SELECT Distinct  checkout_id,

      {2}(producer_created_date) as producer_created_term,
       
      max(cast(CASE
        WHEN (
            CAST(
            coalesce(get_json_object(properties, '$.data.prefill_contact_number'), get_json_object(properties, '$.data.meta.contact_number_entered'))AS string
            ) IS NULL
            OR 
          CAST(
            coalesce(get_json_object(properties, '$.data.prefill_contact_number'), get_json_object(properties, '$.data.meta.contact_number_entered') ) AS string
            ) = ''
        ) THEN 0
        ELSE 1
        END as double)) AS is_contact_number_entered,
      
      max (case when event_name = 'render:complete' then get_json_object(properties,'$.options._.integration_type') end) as checkout_integration_type,

      max(case when event_name = 'behav:contact:fill' and get_json_object(properties, '$.data.valid') = 'true' and ( get_json_object(properties, '$.data.valid') is not NULL or get_json_object(properties, '$.data.valid') <> '') then 1 else 0 end) as is_contact_entered_v1,
      
      max(case when event_name = 'api:truecaller_verification'
and get_json_object(properties,'$.data.success') = 'true' then 1 else 0 end) as is_truecaller_loggedin,
      
      max(coalesce( try_cast(get_json_object(properties,'$.data.meta.initial_loggedIn') as boolean),try_cast(get_json_object(properties,'$.data.meta.is_user_pre_logged_In') as boolean ) ) ) as logged_status
      
    from aggregate_pa.cx_1cc_events_dump_v1
    where producer_created_date  >= date('{0}') 
                and producer_created_date  < date('{1}')
                and (library <> 'magic-x' or library is null)
                and (is_magic_x <> 1 or is_magic_x is NULL)
                
    group by 1,2--,3,4
    )
    
    
    SELECT 
    tt.producer_created_term, 

    count(tt.checkout_id) as total,
    count(ss.checkout_id) as summary_prelogged_in,
    count(saved_add.checkout_id) as saved_Address_screen_loaded
    
    FROM (select * from total
     where (is_contact_number_entered = 1 or is_contact_entered_v1 = 1 or is_truecaller_loggedin = 1 or logged_status = TRUE)
     and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
         ) as tt
    
    LEFT JOIN (select * from summary 
       where logged_status = TRUE and (initial_hasSavedAddress = TRUE or hasSavedAddress = TRUE) 
              )  
              as ss
                ON tt.checkout_id = ss.checkout_id  
    LEFT JOIN saved_add ON tt.checkout_id = saved_add.checkout_id 
    GROUP by 1
    ORDER BY 1 DESC
    """.format(starting_dt,ending_dt, sql_period))


    
    browser_df = browser_db.toPandas()
    if period=='week':
        browser_df['producer_created_term'] = browser_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        browser_df['producer_created_term']= browser_df['producer_created_term'].apply(month_number_to_name)
    else:
        None
    browser_df['Summary Screen'] = (browser_df['summary_prelogged_in']*1.00 / browser_df['total']).apply(percentage_conversion)
    browser_df['Saved Screen'] = (browser_df['saved_Address_screen_loaded']*1.00 / browser_df['total']).apply(percentage_conversion)
    browser_df['Address Prefill Rate'] = ((browser_df['summary_prelogged_in'] + browser_df['saved_Address_screen_loaded'])*1.00 / browser_df['total']).apply(percentage_conversion)
    return browser_df

# COMMAND ----------

def get_saved_address_prefill_rate_new(period,starting_dt,ending_dt):

    if period=='week':
            sql_period = 'weekofyear'
    else:
        sql_period = 'month'

    browser_db = sqlContext.sql(
    """
    WITH summary AS
    (
    SELECT  {2}(producer_created_date) as producer_created_term,
    COUNT(DISTINCT checkout_id) as cid
    from aggregate_pa.magic_checkout_fact
    where summary_screen_loaded = 1
    and lower(initial_loggedin) = 'true'
    and lower(initial_hasSavedAddress) = 'true'
    and producer_created_date  >= date('{0}') 
    and producer_created_date  < date('{1}') 
    and checkout_id not in 
            (
                SELECT DISTINCT checkout_id
                from aggregate_pa.magic_checkout_fact
                where saved_address_screen_loaded = 1
                and producer_created_date  >= date('{0}') 
                and producer_created_date  < date('{1}') 
            
    )
    group by 1
    ),
    saved_add as
    (
    select {2}(producer_created_date) as producer_created_term,
    COUNT(DISTINCT checkout_id) as cid
    from aggregate_pa.magic_checkout_fact
    where saved_address_screen_loaded = 1
    and producer_created_date  >= date('{0}') 
    and producer_created_date  < date('{1}') 
    group by 1
    ),
    total as 
    (
    SELECT {2}(producer_created_date) as producer_created_term, 
    count(distinct checkout_id)  as cid
    from aggregate_pa.magic_checkout_fact
    where summary_screen_loaded = 1
    and producer_created_date  >= date('{0}') 
    and producer_created_date  < date('{1}') 
    group by 1
    )
    
    
    SELECT 
    total.producer_created_term, 
    total.cid as total,
    summary.cid as summary_prelogged_in,
    saved_add.cid as saved_Address_screen_loaded
    FROM total 
    LEFT JOIN summary ON total.producer_created_term = summary.producer_created_term  
    LEFT JOIN saved_add ON total.producer_created_term = saved_add.producer_created_term 

    ORDER BY 1 DESC
    """.format(starting_dt,ending_dt, sql_period))


    
    browser_df = browser_db.toPandas()
    if period=='week':
        browser_df['producer_created_term'] = browser_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        browser_df['producer_created_term']= browser_df['producer_created_term'].apply(month_number_to_name)
    else:
        None
    browser_df['Summary Screen'] = (browser_df['summary_prelogged_in']*1.00 / browser_df['total']).apply(percentage_conversion)
    browser_df['Saved Screen'] = (browser_df['saved_Address_screen_loaded']*1.00 / browser_df['total']).apply(percentage_conversion)
    browser_df['Address Prefill Rate'] = ((browser_df['summary_prelogged_in'] + browser_df['saved_Address_screen_loaded'])*1.00 / browser_df['total']).apply(percentage_conversion)
    return browser_df

# COMMAND ----------

# MAGIC %sql
# MAGIC REFRESH TABLE aggregate_pa.magic_checkout_fact

# COMMAND ----------

def get_combined_cr_table(period,starting_dt,ending_dt, base_df):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        Select  {2}(producer_created_date) as producer_created_term, 
        sum(open) as Checkout_Initiated,
        sum(submit) as Checkout_Submitted
        from aggregate_pa.magic_checkout_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}') 
        and (library <> 'magic-x' or library is null)
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')
        group by 1
        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.astype('int')

    payment_table = base_df.groupby(['payments_created_term']).agg({
            'payment_attempt': 'sum',
            'payment_success': 'sum',
            
        }).reset_index()

    merged_df = cx_base_df.merge(payment_table, how='inner', left_on='producer_created_term', right_on='payments_created_term')
    if period=='week':
        merged_df['producer_created_term'] = merged_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        merged_df['producer_created_term']= merged_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    merged_df = merged_df.drop(columns=['payments_created_term'])
    merged_df['Modal CR'] = (merged_df['Checkout_Submitted']*1.00 / merged_df['Checkout_Initiated']).apply(percentage_conversion)
    merged_df['Payment SR'] = (merged_df['payment_success']*1.00 / merged_df['payment_attempt']).apply(percentage_conversion)
    merged_df['Overall CR'] = (merged_df['payment_success']*1.00 / merged_df['Checkout_Initiated']).apply(percentage_conversion)
    merged_df.iloc[:,1:-3] = merged_df.iloc[:,1:-3].applymap(formatINR)
    merged_df = merged_df.rename(columns={'payment_attempt':'Payment Attempts','payment_success':'Successful Orders'})
    return merged_df

    
    
    


# COMMAND ----------

def get_combined_session_cr_table(period,starting_dt,ending_dt, base_df):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        SELECT
        {2}(producer_created_date) as producer_created_term, 
        count (DISTINCT case when event_name = 'open' then session_id end) as session_Initiated,
        count(DISTINCT case when event_name = 'submit' then session_id end) as session_Submitted
        
        FROM aggregate_pa.cx_1cc_events_dump_v1
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}') 
        and event_name in ('open', 'submit', 'render:complete')
        and (library <> 'magic-x' or library is null)
        --and get_json_object(properties,'$.options._.integration_type') <> 'x'
        and (is_magic_x <> 1 or is_magic_x is NULL)
        
        and checkout_id not in (
            select distinct checkout_id 
            from aggregate_pa.magic_checkout_fact
            where checkout_integration_type = 'x' 
        )
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')
        Group by 1

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.astype('int')

    payment_table = base_df.groupby(['payments_created_term']).agg({
            'payment_attempt': 'sum',
            'payment_success': 'sum',
            
        }).reset_index()

    merged_df = cx_base_df.merge(payment_table, how='inner', left_on='producer_created_term', right_on='payments_created_term')
    if period=='week':
        merged_df['producer_created_term'] = merged_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        merged_df['producer_created_term']= merged_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    merged_df = merged_df.drop(columns=['payments_created_term'])
    merged_df['Modal CR'] = (merged_df['session_Submitted']*1.00 / merged_df['session_Initiated']).apply(percentage_conversion)
    merged_df['Payment SR'] = (merged_df['payment_success']*1.00 / merged_df['payment_attempt']).apply(percentage_conversion)
    merged_df['Overall CR'] = (merged_df['payment_success']*1.00 / merged_df['session_Initiated']).apply(percentage_conversion)
    merged_df.iloc[:,1:-3] = merged_df.iloc[:,1:-3].applymap(formatINR)
    merged_df = merged_df.rename(columns={'payment_attempt':'Payment Attempts','payment_success':'Successful Orders'})
    return merged_df

    
    

    


# COMMAND ----------


def get_cod_intelligence_df(starting_dt,ending_dt, period):
  if period=='week':
    sql_period = 'weekofyear'
  else:
    sql_period = 'month'
  cod_intelligence_db = sqlContext.sql(
      """
          select order_created_date, 
          {2}(order_created_date) as order_created_term,
      result_flag,
      experimentation,
      CASE
    WHEN ml_model_id = 'category_5399' THEN 'Marketplace model'
  WHEN ml_model_id  = 'category_5691' THEN 'Clothing model'
  WHEN ml_model_id  = 'category_5977' THEN 'Cosmetics model'
  WHEN ml_model_id  = 'category_infrequent_mcc_group' THEN 'Generic RTO model'
  WHEN ml_model_id  = 'acm' THEN 'Address model'
    WHEN ml_model_id  = 'merchant_boltt' THEN 'Boltt model'
    ELSE 'Other'
  END
  AS model_id,
  citytier,
      count(distinct order_id) as orders

      from aggregate_pa.magic_rto_reimbursement_fact
      where order_created_date >= date('{0}')
      and order_created_date < date('{1}')
      group by 1,2,3,4,5,6
      order by 1
      """.format(starting_dt,ending_dt, sql_period)
  )
  cod_intelligence_df = cod_intelligence_db.toPandas()

  #Red Flag Rate Calculations
  # Define the conditions
  conditions = [
      (cod_intelligence_df['result_flag'] == "green"),
      (cod_intelligence_df['result_flag'] == "red") & (cod_intelligence_df['experimentation'] == True),
      (cod_intelligence_df['result_flag'] == "red") & (cod_intelligence_df['experimentation'] == False)
  ]

  # Define the values for each condition
  values = ["safe", "safe", "risky"]

  # Create the new column based on the conditions
  cod_intelligence_df['user_facing_risky_flag'] = pd.Series(np.select(conditions, values, default="unknown"))  

  return cod_intelligence_df

# COMMAND ----------

def get_daily_risky_rate(cod_intelligence_temp):
    #Calculating Risky Rate
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_date','user_facing_risky_flag']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby('order_created_date')['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby('order_created_date')['orders'].transform('sum')

    cod_intelligence_risky_rate_df['risky_rate'] = cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']

    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"]
    #cod_intelligence_df

# COMMAND ----------

def get_agg_risky_rate(cod_intelligence_temp, period):
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_term','user_facing_risky_flag']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby('order_created_term')['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby('order_created_term')['orders'].transform('sum')

    cod_intelligence_risky_rate_df['risky_rate'] = (cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']).apply(percentage_conversion)

    if period=='week':
        cod_intelligence_risky_rate_df['order_created_term'] = cod_intelligence_risky_rate_df['order_created_term'].apply(week_to_date)
    elif period=='month':
        cod_intelligence_risky_rate_df['order_created_term']= cod_intelligence_risky_rate_df['order_created_term'].apply(month_number_to_name)
    else:
        None

    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"]
    

# COMMAND ----------


def get_daily_citytier_risky_rate(cod_intelligence_temp):
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_date','user_facing_risky_flag','citytier']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby(['order_created_date','citytier'])['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby(['order_created_date','citytier'])['orders'].transform('sum')

    cod_intelligence_risky_rate_df['risky_rate'] = cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']

    #cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"]
    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].pivot(index='order_created_date', columns='citytier', values=['total','risky_rate']).reset_index()
    #cod_intelligence_df

# COMMAND ----------

def get_agg_citytier_risky_rate(cod_intelligence_temp, period):
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_term','user_facing_risky_flag','citytier']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby(['order_created_term','citytier'])['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby(['order_created_term','citytier'])['orders'].transform('sum')

    cod_intelligence_risky_rate_df['risky_rate'] = (cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']).apply(percentage_conversion)

    if period=='week':
        cod_intelligence_risky_rate_df['order_created_term'] = cod_intelligence_risky_rate_df['order_created_term'].apply(week_to_date)
    elif period=='month':
        cod_intelligence_risky_rate_df['order_created_term']= cod_intelligence_risky_rate_df['order_created_term'].apply(month_number_to_name)

    #cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"]
    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].pivot(index='order_created_term', columns='citytier', values=['total','risky_rate']).reset_index()


# COMMAND ----------


def get_daily_model_risky_rate(cod_intelligence_temp): 
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_date','user_facing_risky_flag','model_id']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby(['order_created_date','model_id'])['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby(['order_created_date','model_id'])['orders'].transform('sum')

    cod_intelligence_risky_rate_df['risky_rate'] = cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']

    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].pivot(index='order_created_date', columns='model_id', values=['total','risky_rate']).reset_index()


# COMMAND ----------

def get_agg_model_risky_rate(cod_intelligence_temp, period): 
    cod_intelligence_risky_rate_df = cod_intelligence_temp.groupby(by=['order_created_term','user_facing_risky_flag','model_id']).agg({'orders':'sum'}).reset_index()

    cod_intelligence_risky_rate_df['risky_orders'] = cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].groupby(['order_created_term','model_id'])['orders'].transform('sum')
    cod_intelligence_risky_rate_df['total'] = round(cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] != "unknown"].groupby(['order_created_term','model_id'])['orders'].transform('sum'),0)

    cod_intelligence_risky_rate_df['risky_rate'] = (cod_intelligence_risky_rate_df['risky_orders']/cod_intelligence_risky_rate_df['total']).apply(percentage_conversion)

    if period=='week':
        cod_intelligence_risky_rate_df['order_created_term'] = cod_intelligence_risky_rate_df['order_created_term'].apply(week_to_date)
    elif period=='month':
        cod_intelligence_risky_rate_df['order_created_term']= cod_intelligence_risky_rate_df['order_created_term'].apply(month_number_to_name)

    #cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"]
    return cod_intelligence_risky_rate_df[cod_intelligence_risky_rate_df['user_facing_risky_flag'] == "risky"].pivot(index='order_created_term', columns='model_id', values='risky_rate').reset_index()


# COMMAND ----------

def get_shipping_data_df(starting_dt,ending_dt, period):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    shipping_data_db = sqlContext.sql(
        """
        select 
        {2}(order_created_date) as order_created_term,
        a.merchant_id,
        case when a.merchant_id='IH7E2OJQGEKKTN' then 'Boltt' else 'Non-Boltt' end as mx_category,
        count(distinct a.order_id) as total_orders,
        count(distinct case when status in ('delivered','rto','lost','cancelled','partially_delivered','returned') then a.order_id else null end) as delivery_status,
        count(distinct case when b.mid is not null then a.order_id else null end) as total_orders_rto_pricing,
        count(distinct case when status in ('delivered','rto','lost','cancelled','partially_delivered','returned') and b.mid is not null then a.order_id else null end) as delivery_status_rto_pricing
        from aggregate_pa.magic_rto_reimbursement_fact a
        left join batch_sheets.magic_checkout_rto_insurance b on a.merchant_id = b.mid
        where order_created_date >= date_sub(date('{0}'), 14)
        and order_created_date < date_sub(date('{1}'), 14)
        and order_status in ('paid','placed')
        and cod_intelligence_enabled = True 
        group by 1,2,3
        """.format(starting_dt,ending_dt, sql_period)
    )
    shipping_data_df = shipping_data_db.toPandas()
    return shipping_data_df


# COMMAND ----------

def get_agg_shipping_data_availability(shipping_data_temp_df, period):
    shipping_data_temp_df = shipping_data_temp_df.drop(columns='merchant_id')
    print(shipping_data_temp_df.columns)
    shipping_data_df = shipping_data_temp_df.groupby(['order_created_term','mx_category',]).agg('sum').reset_index()
    shipping_data_df['shipping_data_availability'] = (shipping_data_df['delivery_status'] / shipping_data_df['total_orders']).apply(percentage_conversion)
    shipping_data_df['shipping_data_rto_pricing'] = (shipping_data_df['delivery_status_rto_pricing'] / shipping_data_df['total_orders_rto_pricing']).apply(percentage_conversion)
    #shipping_data_pivot_df['shipping_data_availability'] = shipping_data_pivot_df['shipping_data_availability']
    #shipping_data_pivot_df['shipping_data_rto_pricing'] = shipping_data_pivot_df['shipping_data_rto_pricing'].apply(percentage_conversion)
    print('lol')
    shipping_data_pivot_df = shipping_data_df.pivot(index='order_created_term', columns='mx_category', values=['total_orders','shipping_data_availability','total_orders_rto_pricing','shipping_data_rto_pricing']).reset_index()

    
    if period=='week':
        shipping_data_pivot_df['order_created_term'] = shipping_data_pivot_df['order_created_term'].apply(week_to_date)
    elif period=='month':
        shipping_data_pivot_df['order_created_term']= shipping_data_pivot_df['order_created_term'].apply(month_number_to_name)
    return shipping_data_pivot_df


# COMMAND ----------

def get_shipping_data_mtu_count(shipping_data_temp_df, period):
    shipping_data_temp_df['shipping_data_rto_pricing'] = (shipping_data_temp_df['delivery_status_rto_pricing'] / shipping_data_temp_df['total_orders_rto_pricing'])
    shipping_data_temp_df['shipping_data_availability'] = (shipping_data_temp_df['delivery_status'] / shipping_data_temp_df['total_orders'])
    shipping_data_temp_df['availability_bins'] = pd.cut(shipping_data_temp_df['shipping_data_availability'],bins=[-0.1,0,0.2,0.5,0.8,0.9,1.0])
    
    shipping_data_temp_df['availability_bins'] = shipping_data_temp_df['availability_bins'].astype('str')
    shipping_data_df = shipping_data_temp_df.groupby(['order_created_term','availability_bins']).agg({'merchant_id':'nunique'}).reset_index()
    print('lol')
    shipping_data_pivot_df = shipping_data_df.pivot(index='order_created_term', columns='availability_bins', values='merchant_id').reset_index()

    if period=='week':
        shipping_data_pivot_df['order_created_term'] = shipping_data_pivot_df['order_created_term'].apply(week_to_date)
    elif period=='month':
        shipping_data_pivot_df['order_created_term']= shipping_data_pivot_df['order_created_term'].apply(month_number_to_name)
    return shipping_data_pivot_df


# COMMAND ----------

def get_address_utilization(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    magic_users_db = sqlContext.sql(
    """
    with checkout_data as (
    select 
    {0}(producer_created_date) as producer_created_term,
   --- checkout_id, 
   --- merchant_id,
    coalesce(
    get_json_object(properties, '$.data.meta.address_id'),
    get_json_object(properties, '$.data.address_id'),
    get_json_object(properties, '$.data.pre_selected_saved_address_id')
    )  as address_id

    from aggregate_pa.cx_1cc_events_dump_v1 
    where producer_created_date >= date('{1}') 
    and producer_created_date < date('{2}')
    and event_name in ('render:1cc_saved_shipping_address_screen_loaded','render:1cc_summary_screen_loaded_completed','behav:1cc_saved_shipping_address_selected')

    )
    select
    case when source_type is null then 'user_saved' else source_type end as address_source,
    producer_created_term,
    count(distinct address_id) as addresses_utilized
    from checkout_data
    left join realtime_hudi_api.addresses a on checkout_data.address_id = a.id
    where address_id is not null
    group by 1,2
    
    ;

    """.format(sql_period,starting_dt,ending_dt)
    )

    magic_users_df = magic_users_db.toPandas()
    

    total_address_db = sqlContext.sql(
    """
    select
        case when source_type is null then 'user_saved' else source_type end as address_source,
        count(distinct id) as total_address_database
        from realtime_hudi_api.addresses 
        where type='shipping_address'
        and created_date < '{2}'
        group by 1
        order by 1 asc
        """.format(sql_period,starting_dt,ending_dt)
    )

    total_address_df = total_address_db.toPandas()


    address_combined_df = total_address_df.merge(magic_users_df, how='left', on='address_source')
    address_combined_df['utilization %'] = address_combined_df['addresses_utilized']*1.0 / address_combined_df['total_address_database']

    address_combined_pivot = address_combined_df.pivot(index='address_source', columns='producer_created_term',values=['addresses_utilized','utilization %']).reset_index()


    return address_combined_pivot

# COMMAND ----------

#get_address_utilization('month', '2023-09-01','2023-11-01')

# COMMAND ----------

def save_plotly_figure_as_image(fig, document,width_inches, height_inches, dpi):
    print(document)
    image_stream = io.BytesIO()
    fig.write_image(image_stream, format='png',width=width_inches * dpi, height=height_inches * dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def create_document_from_multi_level_list(data, date_generated):
    doc = Document()
    for heading in data:
        section_heading = heading[0]
        print(section_heading)
        section_heading_content = heading[1:]
        #print(section_heading_content)
        heading_formating = doc.add_heading(section_heading, level=1)
        heading_run = heading_formating.runs[0]
        heading_run.font.size = Pt(16)  
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_formating.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for section in section_heading_content:
            section_title = section[0]
            print(section_title)
            item = section[1:]
            print(type(item))
            doc.add_heading(section_title, level=1)
            
           
            if isinstance(item, list):
                print(type(item))
                
                subcontent = item[0]
                print(type(subcontent))
                #subcontent = item[1]
                #headings=doc.add_heading(subheading, level=2)
                #headings.paragraph_format.space_before = Pt(8)
                #headings.paragraph_format.space_before = Pt(10)
                if isinstance(subcontent, go.Figure):
                    save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=320)
            
                elif isinstance(subcontent.columns, pd.MultiIndex):
                    paragraph = doc.add_paragraph()
                    #headings.paragraph_format.space_before = Pt(8)
                    #headings.paragraph_format.space_after = Pt(10)
                    table = doc.add_table(subcontent.shape[0]+2, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    last_paragraph = table.rows[-1].cells[0].paragraphs[-1]
                    last_paragraph.paragraph_format.space_after = Pt(10)  

                    
                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(len(col_tuple)==3):
                                    
                                if(j%2 == 0 ):
                                    table.cell(1,i).text = str(col)
                                else:
                                    continue
                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(j%2 != 0):
                                table.cell(0,i).text = str(col)
                            else:
                                continue
                            
                    for i, row_tuple in enumerate(subcontent.itertuples(index=False)):
                        for j, value in enumerate(row_tuple):
                            table.cell(i + 2, j).text = str(value)
                    
                    table_spacing = doc.add_paragraph()
                    table_spacing.paragraph_format.space_before = Pt(8)
                    table_spacing.paragraph_format.space_after = Pt(8)

                elif isinstance(subcontent, pd.DataFrame):
                    doc.add_paragraph()
                    table = doc.add_table(subcontent.shape[0]+1, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    for i, column in enumerate(subcontent.columns):
                        table.cell(0, i).text = str(column)
                        for j, value in enumerate(subcontent[column]):
                            table.cell(j+1, i).text = str(value)
                    
            else:
                raise ValueError("Invalid content format. Expected list.")
    
    doc_title = 'magic_checkout_doc_'+str(date_generated)+'.docx'
    doc.save(doc_title)
    dbutils.fs.cp("file:/databricks/driver/{0}".format(doc_title), "dbfs:/FileStore/shared_transfer/Pallavi_Samodia/{0}".format(doc_title))
    filepath = 'https://razorpay-dev.cloud.databricks.com/files/shared_transfer/Pallavi_Samodia/'+doc_title
    
    return filepath



# COMMAND ----------

def get_shipping_api_availability(period, starting_dt, ending_dt):
    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
        
    shipping_api_db = sqlContext.sql(
        """
        select 
        {0}(producer_created_date) as producer_created_term,
        count(distinct case when pincode_serviceability_successful = 1 then checkout_id else null end)*1.0000/count(distinct case when pincode_serviceability_initiated = 1 then checkout_id else null end) as pincode_api_success_rate
       --- count(distinct case when pincode_serviceability_true = 1 then checkout_id else null end)*1.0000/count(distinct case when pincode_serviceability_successful = 1 then checkout_id else null end) as pincode_delivery_true_rate
        from aggregate_pa.magic_checkout_fact
        where producer_created_date >= date('{1}')
        and producer_created_date < date('{2}')
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and merchant_id not in (
             'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp'
        )
        group by 1
        order by 1
        """.format(sql_period,starting_dt,ending_dt))
    
    shipping_api_df = shipping_api_db.toPandas()

    if period=='week':
        shipping_api_df['producer_created_term'] = shipping_api_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        shipping_api_df['producer_created_term']= shipping_api_df['producer_created_term'].apply(month_number_to_name)

    return shipping_api_df

# COMMAND ----------

def get_payment_availability(period, starting_dt, ending_dt):
    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
        
    payment_availability_db = sqlContext.sql(
        """
        select 
        {0}(producer_created_date) as producer_created_term,
        count(distinct case when get_json_object(properties, '$.data.error.description') is null
         then checkout_id else null end)*1.0000/count(distinct checkout_id) as payment_availability_rate
        from aggregate_pa.cx_1cc_events_dump_v1
        where producer_created_date >= date('{1}')
        and producer_created_date < date('{2}')
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and event_name = 'ajax_response'
        and merchant_id not in ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')
        group by 1
        order by 1
        """.format(sql_period,starting_dt,ending_dt))
    
    payment_availability_df = payment_availability_db.toPandas()

    if period=='week':
        payment_availability_df['producer_created_term'] = payment_availability_df['producer_created_term'].apply(week_to_date)
    elif period=='month':
        payment_availability_df['producer_created_term']= payment_availability_df['producer_created_term'].apply(month_number_to_name)

    return payment_availability_df

# COMMAND ----------

def get_combined_cr_merchant_breakdown(period,starting_dt,ending_dt, base_df):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        Select  {2}(producer_created_date) as producer_created_term, 
        segment,
        sum(open) as Checkout_Initiated,
        sum(submit) as Checkout_Submitted
        from (select * from aggregate_pa.magic_checkout_fact 
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}') 
        and library <> 'magic-x'
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')  ) a
        left join (
            SELECT 
        mid,
        max(segment) as segment
        FROM batch_sheets.magic_merchant_list
        group by 1
        ) b on a.merchant_id = b.mid
        
        group by 1,2
        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    
    cx_base_df['Checkout_Initiated'] = cx_base_df['Checkout_Initiated'].astype('int')
    cx_base_df['Checkout_Submitted'] = cx_base_df['Checkout_Submitted'].astype('int')
    
    payment_table = base_df.groupby(['payments_created_term','segment']).agg({
            'payment_attempt': 'sum',
            'payment_success': 'sum',
            
        }).reset_index()

    merged_df = cx_base_df.merge(payment_table, how='inner', left_on=['producer_created_term','segment'], right_on=['payments_created_term','segment'])

    # merged_df = merged_df.drop(columns=['payments_created_term'])
    # merged_df['Modal CR'] = (merged_df['Checkout_Submitted']*1.00 / merged_df['Checkout_Initiated']).apply(percentage_conversion)
    # merged_df['Payment SR'] = (merged_df['payment_success']*1.00 / merged_df['payment_attempt']).apply(percentage_conversion)
    # merged_df['Overall CR'] = (merged_df['payment_success']*1.00 / merged_df['Checkout_Initiated']).apply(percentage_conversion)
    # merged_df.iloc[:,1:-3] = merged_df.iloc[:,1:-3].applymap(formatINR)
    # merged_df = merged_df.rename(columns={'payment_attempt':'Payment Attempts','payment_success':'Successful Orders'})

    # if period=='week':
    #     merged_df['producer_created_term'] = merged_df['producer_created_term'].map(week_to_date)
    # elif period=='month':
    #     merged_df['producer_created_term']= merged_df['producer_created_term'].map(month_number_to_name)
    # else:
    #     None

    # merged_df = merged_df.pivot(index='producer_created_term', columns='segment', values=['Checkout_Initiated','Modal CR','Overall CR']).reset_index()

    
    return merged_df

# COMMAND ----------

def create_utm_metric_availability (period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
 
        select {2}(producer_created_date) as producer_created_term, 
        case when utm_source is NULL then 'NULL'
        else 'utm_source_found' end as utm_source_availabilty,
        
        COUNT(DISTINCT checkout_id) as cids
        from aggregate_pa.magic_checkout_fact 
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        and library <> 'magic-x'
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp') 
        group by 1,2

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    # cx_base_df = cx_base_df.astype('int')
    
    cx_base_df['utm_source_availabilty'] = cx_base_df['utm_source_availabilty'].fillna('NULL')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None

    sum_cids = cx_base_df.groupby(['producer_created_term', 'utm_source_availabilty'])['cids'].sum().reset_index()

    # Calculate the total cids for each month
    total_cids = cx_base_df.groupby('producer_created_term')['cids'].sum().reset_index()
    total_cids.columns = ['producer_created_term', 'total_cids']

    # Merge sum_cids with total_cids
    merged_df = pd.merge(sum_cids, total_cids, on='producer_created_term')

    # Calculate the desired values
    merged_df['value'] = (merged_df['cids'] / merged_df['total_cids']).apply(percentage_conversion)

    # Pivot the DataFrame to get the desired format
    result_df = merged_df.pivot(index='utm_source_availabilty', columns='producer_created_term', values='value').reset_index()
     
    result_df = result_df.fillna(0)

    # result_df = result_df.drop(columns=['producer_created_term'])

    # Print the result
    print(result_df)
    return result_df


# COMMAND ----------

def create_utm_unavailability_top_mids (period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
 
        select {2}(producer_created_date) as producer_created_term, 
        merchant_id,
        count(DISTINCT case when utm_source is NULL then checkout_id end) as null_cids,  
        COUNT(DISTINCT checkout_id) as total_cids
        
        from aggregate_pa.magic_checkout_fact 
        where producer_created_date >= date('{0}') 
        and producer_created_date  < date('{1}')
        and library <> 'magic-x'
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp') 
        group by 1,2
        

        """.format(starting_dt,ending_dt,sql_period)
    )
    df = base_db.toPandas()
    # cx_base_df = cx_base_df.astype('int')
    
    

    # if period=='week':
    #     df['producer_created_term'] = df['producer_created_term'].map(week_to_date)
    # elif period=='month':
    #     df['producer_created_term']= df['producer_created_term'].map(month_number_to_name)
    # else:
    #     None
    
    # Sort data by week in descending order and then by null_cids in descending order
    df = df.sort_values(by=['producer_created_term', 'null_cids'], ascending=[False, False])

    # Extract the latest producer_created_term
    latest_week = df['producer_created_term'].iloc[0]
    previous_week = df[df['producer_created_term'] != latest_week]['producer_created_term'].iloc[0]

    # Filter top 5 merchants for each week
    top_merchants_latest_week = df[df['producer_created_term'] == latest_week].head(5)
    top_merchants_previous_week = df[df['producer_created_term'] == previous_week].head(5)

    # Calculate sum of null_cids for all merchants for each week
    sum_null_cids_latest_week = df[df['producer_created_term'] == latest_week]['null_cids'].sum()
    sum_null_cids_previous_week = df[df['producer_created_term'] == previous_week]['null_cids'].sum()

    # Calculate the ratios
    top_merchants_latest_week['%null cids'] = (top_merchants_latest_week['null_cids'] / top_merchants_latest_week['total_cids']).apply(percentage_conversion)
    top_merchants_latest_week['%volume share among null'] = (top_merchants_latest_week['null_cids'] / sum_null_cids_latest_week).apply(percentage_conversion)

    top_merchants_previous_week['%null cids'] = (top_merchants_previous_week['null_cids'] / top_merchants_previous_week['total_cids']).apply(percentage_conversion)
    top_merchants_previous_week['%volume share among null'] = (top_merchants_previous_week['null_cids'] / sum_null_cids_previous_week).apply(percentage_conversion)

    # Combine the results
    result_latest = top_merchants_latest_week[['merchant_id', '%null cids', '%volume share among null']].set_index('merchant_id')
    result_previous = top_merchants_previous_week[['merchant_id', '%null cids', '%volume share among null']].set_index('merchant_id')

    # Rename columns to include weeks
    result_latest.columns = [f'{latest_week} %null cids', f'{latest_week} %volume share among null']
    result_previous.columns = [f'{previous_week} %null cids', f'{previous_week} %volume share among null']

    # Merge results on merchant_id
    combined_result = pd.merge(result_latest, result_previous, left_index=True, right_index=True, how='outer').reset_index()

    if period=='week':
        df['producer_created_term'] = df['producer_created_term'].map(week_to_date)
    elif period=='month':
        df['producer_created_term']= df['producer_created_term'].map(month_number_to_name)
    else:
        None
    
    print(combined_result)

    # Reshape the dataframe to have metrics as columns
    # melted = combined_result.melt(id_vars='merchant_id', var_name='metric_week', value_name='value')

    # # Extract week and metric
    # melted[['producer_created_term', 'metric']] = melted['metric_week'].str.extract(r'(\d{4}-\d{2}) (.+)')

    # # Pivot to get metrics as columns under each week
    # final_result = melted.pivot_table(index='merchant_id', columns=['producer_created_term', 'metric'], values='value').reset_index()

    # # Reorder columns to ensure metrics are grouped under each week
    # final_result = final_result.reindex(columns=['merchant_id'] + [(latest_week, '%null cids'), (latest_week, '%volume share among null'), (previous_week, '%null cids'), (previous_week, '%volume share among null')])

    # print(final_result)
    return combined_result


# COMMAND ----------

def create_coupon_mid_metric(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
 
        
        with cte as (Select  
             DISTINCT merchant_id,
        --{2}(producer_created_date) as producer_created_term,
        producer_created_date,
        coupons_available
             
        from aggregate_pa.magic_checkout_fact 
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}') 
        and library <> 'magic-x'
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp') 

        ),
        
    magic_merchant_list AS (
        SELECT 
        mid,
        max(segment) as segment
        FROM batch_sheets.magic_merchant_list
        group by 1
    )
    
    select DISTINCT cte.merchant_id,
    {2}(cte.producer_created_date) as producer_created_term,
    cte.coupons_available
    from cte
    INNER join magic_merchant_list 
    on cte.merchant_id = magic_merchant_list.mid
   
        """.format(starting_dt,ending_dt,sql_period)
    )
    
    cx_base_df = base_db.toPandas()
    # cx_base_df = cx_base_df.astype('int')
    
    # cx_base_df['utm_source_availabilty'] = cx_base_df['utm_source_availabilty'].fillna('NULL')
         
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    
    # Group by both 'producer_created_term' and 'coupons_available' and calculate distinct counts
    distinct_counts = cx_base_df.groupby(['coupons_available', 'producer_created_term'])['merchant_id'].nunique().reset_index()

    # Rename the columns for clarity
    distinct_counts.columns = ['coupons_available', 'producer_created_term', 'distinct_ids_count']

    # Calculate the total distinct merchant_id count per 'producer_created_term'
    total_distinct_ids = cx_base_df.groupby('producer_created_term')['merchant_id'].nunique().reset_index()
    total_distinct_ids.columns = ['producer_created_term', 'total_distinct_ids']

    # Merge the total distinct count back with the distinct_counts DataFrame
    distinct_counts = pd.merge(distinct_counts, total_distinct_ids, on='producer_created_term', how='left')

    # Calculate the ratio and format it as a percentage
    distinct_counts['ratio'] = (distinct_counts['distinct_ids_count'] / distinct_counts['total_distinct_ids']).apply(percentage_conversion)

    # Pivot the DataFrame so that 'producer_created_term' becomes columns
    pivot_df = distinct_counts.pivot(index='coupons_available', columns='producer_created_term', values=['distinct_ids_count', 'ratio'])

    # Flatten the MultiIndex columns
    pivot_df.columns = [f'{metric}_{term}' for metric, term in pivot_df.columns]

    # Reset index to convert it into a regular DataFrame
    pivot_df = pivot_df.reset_index()

    # Print the result
    print(pivot_df)
    
    return pivot_df

    

# COMMAND ----------

def create_coupon_metric_availability (period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
 
        select {2}(producer_created_date) as producer_created_term, 
        coupons_available,
        
        COUNT(DISTINCT checkout_id) as cids
        from aggregate_pa.magic_checkout_fact 
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        and library <> 'magic-x'
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and (is_magic_x <> 1 or is_magic_x is NULL)
        and merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp') 
        group by 1,2

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    # cx_base_df = cx_base_df.astype('int')
    
    cx_base_df['coupons_available'] = cx_base_df['coupons_available'].fillna('NULL')

    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    

    sum_cids = cx_base_df.groupby(['producer_created_term', 'coupons_available'])['cids'].sum().reset_index()

    # Calculate the total cids for each month
    total_cids = cx_base_df.groupby('producer_created_term')['cids'].sum().reset_index()
    total_cids.columns = ['producer_created_term', 'total_cids']

    # Merge sum_cids with total_cids
    merged_df = pd.merge(sum_cids, total_cids, on='producer_created_term')

    # Calculate the desired values
    merged_df['value'] = (merged_df['cids'] / merged_df['total_cids']).apply(percentage_conversion)

    # Pivot the DataFrame to get the desired format
    result_df = merged_df.pivot(index='coupons_available', columns='producer_created_term', values='value').reset_index()

    # Print the result
    print(result_df)
    return result_df
    


# COMMAND ----------

def create_magicx_base_df(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
 
        select {2}(producer_created_date) as producer_created_term, 
        sum(open) as opens,
        sum(submit) as order_created
        from aggregate_pa.magic_x_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        group by 1

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.astype('int')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    
    cx_base_df['Overall CR'] = (cx_base_df['order_created']*1.00 / cx_base_df['opens']).apply(percentage_conversion)
    
    
    return cx_base_df
    


# COMMAND ----------

def create_magicx_gmv(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        select  
        DISTINCT     
        {2}(producer_created_date) as producer_created_term, 
        sum(payment_amount) as GMV
        from aggregate_pa.magic_x_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        group by 1
        
        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.fillna(0)
    cx_base_df = cx_base_df.astype('int')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    cx_base_df['GMV'] = cx_base_df['GMV'].apply(formatINR)
    
    return cx_base_df   


# COMMAND ----------

def create_magicx_gmv_mtu_level(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        select  
        DISTINCT     
        {2}(producer_created_date) as producer_created_term, 
        case when storefront_url in ('https://ihadesigns.in', 'https://ihadesigns.myshopify.com') then 'https://ihadesigns.in'
        else storefront_url end as merchant_id, 
        sum(payment_amount) as GMV
        from aggregate_pa.magic_x_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        group by 1,2
        Having sum(submit) > 0
        
        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.fillna(0)
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    
    cx_base_df['GMV'] = cx_base_df['GMV'].apply(formatINR)
    
    pivot_df = cx_base_df.pivot(index='merchant_id', columns='producer_created_term', values='GMV')

    pivot_df.reset_index(inplace=True)

    return pivot_df


# COMMAND ----------

def create_magicx_mtu(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        select 
        DISTINCT 
        {2}(producer_created_date) as producer_created_term,
        count(distinct case when storefront_url in ('https://ihadesigns.in', 'https://ihadesigns.myshopify.com') then 'https://ihadesigns.in'
        else storefront_url end) as MTUs

        from aggregate_pa.magic_x_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        group by 1
        HAVING sum(submit) > 0

        """.format(starting_dt,ending_dt,sql_period)
    )
    
    cx_base_df = base_db.toPandas()
    cx_base_df = cx_base_df.astype('int')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None
    
    
    return cx_base_df

    
    

    


# COMMAND ----------

def create_magicx_mtu_cr(period, starting_dt, ending_dt):

    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    base_db = sqlContext.sql(
        """
        
        select {2}(producer_created_date) as producer_created_term,
        case when storefront_url in ('https://ihadesigns.in', 'https://ihadesigns.myshopify.com') then 'https://ihadesigns.in'
        else storefront_url end as merchant_id, 
        sum(open) as opens,
        sum(submit) as order_created
        from aggregate_pa.magic_x_fact
        where producer_created_date  >= date('{0}') 
        and producer_created_date  < date('{1}')
        group by 1,2
        HAVING sum(submit) > 0

        """.format(starting_dt,ending_dt,sql_period)
    )
    cx_base_df = base_db.toPandas()
    cx_base_df[['opens', 'order_created']] = cx_base_df[['opens', 'order_created']].astype(int)
    # cx_base_df = cx_base_df.astype('int')
    
    if period=='week':
        cx_base_df['producer_created_term'] = cx_base_df['producer_created_term'].map(week_to_date)
    elif period=='month':
        cx_base_df['producer_created_term']= cx_base_df['producer_created_term'].map(month_number_to_name)
    else:
        None

    

    cx_base_df['Overall CR'] = (cx_base_df['order_created']*1.00 / cx_base_df['opens']).apply(percentage_conversion)
     
    cx_base_df.drop('opens', axis=1, inplace=True)
    cx_base_df.drop('order_created', axis=1, inplace=True)

    pivot_df = cx_base_df.pivot(index='merchant_id', columns='producer_created_term', values='Overall CR')

    # Reset the index if needed
    pivot_df.reset_index(inplace=True)

    return pivot_df
    


# COMMAND ----------

def get_total_addresses(period, starting_dt, ending_dt):
    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    
    total_addresses_db = sqlContext.sql(
        """
        select
        case when created_date < '{1}' then 'Addresses so far' else {0}(created_date) end as created_date_group,
        case when source_type is null then 'user_saved' else source_type end as address_source,
        count(distinct id) as Count_of_Addresses_Added
        from realtime_hudi_api.addresses 
        where type='shipping_address'
        and created_date < '{2}'
        group by 1,2
        order by 1 asc
        """.format(sql_period,starting_dt,ending_dt)
    )
    total_addresses_df = total_addresses_db.toPandas()
    total_addresses_df = total_addresses_df.pivot(index='created_date_group', columns='address_source', values='Count_of_Addresses_Added').reset_index()
    total_addresses_df = total_addresses_df.fillna(0)
    total_addresses_df = total_addresses_df.sort_values(by='created_date_group', ascending=False)
    total_addresses_df['Total Addresses'] = total_addresses_df.iloc[:,1:].sum(axis=1)
    total_addresses_df['Running Total'] = total_addresses_df['Total Addresses'].cumsum()
    

    for col in total_addresses_df.columns[1:]:
        total_addresses_df.loc[:,col] = [f'{x:,.0f}' for x in total_addresses_df.loc[:,col]]
    
    #total_addresses_df.iloc[1:, 1:] = total_addresses_df.iloc[1:, 1:].format(lambda x: "{:,.0f}".format(x))
    #total_addresses_df.iloc[1:, 1:] = total_addresses_df.iloc[1:, 1:].astype(int)
    #total_addresses_df.iloc[1:, 1:] = total_addresses_df.iloc[1:, 1:].apply(formatINR)
    #n = total_addresses_df.shape[1]
    #for i in range(1,n):
     #   total_addresses_df[i] = total_addresses_df[i].map(formatINR)

    if period=='week':
        total_addresses_df['created_date_group'] = total_addresses_df['created_date_group'].apply(week_to_dates)
    elif period=='month':
        total_addresses_df['created_date_group']= total_addresses_df['created_date_group'].apply(month_number_to_names)
    
    return total_addresses_df.T



# COMMAND ----------

current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months
# current_dt = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')
period = 'week'
term = 4
starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)
print(starting_dt)
print(waterfall_starting_dt)
print(ending_dt)
base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
print(base_df.columns)

merchant_sql = create_merchant_df()


# COMMAND ----------


file_path = '/Workspace/Users/manjeet.singh@razorpay.com/AWT 21st Jan.csv'

base_df = pd.read_csv(file_path)

# COMMAND ----------

file_path = '/Workspace/Users/manjeet.singh@razorpay.com/mid_website_apollo.csv'

merchant_sql = pd.read_csv(file_path)

# COMMAND ----------

def get_disabled_reasons(period, ending_dt, churn_starting_dt, merchant_df, payments_df):
    #1. identify merchants where churn happened
    #2. Get their disabled reasons
    if period=='week':
        sql_period = 'weekofyear'
    else:
        sql_period = 'month'
    disabled_reasons_db = sqlContext.sql(
    """
    select mid,
    case when disabled_reasons is null or disabled_reasons = '' then 'Reason not available' else disabled_reasons end as disabled_reasons,
    case when disable_category is null or disable_category = '' then 'Reason not available' else disable_category end as disable_category
    from batch_sheets.magic_merchant_list
    where lower(sales_merchant_status) = 'disabled'
    and segment <> 'SME'
    """
    )

    base_db = sqlContext.sql("""
    WITH magic_merchant_list AS (
        SELECT 
        mid,
        max(segment) as segment
        FROM batch_sheets.magic_merchant_list
        group by 1
    )
    SELECT 
    {2}(payments.created_date) as producer_created_term,
    payments.merchant_id as merchant_id,
    --magic_merchant_list.segment as segment,
    COUNT(DISTINCT payments.id) AS payment_attempt
    
    FROM realtime_hudi_api.payments AS payments
    
    LEFT JOIN realtime_hudi_api.payment_analytics as pa
    on payments.id = pa.payment_id

    inner JOIN realtime_hudi_api.order_meta  AS order_meta ON payments.order_id =order_meta.order_id
    LEFT JOIN magic_merchant_list AS magic_merchant_list ON payments.merchant_id = magic_merchant_list.mid
    where payments.created_date >= '{0}'
    and payments.created_date < '{1}'
    and order_meta.created_date >= '{0}'
    and order_meta.created_date < '{1}'
    and pa.checkout_id in (
        SELECT checkout_id
        FROM aggregate_pa.magic_checkout_fact
        WHERE library <> 'magic-x' 
        and (checkout_integration_type <> 'x' or checkout_integration_type is NULL)
        and  (is_magic_x <> 1 or is_magic_x is NULL)
            AND producer_created_date >= date('{0}') AND producer_created_date <  date('{1}')
    )
    and payments.merchant_id not in  ( 'Jk0cBn86K2qcLW',
'HqVjVuDqQxYp83',
'NGSsnxCaBdvFkZ',
'6nbwElTx3T6BjV',
'LYAFbUjupqdfuZ',
'CEJeWFPkKwBRUI',
'KG5Lj3iYF5byyW',
'FYqoJtceB8PxFv',
'Fn9n3cFtdw3EnS',
'J3NPvUSyXM8xyw',
'G4vxH4Gg4qZxoe',
'JFy2dMQgggcNZS',
'DKM1Z5YYPsx3zD',
'GPXVVnssg0W0q3',
'JM0Ka1zhZzYdb2',
'NoEF8hteiWrSyI',
'F8QKNfMBz98uqa',
'GsPHV1fS27O7Nw',
'IC1De6fM3nNU76',
'4af5pL6Gz4AElE',
'KQ5WA1k7vnf1gH',
'DeNmZIZGXvaN6f',
'EfPwexBImj5oxv',
'J4ZyeBrzHQzgGg',
'I4jUfmyZv18qqa',
'80pX1MMiXMC93y',
'LaQfoWnvltVJlm',
'IgRWARSBxW5BEQ',
'I4jVPuT0AQUqOp')
    and type='one_click_checkout'
    GROUP BY 1,2
    """.format(starting_dt,ending_dt,sql_period))

    payments_grouped_df = payments_df.groupby(by=['payments_created_term','merchant_id',]).agg({'gmv_authorized':'sum','payment_success':'sum'}).reset_index()
    payments_pivot = pd.pivot_table(payments_grouped_df, values=['gmv_authorized'], columns='payments_created_term', index=['merchant_id',]).reset_index()
    payments_pivot = payments_pivot.fillna(0)
    last_two_weeks = payments_pivot.columns[-2:]
    

    payments_pivot['GMV impact'] = payments_pivot[last_two_weeks[1]] - payments_pivot[last_two_weeks[0]]
    
    try: 
        payments_pivot.columns = ['merchant_id',last_two_weeks[1], last_two_weeks[0],'GMV impact']
    except:
        payments_pivot.columns = ['merchant_id','T-2','T-1',last_two_weeks[1],last_two_weeks[0],'GMV impact']    
    # if period=='week':
    #     # payments_pivot.columns = ['merchant_id','T-2','T-1',last_two_weeks[1],last_two_weeks[0],'GMV impact']
    #     payments_pivot.columns = ['merchant_id','T-2',last_two_weeks[1], last_two_weeks[0],'GMV impact']
    # else:
    #     payments_pivot.columns = ['merchant_id','T-2',last_two_weeks[1], last_two_weeks[0],'GMV impact']
    
    # payments_pivot.columns = ['merchant_id','T-2','T-1',last_two_weeks[1],last_two_weeks[0],'GMV impact']

    # payments_pivot.columns = ['merchant_id','T-2',last_two_weeks[1], last_two_weeks[0],'GMV impact']
    
    #print(payments_pivot.shape)
    #print(payments_pivot[payments_pivot['merchant_id']=='BUOPLc8zVITWGR'])
    print('disabled_reasons df created but not converted to pandas yet')
    disabled_reasons_df = disabled_reasons_db.toPandas()
    base_df = base_db.toPandas()
    #disabled_reasons_df.head()
    #Pivoting base table
    base_pivot = base_df.pivot(index='merchant_id', columns='producer_created_term',values= 'payment_attempt').reset_index()
    pre_term = min(base_df['producer_created_term'])
    post_term = max(base_df['producer_created_term'])
    base_pivot = base_pivot.fillna(0)
    churned_mx = base_pivot[(base_pivot[pre_term] != 0) & (base_pivot[post_term] == 0)]
    disabled_reason_merged = pd.merge(churned_mx, disabled_reasons_df, how='left', left_on='merchant_id',  right_on='mid',).reset_index()

    print(disabled_reason_merged)

    disabled_reason_merged['disabled_reasons'] = disabled_reason_merged['disabled_reasons'].replace('', 'Reason not available')
    disabled_reason_merged['disable_category'] = disabled_reason_merged['disable_category'].replace('', 'Reason not available')
    disabled_reason_merged['disabled_reasons'] = disabled_reason_merged['disabled_reasons'].fillna('Reason not available')
    disabled_reason_merged['disable_category'] = disabled_reason_merged['disable_category'].fillna('Reason not available')
    merchant_df = merchant_df.reset_index()
    
    #payments_pivot = payments_pivot.reset_index()
    #print(disabled_reason_merged.columns)
    #print(merchant_df.columns)
    #print(disabled_reason_merged.columns)
    inter_df = pd.merge(disabled_reason_merged, merchant_df, how='left', on='merchant_id').reset_index()
    #print(inter_df.columns)
    print(inter_df)

    merchant_and_payment_table_merged = pd.merge(inter_df, payments_pivot, how='left', on='merchant_id', ).reset_index()
    #print(merchant_and_payment_table_merged.columns)
    disabled_detailed_df = merchant_and_payment_table_merged[merchant_and_payment_table_merged['disabled_reasons'] != 'Reason not available'][['merchant_id','website','disabled_reasons','GMV impact']]
    disabled_category_df = merchant_and_payment_table_merged.groupby(by=['disable_category']).agg({'merchant_id':'nunique','GMV impact':'sum'}).reset_index()
    print (base_df.shape[0])
    print (churned_mx.shape[0])
    print (disabled_detailed_df.head())
    print (disabled_category_df.head())
    print (merchant_and_payment_table_merged.head())
    print('disabled_detailed df created and converted to pandas')
    return disabled_detailed_df, disabled_category_df

# COMMAND ----------

def save_plotly_figure_as_image(fig, document,width_inches, height_inches, dpi):
    print(document)
    image_stream = io.BytesIO()
    fig.write_image(image_stream, format='png',width=width_inches * dpi, height=height_inches * dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def create_document_from_multi_level_list(data, date):
    doc = Document()
    for heading in data:
        section_heading = heading[0]
        print(section_heading)
        section_heading_content = heading[1:]
        #print(section_heading_content)
        heading_formating = doc.add_heading(section_heading, level=1)
        heading_run = heading_formating.runs[0]
        heading_run.font.size = Pt(16)  
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_formating.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for section in section_heading_content:
            section_title = section[0]
            print(section_title)
            item = section[1:]
            print(type(item))
            doc.add_heading(section_title, level=1)
            
           
            if isinstance(item, list):
                subcontent = item[0]
                #subcontent = item[1]
                #headings=doc.add_heading(subheading, level=2)
                #headings.paragraph_format.space_before = Pt(8)
                #headings.paragraph_format.space_before = Pt(10)
                if isinstance(subcontent, go.Figure):
                    save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=320)
            
                elif isinstance(subcontent.columns, pd.MultiIndex):
                    paragraph = doc.add_paragraph()
                    headings.paragraph_format.space_before = Pt(8)
                    headings.paragraph_format.space_after = Pt(10)
                    table = doc.add_table(subcontent.shape[0]+2, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    last_paragraph = table.rows[-1].cells[0].paragraphs[-1]
                    last_paragraph.paragraph_format.space_after = Pt(10)  

                    
                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(len(col_tuple)==3):
                                    
                                if(j%2 == 0 ):
                                    table.cell(1,i).text = str(col)
                                else:
                                    continue
                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(j%2 != 0):
                                table.cell(0,i).text = str(col)
                            else:
                                continue
                            
                    for i, row_tuple in enumerate(subcontent.itertuples(index=False)):
                        for j, value in enumerate(row_tuple):
                            table.cell(i + 2, j).text = str(value)
                    
                    table_spacing = doc.add_paragraph()
                    table_spacing.paragraph_format.space_before = Pt(8)
                    table_spacing.paragraph_format.space_after = Pt(8)

                elif isinstance(subcontent, pd.DataFrame):
                    doc.add_paragraph()
                    table = doc.add_table(subcontent.shape[0]+1, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    for i, column in enumerate(subcontent.columns):
                        table.cell(0, i).text = str(column)
                        for j, value in enumerate(subcontent[column]):
                            table.cell(j+1, i).text = str(value)
                    
            else:
                raise ValueError("Invalid content format. Expected list.")
    doc.save('hm_data_list.docx')
    dbutils.fs.cp("file:/databricks/driver/hm_data_list.docx", "dbfs:/FileStore/shared_transfer/Pallavi_Samodia/hm_data_list.docx")
    # https://razorpay-dev.cloud.databricks.com/files/shared_transfer/Pallavi_Samodia/hm_data_list.docx
    /dbfs/FileStore/
    return doc
#document = create_document_from_multi_level_list(data)


# COMMAND ----------

# def save_plotly_figure_as_image(fig, document,width_inches, height_inches, dpi):
#     print(document)
#     image_stream = io.BytesIO()
#     fig.write_image(image_stream, format='png',width=width_inches * dpi, height=height_inches * dpi)
#     image_stream.seek(0)
#     document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

# def create_document_from_multi_level_list(data, date):
#     doc = Document()
#     for heading in data:
#         section_heading = heading[0]
#         print(section_heading)
#         section_heading_content = heading[1:]
#         #print(section_heading_content)
#         heading_formating = doc.add_heading(section_heading, level=1)
#         heading_run = heading_formating.runs[0]
#         heading_run.font.size = Pt(16)  
#         heading_run.font.color.rgb = RGBColor(0, 0, 0)
#         heading_formating.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         for section in section_heading_content:
#             section_title = section[0]
#             print(section_title)
#             item = section[1:]
#             print(type(item))
#             doc.add_heading(section_title, level=1)
            
           
#             if isinstance(item, list):
#                 subcontent = item[0]
#                 #subcontent = item[1]
#                 #headings=doc.add_heading(subheading, level=2)
#                 #headings.paragraph_format.space_before = Pt(8)
#                 #headings.paragraph_format.space_before = Pt(10)
#                 if isinstance(subcontent, go.Figure):
#                     save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=320)
            
#                 elif isinstance(subcontent.columns, pd.MultiIndex):
#                     paragraph = doc.add_paragraph()
#                     headings.paragraph_format.space_before = Pt(8)
#                     headings.paragraph_format.space_after = Pt(10)
#                     table = doc.add_table(subcontent.shape[0]+2, subcontent.shape[1])
#                     table.style = 'LightGrid-Accent1'
#                     last_paragraph = table.rows[-1].cells[0].paragraphs[-1]
#                     last_paragraph.paragraph_format.space_after = Pt(10)  

                    
#                     for i,col_tuple in enumerate(subcontent.columns):
#                         for j, col in enumerate(col_tuple):
#                             if(len(col_tuple)==3):
                                    
#                                 if(j%2 == 0 ):
#                                     table.cell(1,i).text = str(col)
#                                 else:
#                                     continue
#                     for i,col_tuple in enumerate(subcontent.columns):
#                         for j, col in enumerate(col_tuple):
#                             if(j%2 != 0):
#                                 table.cell(0,i).text = str(col)
#                             else:
#                                 continue
                            
#                     for i, row_tuple in enumerate(subcontent.itertuples(index=False)):
#                         for j, value in enumerate(row_tuple):
#                             table.cell(i + 2, j).text = str(value)
                    
#                     table_spacing = doc.add_paragraph()
#                     table_spacing.paragraph_format.space_before = Pt(8)
#                     table_spacing.paragraph_format.space_after = Pt(8)

#                 elif isinstance(subcontent, pd.DataFrame):
#                     doc.add_paragraph()
#                     table = doc.add_table(subcontent.shape[0]+1, subcontent.shape[1])
#                     table.style = 'LightGrid-Accent1'
#                     for i, column in enumerate(subcontent.columns):
#                         table.cell(0, i).text = str(column)
#                         for j, value in enumerate(subcontent[column]):
#                             table.cell(j+1, i).text = str(value)
                    
#             else:
#                 raise ValueError("Invalid content format. Expected list.")
#     doc.save('hm_data_list.docx')
#     dbutils.fs.cp("/dbfs/FileStore/hm_data_list.docx", "/dbfs/FileStore/hm_data_list.docx")
#     # https://razorpay-dev.cloud.databricks.com/dbfs/FileStore/hm_data_list.docx
    
#     return doc
# #document = create_document_from_multi_level_list(data)


# COMMAND ----------

import io
import plotly.graph_objs as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

def save_plotly_figure_as_image(fig, document, width_inches, height_inches, dpi):
    image_stream = io.BytesIO()
    fig.write_image(image_stream, format='png', width=width_inches * dpi, height=height_inches * dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def create_document_from_multi_level_list(data, date):
    doc = Document()
    for heading in data:
        section_heading = heading[0]
        section_heading_content = heading[1:]
        heading_formatting = doc.add_heading(section_heading, level=1)
        heading_run = heading_formatting.runs[0]
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_formatting.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for section in section_heading_content:
            section_title = section[0]
            item = section[1:]
            doc.add_heading(section_title, level=1)

            if isinstance(item, list):
                subcontent = item[0]
                if isinstance(subcontent, go.Figure):
                    save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=320)
                elif isinstance(subcontent, pd.DataFrame):
                    table = doc.add_table(subcontent.shape[0] + 1, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    for i, column in enumerate(subcontent.columns):
                        table.cell(0, i).text = str(column)
                        for j, value in enumerate(subcontent[column]):
                            table.cell(j + 1, i).text = str(value)
                else:
                    raise ValueError("Invalid content format. Expected list or DataFrame.")
            else:
                raise ValueError("Invalid content format. Expected list.")
    
    # Save the document locally
    local_output_file = '/tmp/hm_data_list.docx'
    doc.save(local_output_file)

    # Copy the file to DBFS
    dbfs_output_file = '/FileStore/hm_data_list.docx'
    dbutils.fs.cp(f"file://{local_output_file}", f"dbfs:{dbfs_output_file}")

    # Return the DBFS path for reference
    return dbfs_output_file

# Example usage:
# dbfs_path = create_document_from_multi_level_list(data, '2025-01-06')
# print(f"Document saved at: {dbfs_path}")


# COMMAND ----------

## DOC CREATION FUNCTION
#Period by default with be Monthly 'month'[Apollo] For Weekly use 'week' [AWT]
#Term is the last number of months that will be compared

def create_doc(period = 'week', term=4):
    current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months
    # current_dt = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')

    starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)

    base_df_3 = create_coupon_mid_metric(period, starting_dt, ending_dt)
    print("create_coupon_mid_metric")
    # # address_prefill_new = get_saved_address_prefill_rate(period, starting_dt, ending_dt)
    # # print("Address Prefill Table Created")
    
    # # contact_prefill_new = get_contact_prefill_rate(period, starting_dt, ending_dt)
    # # print("Contact Prefill Table Created")
    
    

    total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    total_addresses_df = total_addresses_df.reset_index()

    # # # Magic X
    # # # cr_magicx = create_magicx_base_df(period, starting_dt, ending_dt)
    # # # magicx_gmv = create_magicx_gmv(period, starting_dt, ending_dt)
    # # # magicx_mtu_gmv = create_magicx_gmv_mtu_level(period, starting_dt, ending_dt)
    # # # magicx_mtu = create_magicx_mtu(period, starting_dt, ending_dt)
    # # # magicx_mtu_cr = create_magicx_mtu_cr(period, starting_dt, ending_dt)

    # # # Acquisition Tables

    # # # base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
    # # # print(base_df.columns)
    # # # merchant_sql = create_merchant_df()
    # # # print("Merchant DF created")


    # # # disabled_reasons_table, disabled_category_df = get_disabled_reasons(period, ending_dt,waterfall_starting_dt, merchant_sql,base_df)
    # # print("Disabled DF created")
    # cod_funnel = create_cod_funnel(base_df, period, starting_dt, ending_dt)
    mtu = create_mtu(period, starting_dt, ending_dt,base_df)


    print("MTU DF created")
    payment_table = create_payment_table(base_df, period, term)
    print("payment DF created")
    
    # orders_waterfall = create_orders_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("orders_waterfall created")
    # gmv_waterfall = create_gmv_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("gmv_waterfall created")

    
    
   

    combined_cr = get_combined_cr_table(period,starting_dt,ending_dt, base_df)
    print("Combined CR table created")
    # combined_session_cr = get_combined_session_cr_table(period,starting_dt,ending_dt, base_df)
    print("Combined Session CR table created")
    cod_intelligence_df = get_cod_intelligence_df(starting_dt,ending_dt, period)
    print("COD Intelligence table created")
    risky_rate_df = get_agg_risky_rate(cod_intelligence_df, period)
    print("Risky rate table created")
    citytier_risky_rate_df = get_agg_citytier_risky_rate(cod_intelligence_df, period)
    print("Citytier risky rate table created")
    model_risky_rate_df = get_agg_model_risky_rate(cod_intelligence_df, period)
    print("Model risky rate table created")
    shipping_data_df = get_shipping_data_df(starting_dt,ending_dt, period)
    print("Shipping data table created")
    shipping_data_orders_df = get_agg_shipping_data_availability(shipping_data_df, period)
    print("Shipping data % table created")
    get_shipping_data_mtu_count_df = get_shipping_data_mtu_count(shipping_data_df,period)
    print("Shipping data MTU table created")
    # # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # # print("Total addresses table created")
    payment_availability = get_payment_availability(period, starting_dt, ending_dt)
    print("payment availability table created")
    shipping_api_availability = get_shipping_api_availability(period, starting_dt, ending_dt)
    print("shipping api availability table created")
    mid_segment_level_CR = get_combined_cr_merchant_breakdown(period,starting_dt,ending_dt,base_df)
    # print("MID Segment level CR table created")
    # # mhi_cr = create_mhi_cr_bucket(period, starting_dt, ending_dt)

    base_df_1 = create_utm_unavailability_top_mids(period, starting_dt, ending_dt)
    base_df_2 = create_utm_metric_availability(period, starting_dt, ending_dt)

    
    base_df_4 = create_coupon_metric_availability(period, waterfall_starting_dt, ending_dt)

    visualization_list = [
        ['Acquisition',
    
           
            
            ['Orders', payment_table],
            ['MTU', mtu],
            # ['Disabled Reasons by Category', disabled_category_df],
            # ['Disabled Reasons Detailed', disabled_reasons_table],
            # ['Orders: Merchant Waterfall',orders_waterfall],
            # ['GMV: Merchant Waterfall',gmv_waterfall],
        # #     # ['MHI: CR% Change Bucket',mhi_cr],
       
        ],
        ['Consumer Experience',

            ['Combined CR Table', combined_cr],
            # ['Combined Session CR Table', combined_session_cr],
            # ['Contact Prefill Rate', contact_prefill_new],
            # ['Address Prefill Rate New', address_prefill_new],  
        ['MID Segment level CR Table', mid_segment_level_CR],
     
     
        ],

        # ['Magic-X Metrics',

        #     ['Combined Magic-X Session CR Table', cr_magicx],
        #     ['Magic-X GMV', magicx_gmv], 
        #     ['Magic-X MTU level GMV', magicx_mtu_gmv], 
        #     ['Magic-X MTUs', magicx_mtu], 
        #     ['Magic-X MTU level CR', magicx_mtu_cr], 
        
        # ],


        ['Availability Metrics',
            ['Payment availability Table', payment_availability],
            ['Shipping Api Availability Table', shipping_api_availability],
            ['utm_unavailability_top_mids', base_df_1], 
            ['utm_metric_availability', base_df_2],  
            ['coupon_mid_metric', base_df_3],  
            ['coupon_metric_availability', base_df_4],  

        ],

        ['COD Intelligence',
            ['Risky Rate', risky_rate_df],  
            ['Citytier Risky Rate', citytier_risky_rate_df],  
            ['Model Risky Rate', model_risky_rate_df],  
            ['Shipping Data Availability - Orders', shipping_data_orders_df],  
            ['Shipping Data Availability - MTU', get_shipping_data_mtu_count_df], 
        #     # ['Address Prefill Rate New', address_prefill_new],  
            ['Address Count in Database', total_addresses_df],   
            # ['COD Funnel', cod_funnel],
        
        ],
      
        
    ]
    
    return create_document_from_multi_level_list(visualization_list, current_dt)
    


# COMMAND ----------

create_doc('week',4)

# COMMAND ----------

# https://razorpay-dev.cloud.databricks.com/files/hm_data_list.docx



# COMMAND ----------

## DOC CREATION FUNCTION
#Period by default with be Monthly 'month'[Apollo] For Weekly use 'week' [AWT]
#Term is the last number of months that will be compared

def create_doc(period = 'week', term=4):
    current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months

    starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)

    # base_df_3 = create_coupon_mid_metric(period, starting_dt, ending_dt)
    # print("create_coupon_mid_metric")
    # address_prefill_new = get_saved_address_prefill_rate(period, starting_dt, ending_dt)
    # print("Address Prefill Table Created")
    
    # contact_prefill_new = get_contact_prefill_rate(period, starting_dt, ending_dt)
    # print("Contact Prefill Table Created")
    
    

    # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # total_addresses_df = total_addresses_df.reset_index()

    # # Magic X
    # # cr_magicx = create_magicx_base_df(period, starting_dt, ending_dt)
    # # magicx_gmv = create_magicx_gmv(period, starting_dt, ending_dt)
    # # magicx_mtu_gmv = create_magicx_gmv_mtu_level(period, starting_dt, ending_dt)
    # # magicx_mtu = create_magicx_mtu(period, starting_dt, ending_dt)
    # # magicx_mtu_cr = create_magicx_mtu_cr(period, starting_dt, ending_dt)

    # # Acquisition Tables

    # # base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
    # # print(base_df.columns)
    merchant_sql = create_merchant_df()
    # # print("Merchant DF created")


    
    disabled_reasons_table, disabled_category_df = get_disabled_reasons(period, ending_dt,waterfall_starting_dt, merchant_sql,base_df)
    # print("Disabled DF created")
    # cod_funnel = create_cod_funnel(base_df, period, starting_dt, ending_dt)
    # mtu = create_mtu(period, starting_dt, ending_dt,base_df)


    # print("MTU DF created")
    # payment_table = create_payment_table(base_df, period, term)
    # print("payment DF created")
    
    # orders_waterfall = create_orders_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("orders_waterfall created")
    # gmv_waterfall = create_gmv_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("gmv_waterfall created")

    
    
   

    # combined_cr = get_combined_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined CR table created")
    # combined_session_cr = get_combined_session_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined Session CR table created")
    # cod_intelligence_df = get_cod_intelligence_df(starting_dt,ending_dt, period)
    # print("COD Intelligence table created")
    # risky_rate_df = get_agg_risky_rate(cod_intelligence_df, period)
    # print("Risky rate table created")
    # citytier_risky_rate_df = get_agg_citytier_risky_rate(cod_intelligence_df, period)
    # print("Citytier risky rate table created")
    # model_risky_rate_df = get_agg_model_risky_rate(cod_intelligence_df, period)
    # print("Model risky rate table created")
    # shipping_data_df = get_shipping_data_df(starting_dt,ending_dt, period)
    # print("Shipping data table created")
    # shipping_data_orders_df = get_agg_shipping_data_availability(shipping_data_df, period)
    # print("Shipping data % table created")
    # get_shipping_data_mtu_count_df = get_shipping_data_mtu_count(shipping_data_df,period)
    # print("Shipping data MTU table created")
    # # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # # print("Total addresses table created")
    # payment_availability = get_payment_availability(period, starting_dt, ending_dt)
    # print("payment availability table created")
    # shipping_api_availability = get_shipping_api_availability(period, starting_dt, ending_dt)
    # print("shipping api availability table created")
    # mid_segment_level_CR = get_combined_cr_merchant_breakdown(period,starting_dt,ending_dt,base_df)
    # print("MID Segment level CR table created")
    # # mhi_cr = create_mhi_cr_bucket(period, starting_dt, ending_dt)

    # base_df_1 = create_utm_unavailability_top_mids(period, starting_dt, ending_dt)
    # base_df_2 = create_utm_metric_availability(period, starting_dt, ending_dt)

    
    # base_df_4 = create_coupon_metric_availability(period, waterfall_starting_dt, ending_dt)

    visualization_list = [
        # ['Acquisition',
    
           
            
        #     ['Orders', payment_table],
        #     ['MTU', mtu],
        #     # ['Disabled Reasons by Category', disabled_category_df],
        #     # ['Disabled Reasons Detailed', disabled_reasons_table],
        #     ['Orders: Merchant Waterfall',orders_waterfall],
        #     ['GMV: Merchant Waterfall',gmv_waterfall],
        # #     # ['MHI: CR% Change Bucket',mhi_cr],
       
        # ],
        ['Consumer Experience',

            # ['Combined CR Table', combined_cr],
            # ['Combined Session CR Table', combined_session_cr],
            # ['Contact Prefill Rate', contact_prefill_new],
            # ['Address Prefill Rate New', address_prefill_new],  
            ['Disabled Reasons by Category', disabled_category_df],
            ['Disabled Reasons Detailed', disabled_reasons_table],
            # ['MID Segment level CR Table', mid_segment_level_CR],
     
     
        ],

        # ['Magic-X Metrics',

        #     ['Combined Magic-X Session CR Table', cr_magicx],
        #     ['Magic-X GMV', magicx_gmv], 
        #     ['Magic-X MTU level GMV', magicx_mtu_gmv], 
        #     ['Magic-X MTUs', magicx_mtu], 
        #     ['Magic-X MTU level CR', magicx_mtu_cr], 
        
        # ],


        # ['Availability Metrics',
        #     ['Payment availability Table', payment_availability],
        #     ['Shipping Api Availability Table', shipping_api_availability],
        #     ['utm_unavailability_top_mids', base_df_1], 
        #     ['utm_metric_availability', base_df_2],  
        #     ['coupon_mid_metric', base_df_3],  
        #     ['coupon_metric_availability', base_df_4],  

        # ],

        # ['COD Intelligence',
        #     ['Risky Rate', risky_rate_df],  
        #     ['Citytier Risky Rate', citytier_risky_rate_df],  
        #     ['Model Risky Rate', model_risky_rate_df],  
        #     ['Shipping Data Availability - Orders', shipping_data_orders_df],  
        #     ['Shipping Data Availability - MTU', get_shipping_data_mtu_count_df], 
        #     # ['Address Prefill Rate New', address_prefill_new],  
        #     ['Address Count in Database', total_addresses_df],   
        #     ['COD Funnel', cod_funnel],
        
        # ],
      
        
    ]
    
    return create_document_from_multi_level_list(visualization_list, current_dt)
    


# COMMAND ----------

create_doc('month',2)

# COMMAND ----------

## DOC CREATION FUNCTION
#Period by default with be Monthly 'month'[Apollo] For Weekly use 'week' [AWT]
#Term is the last number of months that will be compared

def create_doc(period = 'week', term=4):
    current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months

    starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)

    # base_df_3 = create_coupon_mid_metric(period, starting_dt, ending_dt)
    # print("create_coupon_mid_metric")
    address_prefill_new = get_saved_address_prefill_rate(period, starting_dt, ending_dt)
    print("Address Prefill Table Created")
    
    # contact_prefill_new = get_contact_prefill_rate(period, starting_dt, ending_dt)
    # print("Contact Prefill Table Created")
    
    

    # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # total_addresses_df = total_addresses_df.reset_index()

    # # Magic X
    # # cr_magicx = create_magicx_base_df(period, starting_dt, ending_dt)
    # # magicx_gmv = create_magicx_gmv(period, starting_dt, ending_dt)
    # # magicx_mtu_gmv = create_magicx_gmv_mtu_level(period, starting_dt, ending_dt)
    # # magicx_mtu = create_magicx_mtu(period, starting_dt, ending_dt)
    # # magicx_mtu_cr = create_magicx_mtu_cr(period, starting_dt, ending_dt)

    # # Acquisition Tables

    # # base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
    # # print(base_df.columns)
    # # merchant_sql = create_merchant_df()
    # # print("Merchant DF created")


    # # disabled_reasons_table, disabled_category_df = get_disabled_reasons(period, ending_dt,waterfall_starting_dt, merchant_sql,base_df)
    # print("Disabled DF created")
    # cod_funnel = create_cod_funnel(base_df, period, starting_dt, ending_dt)
    # mtu = create_mtu(period, starting_dt, ending_dt,base_df)


    # print("MTU DF created")
    # payment_table = create_payment_table(base_df, period, term)
    # print("payment DF created")
    
    # orders_waterfall = create_orders_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("orders_waterfall created")
    # gmv_waterfall = create_gmv_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("gmv_waterfall created")

    
    
   

    # combined_cr = get_combined_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined CR table created")
    # combined_session_cr = get_combined_session_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined Session CR table created")
    # cod_intelligence_df = get_cod_intelligence_df(starting_dt,ending_dt, period)
    # print("COD Intelligence table created")
    # risky_rate_df = get_agg_risky_rate(cod_intelligence_df, period)
    # print("Risky rate table created")
    # citytier_risky_rate_df = get_agg_citytier_risky_rate(cod_intelligence_df, period)
    # print("Citytier risky rate table created")
    # model_risky_rate_df = get_agg_model_risky_rate(cod_intelligence_df, period)
    # print("Model risky rate table created")
    # shipping_data_df = get_shipping_data_df(starting_dt,ending_dt, period)
    # print("Shipping data table created")
    # shipping_data_orders_df = get_agg_shipping_data_availability(shipping_data_df, period)
    # print("Shipping data % table created")
    # get_shipping_data_mtu_count_df = get_shipping_data_mtu_count(shipping_data_df,period)
    # print("Shipping data MTU table created")
    # # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # # print("Total addresses table created")
    # payment_availability = get_payment_availability(period, starting_dt, ending_dt)
    # print("payment availability table created")
    # shipping_api_availability = get_shipping_api_availability(period, starting_dt, ending_dt)
    # print("shipping api availability table created")
    # mid_segment_level_CR = get_combined_cr_merchant_breakdown(period,starting_dt,ending_dt,base_df)
    # print("MID Segment level CR table created")
    # # mhi_cr = create_mhi_cr_bucket(period, starting_dt, ending_dt)

    # base_df_1 = create_utm_unavailability_top_mids(period, starting_dt, ending_dt)
    # base_df_2 = create_utm_metric_availability(period, starting_dt, ending_dt)

    
    # base_df_4 = create_coupon_metric_availability(period, waterfall_starting_dt, ending_dt)

    visualization_list = [
        # ['Acquisition',
    
           
            
        #     ['Orders', payment_table],
        #     ['MTU', mtu],
        #     # ['Disabled Reasons by Category', disabled_category_df],
        #     # ['Disabled Reasons Detailed', disabled_reasons_table],
        #     ['Orders: Merchant Waterfall',orders_waterfall],
        #     ['GMV: Merchant Waterfall',gmv_waterfall],
        # #     # ['MHI: CR% Change Bucket',mhi_cr],
       
        # ],
        ['Consumer Experience',

            # ['Combined CR Table', combined_cr],
            # ['Combined Session CR Table', combined_session_cr],
            # ['Contact Prefill Rate', contact_prefill_new],
            ['Address Prefill Rate New', address_prefill_new],  
            # ['Disabled Reasons by Category', disabled_category_df],
            # ['Disabled Reasons Detailed', disabled_reasons_table],
            # ['MID Segment level CR Table', mid_segment_level_CR],
     
     
        ],

        # ['Magic-X Metrics',

        #     ['Combined Magic-X Session CR Table', cr_magicx],
        #     ['Magic-X GMV', magicx_gmv], 
        #     ['Magic-X MTU level GMV', magicx_mtu_gmv], 
        #     ['Magic-X MTUs', magicx_mtu], 
        #     ['Magic-X MTU level CR', magicx_mtu_cr], 
        
        # ],


        # ['Availability Metrics',
        #     ['Payment availability Table', payment_availability],
        #     ['Shipping Api Availability Table', shipping_api_availability],
        #     ['utm_unavailability_top_mids', base_df_1], 
        #     ['utm_metric_availability', base_df_2],  
        #     ['coupon_mid_metric', base_df_3],  
        #     ['coupon_metric_availability', base_df_4],  

        # ],

        # ['COD Intelligence',
        #     ['Risky Rate', risky_rate_df],  
        #     ['Citytier Risky Rate', citytier_risky_rate_df],  
        #     ['Model Risky Rate', model_risky_rate_df],  
        #     ['Shipping Data Availability - Orders', shipping_data_orders_df],  
        #     ['Shipping Data Availability - MTU', get_shipping_data_mtu_count_df], 
        #     # ['Address Prefill Rate New', address_prefill_new],  
        #     ['Address Count in Database', total_addresses_df],   
        #     ['COD Funnel', cod_funnel],
        
        # ],
      
        
    ]
    
    return create_document_from_multi_level_list(visualization_list, current_dt)
    


# COMMAND ----------

# MAGIC %sql
# MAGIC REFRESH TABLE aggregate_pa.magic_checkout_fact

# COMMAND ----------

# MAGIC %sql
# MAGIC REFRESH TABLE aggregate_pa.cx_1cc_events_dump_v1

# COMMAND ----------

create_doc('week',4) 

# COMMAND ----------

## DOC CREATION FUNCTION
#Period by default with be Monthly 'month'[Apollo] For Weekly use 'week' [AWT]
#Term is the last number of months that will be compared

def create_doc(period = 'week', term=4):
    current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months

    starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)

    # base_df_3 = create_coupon_mid_metric(period, starting_dt, ending_dt)
    # print("create_coupon_mid_metric")
    # address_prefill_new = get_saved_address_prefill_rate(period, starting_dt, ending_dt)
    # print("Address Prefill Table Created")
    
    contact_prefill_new = get_contact_prefill_rate(period, starting_dt, ending_dt)
    print("Contact Prefill Table Created")
    
    

    # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # total_addresses_df = total_addresses_df.reset_index()

    # # Magic X
    # # cr_magicx = create_magicx_base_df(period, starting_dt, ending_dt)
    # # magicx_gmv = create_magicx_gmv(period, starting_dt, ending_dt)
    # # magicx_mtu_gmv = create_magicx_gmv_mtu_level(period, starting_dt, ending_dt)
    # # magicx_mtu = create_magicx_mtu(period, starting_dt, ending_dt)
    # # magicx_mtu_cr = create_magicx_mtu_cr(period, starting_dt, ending_dt)

    # # Acquisition Tables

    # # base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
    # # print(base_df.columns)
    # # merchant_sql = create_merchant_df()
    # # print("Merchant DF created")


    # # disabled_reasons_table, disabled_category_df = get_disabled_reasons(period, ending_dt,waterfall_starting_dt, merchant_sql,base_df)
    # print("Disabled DF created")
    # cod_funnel = create_cod_funnel(base_df, period, starting_dt, ending_dt)
    # mtu = create_mtu(period, starting_dt, ending_dt,base_df)


    # print("MTU DF created")
    # payment_table = create_payment_table(base_df, period, term)
    # print("payment DF created")
    
    # orders_waterfall = create_orders_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("orders_waterfall created")
    # gmv_waterfall = create_gmv_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
    # print("gmv_waterfall created")

    
    
   

    # combined_cr = get_combined_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined CR table created")
    # combined_session_cr = get_combined_session_cr_table(period,starting_dt,ending_dt, base_df)
    # print("Combined Session CR table created")
    # cod_intelligence_df = get_cod_intelligence_df(starting_dt,ending_dt, period)
    # print("COD Intelligence table created")
    # risky_rate_df = get_agg_risky_rate(cod_intelligence_df, period)
    # print("Risky rate table created")
    # citytier_risky_rate_df = get_agg_citytier_risky_rate(cod_intelligence_df, period)
    # print("Citytier risky rate table created")
    # model_risky_rate_df = get_agg_model_risky_rate(cod_intelligence_df, period)
    # print("Model risky rate table created")
    # shipping_data_df = get_shipping_data_df(starting_dt,ending_dt, period)
    # print("Shipping data table created")
    # shipping_data_orders_df = get_agg_shipping_data_availability(shipping_data_df, period)
    # print("Shipping data % table created")
    # get_shipping_data_mtu_count_df = get_shipping_data_mtu_count(shipping_data_df,period)
    # print("Shipping data MTU table created")
    # # total_addresses_df = get_total_addresses(period, starting_dt, ending_dt)
    # # print("Total addresses table created")
    # payment_availability = get_payment_availability(period, starting_dt, ending_dt)
    # print("payment availability table created")
    # shipping_api_availability = get_shipping_api_availability(period, starting_dt, ending_dt)
    # print("shipping api availability table created")
    # mid_segment_level_CR = get_combined_cr_merchant_breakdown(period,starting_dt,ending_dt,base_df)
    # print("MID Segment level CR table created")
    # # mhi_cr = create_mhi_cr_bucket(period, starting_dt, ending_dt)

    # base_df_1 = create_utm_unavailability_top_mids(period, starting_dt, ending_dt)
    # base_df_2 = create_utm_metric_availability(period, starting_dt, ending_dt)

    
    # base_df_4 = create_coupon_metric_availability(period, waterfall_starting_dt, ending_dt)

    visualization_list = [
        # ['Acquisition',
    
           
            
        #     ['Orders', payment_table],
        #     ['MTU', mtu],
        #     # ['Disabled Reasons by Category', disabled_category_df],
        #     # ['Disabled Reasons Detailed', disabled_reasons_table],
        #     ['Orders: Merchant Waterfall',orders_waterfall],
        #     ['GMV: Merchant Waterfall',gmv_waterfall],
        # #     # ['MHI: CR% Change Bucket',mhi_cr],
       
        # ],
        ['Consumer Experience',

            # ['Combined CR Table', combined_cr],
            # ['Combined Session CR Table', combined_session_cr],
            ['Contact Prefill Rate', contact_prefill_new],
            # ['Address Prefill Rate New', address_prefill_new],  
            # ['Disabled Reasons by Category', disabled_category_df],
            # ['Disabled Reasons Detailed', disabled_reasons_table],
            # ['MID Segment level CR Table', mid_segment_level_CR],
     
     
        ],

        # ['Magic-X Metrics',

        #     ['Combined Magic-X Session CR Table', cr_magicx],
        #     ['Magic-X GMV', magicx_gmv], 
        #     ['Magic-X MTU level GMV', magicx_mtu_gmv], 
        #     ['Magic-X MTUs', magicx_mtu], 
        #     ['Magic-X MTU level CR', magicx_mtu_cr], 
        
        # ],


        # ['Availability Metrics',
        #     ['Payment availability Table', payment_availability],
        #     ['Shipping Api Availability Table', shipping_api_availability],
        #     ['utm_unavailability_top_mids', base_df_1], 
        #     ['utm_metric_availability', base_df_2],  
        #     ['coupon_mid_metric', base_df_3],  
        #     ['coupon_metric_availability', base_df_4],  

        # ],

        # ['COD Intelligence',
        #     ['Risky Rate', risky_rate_df],  
        #     ['Citytier Risky Rate', citytier_risky_rate_df],  
        #     ['Model Risky Rate', model_risky_rate_df],  
        #     ['Shipping Data Availability - Orders', shipping_data_orders_df],  
        #     ['Shipping Data Availability - MTU', get_shipping_data_mtu_count_df], 
        #     # ['Address Prefill Rate New', address_prefill_new],  
        #     ['Address Count in Database', total_addresses_df],   
        #     ['COD Funnel', cod_funnel],
        
        # ],
      
        
    ]
    
    return create_document_from_multi_level_list(visualization_list, current_dt)
    


# COMMAND ----------

create_doc('week',4)

# COMMAND ----------

disabled_reasons_table, disabled_category_df = get_disabled_reasons(period, ending_dt,waterfall_starting_dt, merchant_sql,base_df)

# COMMAND ----------

total_addresses_df = total_addresses_df.reset_index()
print(total_addresses_df)
total_addresses_df.columns

# COMMAND ----------

¿?

# COMMAND ----------

# Orders 
get std_ecommerce_sr(starting_dt, ending_dt):
    std_ecommerce_db = sqlContext.sql("""
    Select 
    --payments.created_date,
    weekofyear(payments_fact_druid.payments_created_date),
    ROUND(( COUNT(DISTINCT CASE WHEN payments_fact_druid.payments_authorized_at IS NOT NULL then payments_fact_druid.payments_id END ) )*1.000000/( COUNT(DISTINCT payments_fact_druid.payments_id ) ),4) as Std_checkout_Ecommerce_Sr
    FROM warehouse.payments  AS payments_fact_druid
    where payments_fact_druid.payments_created_date between '{0}' and '{1}'
    --and payments_fact_druid.payments_created_date < '{1}'
    and (CASE
            WHEN payments_fact_druid.payment_analytics_library = 1 THEN 'CHECKOUTJS'
            WHEN payments_fact_druid.payment_analytics_library = 2 THEN 'RAZORPAYJS'
            WHEN payments_fact_druid.payment_analytics_library = 3 THEN 'S2S'
            WHEN payments_fact_druid.payment_analytics_library = 4 THEN 'CUSTOM'
            WHEN payments_fact_druid.payment_analytics_library = 5 THEN 'DIRECT'
            WHEN payments_fact_druid.payment_analytics_library = 6 THEN 'PUSH'
            WHEN payments_fact_druid.payment_analytics_library = 7 THEN 'LEGACYJS'
            WHEN payments_fact_druid.payment_analytics_library = 8 THEN 'HOSTED' /* Hosted and embedded checkout identifers are added at the end of Nov, 2020 */
            WHEN payments_fact_druid.payment_analytics_library = 9 THEN 'EMBEDDED'
            ELSE 'UNKNOWN'
            END ) IN ('CHECKOUTJS', 'HOSTED') 
    AND (payments_fact_druid.merchants_category2 = 'ecommerce') 
    group by 1
    ORDER BY 1
        """.format(starting_dt.strftime('%Y-%m-%d'),ending_dt))
    std_ecommerce_df = std_ecommerce_db.toPandas()
    return std_ecommerce_df


# COMMAND ----------

# DBTITLE 1,WIP: Summary Screen Analysis 

summary_base_db = sqlContext.sql(
    """
    select
    --month(producer_created_date) as mnth,
    case when producer_created_date <date('2023-10-02') then 'Prev' else 'Post' end as mnth,
    merchant_id,
    browser_name,
    os_brand_family,
    initial_loggedin,
    case 
    when original_amount is null then null
    when original_amount < 500 then '<500'
    when original_amount < 1000 then '500 - 1k'
    when original_amount < 2000 then '1k - 2k'
    when original_amount < 5000 then '2k - 5k'
    when original_amount < 10000 then '5k - 10k'
    else '>10K' end as aov,
    sum(summary_screen_loaded) as summary_screen_loaded,
    sum(summary_screen_continue_cta_clicked) as summary_cta
    from aggregate_pa.magic_checkout_fact
    where producer_created_date between date('2023-09-25') and  date('2023-10-08') 
    group by 1,2,3,4,5,6

    """
)
summary_base_df = summary_base_db.toPandas()
summary_temp = summary_base_df
summary_temp.head()

# COMMAND ----------

def map_platform(platform):
    if platform in ['Android', 'iOS']:
        return platform
    else:
        return 'Others'

# Apply the function to the 'Platform' column
summary_temp['os_brand_family'] = summary_temp['os_brand_family'].apply(map_platform)


# COMMAND ----------

def table_calculations(table):
    table['Prev_Summary_CR'] = (table['summary_cta_Prev']/table['summary_screen_loaded_Prev'])
    table['Post_Summary_CR'] = (table['summary_cta_Post']/table['summary_screen_loaded_Post'])
    table['Prev_Summary_CR'].replace([np.inf, -np.inf], 0, inplace=True)
    table['Post_Summary_CR'].replace([np.inf, -np.inf], 0, inplace=True)
    table['Prev_volume'] = (table['summary_screen_loaded_Prev'] / table['summary_screen_loaded_Prev'].sum())
    table['Post_volume'] = (table['summary_screen_loaded_Post'] / table['summary_screen_loaded_Post'].sum())
    table['Vol_impact'] = table['Prev_Summary_CR']*(table['Post_volume'] - table['Prev_volume'])
    table['CR_impact'] = table['Post_volume']*(table['Post_Summary_CR']  - table['Prev_Summary_CR'])
    table['total_impact'] = table['Vol_impact'] + table['CR_impact'] 
    table['abs_total_impact'] = abs(table['total_impact'])
    return table



# COMMAND ----------



summary_grouped = summary_temp.groupby(by=['mnth','merchant_id','browser_name','os_brand_family','initial_loggedin','aov']).agg('sum').reset_index()


#summary_grouped['mnth'] = summary_grouped['mnth'].map(month_number_to_name)
summary_grouped['summary_cr'] = (summary_grouped['summary_cta'] / summary_grouped['summary_screen_loaded']).apply(percentage_conversion)
summary_df = summary_grouped.pivot(index=['merchant_id','browser_name','os_brand_family','initial_loggedin','aov'], columns='mnth', values=['summary_screen_loaded','summary_cta','summary_cr']).reset_index()
summary_df = summary_df.fillna(0)
summary_df.columns = ['_'.join(col).strip() for col in summary_df.columns.values]
#summary_df['Prev_volume'] = (summary_df['summary_screen_loaded_Prev'] / summary_df['summary_screen_loaded_Prev'].sum()).apply(percentage_conversion)
#summary_df['Post_volume'] = (summary_df['summary_screen_loaded_Post'] / summary_df['summary_screen_loaded_Post'].sum()).apply(percentage_conversion)
summary_df = table_calculations(summary_df)
summary_df.sort_values(by='abs_total_impact', ascending=False)



# COMMAND ----------

summary_df.to_csv('/dbfs/FileStore/summary_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/summary_df.csv"

# COMMAND ----------

summary_df = pd.read_csv('/dbfs/FileStore/summary_df.csv')
summary_df.head()

# COMMAND ----------



# COMMAND ----------

aov_summary_df = summary_df.groupby(by=['aov_',]).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_summary_df = table_calculations(aov_summary_df)
aov_summary_df.sort_values(by='abs_total_impact', ascending=False)


# COMMAND ----------

aov_summary_df.to_csv('/dbfs/FileStore/aov_summary_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_summary_df.csv"

# COMMAND ----------

aov_os_summary_df = summary_df.groupby(by=['aov_','os_brand_family_']).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_os_summary_df = table_calculations(aov_os_summary_df)
aov_os_summary_df.sort_values(by='abs_total_impact', ascending=False)


# COMMAND ----------

aov_os_summary_df.to_csv('/dbfs/FileStore/aov_os_summary_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_os_summary_df.csv"

# COMMAND ----------

aov_os_mid_df = summary_df.groupby(by=['aov_','os_brand_family_','merchant_id_']).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_os_mid_df = table_calculations(aov_os_mid_df)
aov_os_mid_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------

aov_os_mid_df.to_csv('/dbfs/FileStore/aov_os_mid_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_os_mid_df.csv"

# COMMAND ----------

aov_mid_df = summary_df.groupby(by=['aov_','merchant_id_',]).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_mid_df = table_calculations(aov_mid_df)
aov_mid_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------



# COMMAND ----------

aov_mid_df.to_csv('/dbfs/FileStore/aov_mid_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_mid_df.csv"

# COMMAND ----------

mid_df = summary_df.groupby(by=['merchant_id_',]).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
mid_df = table_calculations(mid_df)
mid_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------

mid_df.to_csv('/dbfs/FileStore/mid_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/mid_df.csv"

# COMMAND ----------

aov_mid_final_df = aov_mid_df.merge((mid_df.rename(columns={'summary_screen_loaded_Prev':'Prev_mid_total','summary_screen_loaded_Post':'Post_mid_total'}))[['merchant_id_','Prev_mid_total','Post_mid_total']], on='merchant_id_', how='left')
aov_mid_final_df['Prev: Vol of AOV within merchant'] = aov_mid_final_df['summary_screen_loaded_Prev'] / aov_mid_final_df['Prev_mid_total']
aov_mid_final_df['Post: Vol of AOV within merchant'] = aov_mid_final_df['summary_screen_loaded_Post'] / aov_mid_final_df['Post_mid_total']
aov_mid_final_df.sort_values(by='merchant_id_')

# COMMAND ----------

aov_mid_final_df.to_csv('/dbfs/FileStore/aov_mid_final_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_mid_final_df.csv"

# COMMAND ----------

aov_mid_logged_df = summary_df.groupby(by=['aov_','merchant_id_','initial_loggedin_']).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_mid_logged_df = table_calculations(aov_mid_logged_df)
aov_mid_logged_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------

aov_mid_logged_final_df = aov_mid_logged_df.merge((aov_mid_final_df.rename(columns={'summary_screen_loaded_Prev':'Prev_aov_mid_total','summary_screen_loaded_Post':'Post_aov_mid_total'}))[['merchant_id_','Prev_aov_mid_total','Post_aov_mid_total','Prev: Vol of AOV within merchant','Post: Vol of AOV within merchant','aov_']], on=['merchant_id_','aov_'], how='left')
aov_mid_logged_final_df['Prev: Vol of loggedIn within AOV/merchant'] = aov_mid_logged_final_df['summary_screen_loaded_Prev'] / aov_mid_logged_final_df['Prev_aov_mid_total']
aov_mid_logged_final_df['Post: Vol of loggedIn within AOV/merchant'] = aov_mid_logged_final_df['summary_screen_loaded_Post'] / aov_mid_logged_final_df['Post_aov_mid_total']
aov_mid_logged_final_df.sort_values(by='merchant_id_')

# COMMAND ----------

aov_mid_logged_final_df.to_csv('/dbfs/FileStore/aov_mid_logged_final_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_mid_logged_final_df.csv"

# COMMAND ----------

aov_mid_logged_os_df = summary_df.groupby(by=['aov_','merchant_id_','initial_loggedin_','os_brand_family_',]).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_mid_logged_os_df = table_calculations(aov_mid_logged_os_df)
aov_mid_logged_os_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------

aov_mid_logged_os_final = aov_mid_logged_os_df.merge((aov_mid_logged_final_df.rename(columns={'summary_screen_loaded_Prev':'Prev_aov_mid_logged_total',
'summary_screen_loaded_Post':'Post_aov_mid_logged_total'}))[['merchant_id_','Prev_aov_mid_total','Post_aov_mid_total','Prev: Vol of AOV within merchant','Post: Vol of AOV within merchant','aov_','initial_loggedin_','Prev: Vol of loggedIn within AOV/merchant','Post: Vol of loggedIn within AOV/merchant','Prev_aov_mid_logged_total','Post_aov_mid_logged_total']], on=['merchant_id_','aov_','initial_loggedin_'], how='left')
aov_mid_logged_os_final['Prev: Vol of OS within AOV/merchant/loggedin'] = aov_mid_logged_os_final['summary_screen_loaded_Prev'] / aov_mid_logged_os_final['Prev_aov_mid_logged_total']
aov_mid_logged_os_final['Post: Vol of OS within AOV/merchant/loggedin'] = aov_mid_logged_os_final['summary_screen_loaded_Post'] / aov_mid_logged_os_final['Post_aov_mid_logged_total']
aov_mid_logged_os_final.sort_values(by='merchant_id_')

# COMMAND ----------

aov_mid_logged_os_final.to_csv('/dbfs/FileStore/aov_mid_logged_os_final.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_mid_logged_os_final.csv"

# COMMAND ----------

aov_os_logged_df = summary_df.groupby(by=['aov_','os_brand_family_','initial_loggedin_']).agg({'summary_screen_loaded_Prev':'sum','summary_screen_loaded_Post':'sum',	'summary_cta_Prev':'sum',	'summary_cta_Post':'sum'}).reset_index()
aov_os_logged_df = table_calculations(aov_os_logged_df)
aov_os_logged_df.sort_values(by='abs_total_impact', ascending=False)

# COMMAND ----------

aov_mid_logged_final_df.to_csv('/dbfs/FileStore/aov_mid_logged_final_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_mid_logged_final_df.csv"

# COMMAND ----------

aov_os_logged_df.to_csv('/dbfs/FileStore/aov_os_logged_df.csv', index=False)
"https://razorpay-dev.cloud.databricks.com/files/aov_os_logged_df.csv"

# COMMAND ----------

funnel_db = sqlContext.sql(
    """
    select
    month(producer_created_date) as mnth,
    merchant_id,
    browser_name,
    os_brand_family,
    initial_loggedin,
    case 
    when original_amount/100 < 1000 then '<1000'
    when original_amount/100 < 2000 then '1k - 2k'
    when original_amount/100 < 5000 then '2k - 5k'
    when original_amount/100 < 10000 then '5k - 10k'
    else '>10K' end as aov,
    sum(open) as open,
     sum(summary_screen_loaded) as summary_screen_loaded,
    sum(summary_screen_continue_cta_clicked) as summary_cta,
    sum(payment_home_screen_loaded) as payment_home_screen_loaded,
    sum(submit) as submit
    from aggregate_pa.magic_checkout_fact
    where producer_created_date between date('2023-08-01') and  date('2023-09-30') 
    group by 1,2,3,4,5,6

    """
)
summary_base_df = summary_base_db.toPandas()

# COMMAND ----------

current_dt = date.today().strftime('%Y-%m-%d') #Input date for the current week's Monday as '2023-04-17' if data needs to be run for period other than last complete weeks and months
period='month'
term = 3

starting_dt, waterfall_starting_dt, ending_dt = get_starting_ending_dates(current_dt, period, term)

#Acquisition Tables
#base_df = create_acquisition_base_df(period, starting_dt, ending_dt)
print(base_df.columns)
#mtu = create_mtu(period, starting_dt, ending_dt,base_df)
print("MTU DF created")
#payment_table = create_payment_table(base_df, period, term)
print("payment DF created")
#merchant_sql = create_merchant_df()
print("Merchant DF created")
#orders_waterfall = create_orders_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
print("orders_waterfall created")
#gmv_waterfall = create_gmv_waterfall(base_df, merchant_sql, waterfall_starting_dt, period)
print("gmv_waterfall created")
#contact_prefill = get_contact_prefill_rate(period, starting_dt, ending_dt)
print("Contact Prefill Table Created")
#address_prefill = get_saved_address_prefill_rate(period, starting_dt, ending_dt)
print("Address Prefill Table Created")
#combined_cr = get_combined_cr_table(period,starting_dt,ending_dt, base_df)
print("Combined CR table created")





#create_doc()

# COMMAND ----------

visualization_list = [

    ['Acquisition',
        ['Orders', payment_table],
        ['MTU', mtu],
        ['Orders: Merchant Waterfall',orders_waterfall],
        ['GMV: Merchant Waterfall',gmv_waterfall],
    ],
    ['Consumer Experience',
        ['Combined CR Table', combined_cr],
       # ['Contact Prefill Rate', contact_prefill],
       # ['Address Prefill Rate', address_prefill],    
    ],
    

]
create_document_from_multi_level_list(visualization_list)

# COMMAND ----------

#contact prefill rate

contact_prefill_db = sqlContext.sql(
"""
WITH non_prefill AS(
  SELECT
    -- count(distinct checkout_id)
  merchant_id,
    checkout_id,
producer_created_date,
    CASE
      WHEN (
        CAST(
          get_json_object(properties, '$.data.prefill_contact_number') AS string
        ) IS NULL
        OR CAST(
           get_json_object(properties, '$.data.prefill_contact_number') AS string
        ) = ''
      ) THEN 0
      ELSE 1
    END AS prefill_contact_number,
  cast(get_json_object(properties, '$.data.meta.contact_prefill_source') as string) as contact_prefill_source,
  case when lower(get_json_object(context,'$.user_agent_parsed.os.family'))='ios' then 'iOS' else 'non-iOS' end as os_family,
  case when lower(browser_name) like '%safari%' then 'safari' else 'non-safari' end as browser,
  case when lower(get_json_object(context,'$.user_agent_parsed.os.family'))='ios' or lower(browser_name) like '%safari%' then 'not-applicable' else 'applicable' end as prefill_applicable
  FROM
    aggregate_pa.cx_1cc_events_dump_v1
  WHERE
    event_name = 'render:1cc_summary_screen_loaded_completed'
    --AND producer_created_date = date('2023-02-01') 
  and producer_created_date >= date('{0}')
   and producer_created_date <=  date('{1}')
),
contact_screen AS(
  SELECT
    checkout_id,
    get_json_object(properties, '$.data.contact_number') AS contact_number,
    producer_created_date
  FROM
    aggregate_pa.cx_1cc_events_dump_v1
  WHERE
    event_name = 'behav:1cc_summary_screen_contact_number_entered'
   -- AND producer_created_date = date('2023-02-01') 
  and producer_created_date >= date('{0}')
   and producer_created_date <= date('{1}')
    AND (
      get_json_object(properties, '$.data.contact_number') IS NOT NULL
      AND get_json_object(properties, '$.data.contact_number') <> ''
    )
)
SELECT

weekofyear(producer_created_date),


  round(COUNT(DISTINCT case when prefill_contact_number=1 then  non_prefill.checkout_id else null end)*1.0/count(DISTINCT non_prefill.checkout_id),4) as prefill_rate

FROM
  non_prefill
where producer_created_date >= date('{0}') and producer_created_date  <= date('{1}')
and prefill_applicable = 'applicable'
group by 1
order by 1 desc 
 
""".format(starting_dt.strftime('%Y-%m-%d'),ending_dt))

contact_prefill_df = contact_prefill_db.toPandas()
contact_prefill_df.head()

# COMMAND ----------

pre_magic_cr_db = sqlContext.sql(
"""
WITH Opens 
AS 
(
SELECT DISTINCT merchant_id, checkout_id
FROM aggregate_pa.cx_1cc_events_dump_v1
WHERE event_name = 'open'
AND producer_created_date BETWEEN DATE('{0}') AND DATE('{1}')
),

Payments AS 
(
SELECT p.id, pa.checkout_id, p.authorized_at, p.method
FROM realtime_hudi_api.payments p
LEFT JOIN (
            SELECT payment_id, checkout_id
            FROM realtime_hudi_api.payment_analytics
            WHERE created_date BETWEEN '{0}' AND '{1}'
          ) pa ON p.id = pa.payment_id
WHERE created_date BETWEEN '{0}' AND '{1}'
),

list as (Select mid ,segment,pre_magic_cr from batch_sheets.magic_merchant_list)


SELECT merchant_id, segment, pre_magic_cr,
round((((payment_success*100)/Open)*0.01),2)
FROM 
(
SELECT Opens.merchant_id,
       list.segment,
       list.pre_magic_cr,
       COUNT(
    DISTINCT CASE
      WHEN authorized_at IS NOT NULL
      AND lower(method) <> 'cod' THEN id
      WHEN lower(method) = 'cod' THEN id
      ELSE null
    END
  ) AS payment_success,
  COUNT(DISTINCT Opens.checkout_id) Open
FROM list
LEFT JOIN Opens ON Opens.merchant_id = list.mid
LEFT JOIN Payments ON Opens.checkout_id = Payments.checkout_id
WHERE Opens.checkout_id IS NOT NULL
GROUP BY 1, 2, 3
)
""".format(starting_dt.strftime('%Y-%m-%d'),ending_dt))

pre_magic_cr_df = pre_magic_cr_db.toPandas()
pre_magic_cr_df.head()

# COMMAND ----------

address_prefill_db = sqlContext.sql(
"""
WITH Opens 
AS 
(
SELECT DISTINCT merchant_id, checkout_id
FROM aggregate_pa.cx_1cc_events_dump_v1
WHERE event_name = 'open'
AND producer_created_date BETWEEN DATE('{0}') AND DATE('{1}')
),

Payments AS 
(
SELECT p.id, pa.checkout_id, p.authorized_at, p.method
FROM realtime_hudi_api.payments p
LEFT JOIN (
            SELECT payment_id, checkout_id
            FROM realtime_hudi_api.payment_analytics
            WHERE created_date BETWEEN '{0}' AND '{1}'
          ) pa ON p.id = pa.payment_id
WHERE created_date BETWEEN '{0}' AND '{1}'
),

list as (Select mid ,segment,pre_magic_cr from batch_sheets.magic_merchant_list)


SELECT merchant_id, segment, pre_magic_cr,
round((((payment_success*100)/Open)*0.01),2)
FROM 
(
SELECT Opens.merchant_id,
       list.segment,
       list.pre_magic_cr,
       COUNT(
    DISTINCT CASE
      WHEN authorized_at IS NOT NULL
      AND lower(method) <> 'cod' THEN id
      WHEN lower(method) = 'cod' THEN id
      ELSE null
    END
  ) AS payment_success,
  COUNT(DISTINCT Opens.checkout_id) Open
FROM list
LEFT JOIN Opens ON Opens.merchant_id = list.mid
LEFT JOIN Payments ON Opens.checkout_id = Payments.checkout_id
WHERE Opens.checkout_id IS NOT NULL
GROUP BY 1, 2, 3
)
""".format(starting_dt.strftime('%Y-%m-%d'),ending_dt))

address_prefill_df = address_prefill_db.toPandas()
address_prefill_df.head()

# COMMAND ----------



# COMMAND ----------

combined_cr_table = get_combined_cr_table('month',starting_dt,ending_dt, base_df)
combined_cr_table

# COMMAND ----------

data = [
   
    ['Acquisition',
        ['Orders', payment_table],
        ['Orders: Merchant Waterfall',orders_waterfall],
        ['GMV: Merchant Waterfall',gmv_waterfall],

    
    ],
    ['COD Intelligence',
        ['Orders', payment_table],
        ['Orders: Merchant Waterfall',orders_waterfall],
    ],
    [
        'CX'
    ]
        
]

# COMMAND ----------

data = [
    ['Acquisition',
        ['WTU','wtu table'],
        ['Orders', 'orders_df'],
        ['Std ECommerce SR','std_ecommerce_df'],

    
    ],
    [
        'COD Intelligence'
    ],
    [
        'CX'
    ]
]
for section in data:
    print(section)
    for item in section[1:]:
        print(item[0])
        print(item[1])

# COMMAND ----------

"https://razorpay-dev.cloud.databricks.com/files/shared_transfer/" + "Pallavi_Samodia/temp.docx"

# COMMAND ----------

dbutils.fs.ls("dbfs:/FileStore/Pallavi_Samodia/")

# COMMAND ----------

dbutils.fs.ls("file:/databricks/driver/")

# COMMAND ----------

display(wtu_grouped)
