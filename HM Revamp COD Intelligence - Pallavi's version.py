# Databricks notebook source
!/databricks/python3/bin/python -m pip install --upgrade pip
!/databricks/python3/bin/python -m pip install -U kaleido

# COMMAND ----------

!pip install python-docx
!pip install plotly
!pip install -U kaleido

# COMMAND ----------

from datetime import datetime, timedelta
from datetime import date
from dateutil.relativedelta import relativedelta
import pandas as pd
import os
import io
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.offline as pyo
import docx
import plotly.io as pio
from docx import Document
from docx.shared import Inches
from PIL import Image
import math
from docx.shared import Pt

# COMMAND ----------

def percentage_conversion(value):
    return f'{value/100:.2%}'

def convert_week_to_date_month(value):
    date_values = []
    for i in range(len(value)):
        week_int = int(value[i])
        date = datetime.date(datetime(current_year, 1, 1) + relativedelta(weeks=week_int-1,weekday=0))
        formated_date = date.strftime('%d %B')
        date_values.append(formated_date)
    return date_values

def convert_week_to_date_month_for_moving_avg(value):
    date_values = []
    for i in range(len(value)):
        week_int = int(value[i])
        start_date = datetime.date(datetime(current_year, 1, 1) + relativedelta(weeks=week_int,weekday=0))
        end_date = start_date + timedelta(days=-28) 
        formated_date_range = f"{end_date.strftime('%d %B')} - {start_date.strftime('%d %B')}"
        date_values.append(formated_date_range)
    return date_values

def week_to_date_running_avg(week_number):
    start_date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number-1,days=1))
    end_date = start_date + timedelta(days=-28)
    return f"{end_date.strftime('%d %B')} - {start_date.strftime('%d %B')}"

def week_to_date(week_number):
    date = datetime.date(datetime(current_year, 1, 1) + timedelta(weeks=week_number-1,days=1))
    return date.strftime('%d %B')

current_dt = date.today().strftime('%Y-%m-%d')
num_weeks = 6 #Input number of weeks we will be comparing the data for
waterfall_num_weeks = 2 #Input number of weeks we will be comparing the data for in the waterfall charts
num_days = num_weeks * 7
dt = datetime.strptime(current_dt, '%Y-%m-%d')
current_monday = (dt - timedelta(days = dt.weekday()))
week = dt.isocalendar()[1] #current_monday is also the ending date
starting_dt = (current_monday - timedelta(days=num_days))
ending_dt = (current_monday - timedelta(days=1)).strftime('%Y-%m-%d')
waterfall_starting_dt = (current_monday - timedelta(days=waterfall_num_weeks*7))
current_date = datetime.now().date()
previous_week_monday = current_date - timedelta(days=current_date.weekday(), weeks=1)
previous_week_month = previous_week_monday.strftime("%B")
previous_week_day = previous_week_monday.strftime("%d")
current_year = datetime.now().year


print('This report was generated on '+ str(current_dt) + ', in the week of '+ str(current_monday.strftime('%Y-%m-%d'))+ ' that is the week '+str(week)+' of the current year')
print('This report is comparing data between weeks of '+ str(starting_dt.strftime('%Y-%m-%d')) + ' and '+ str(ending_dt))
print('The waterfall charts will be comparing data between weeks of '+ str(waterfall_starting_dt.strftime('%Y-%m-%d')) + ' and '+ str(ending_dt))


# COMMAND ----------

base_db = sqlContext.sql("""
    SELECT *,
    month(rto_reimbursement.order_created_date) as order_created_week,
    CASE
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('category_5399')) THEN 'Marketplace model'
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('category_5691')) THEN 'Clothing model '
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('category_5977')) THEN 'Cosmetics model'
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('category_5732')) THEN 'Electronics model'
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('category_infrequent_mcc_group')) THEN 'Generic RTO model'
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('acm')) THEN 'Address model'
        WHEN (UPPER(rto_reimbursement.ml_model_id ) = UPPER('merchant_boltt')) THEN 'Boltt model'
        -- ELSE 'Other'
    END AS model_id,
    CASE 
        WHEN (UPPER(CASE
            WHEN UPPER(UPPER(rto_reimbursement.result_flag)) = UPPER('green') THEN 'safe'
            WHEN (UPPER(UPPER(rto_reimbursement.result_flag)) = UPPER('red')) AND rto_reimbursement.experimentation THEN 'safe'
            WHEN (UPPER(UPPER(rto_reimbursement.result_flag)) = UPPER('red')) AND (NOT (rto_reimbursement.experimentation)) THEN 'risky'
            ELSE 'unknown'
        END) = UPPER('risky')) THEN rto_reimbursement.order_id 
        ELSE NULL 
    END 
    AS risky,
    
    CASE 
        WHEN ((UPPER(rto_reimbursement.status ) = UPPER('rto'))) AND ((UPPER(
            CASE 
                WHEN CASE 
                        WHEN UPPER(UPPER(rto_reimbursement.ml_flag)) = UPPER('green') THEN true 
                        ELSE false 
                    END AND 
                    CASE 
                        WHEN UPPER(UPPER(rto_reimbursement.rule_flag)) = UPPER('green') THEN true 
                        ELSE false 
                    END THEN 'Safe' 
                ELSE 'Risky' 
            END) = UPPER('Risky'))) THEN rto_reimbursement.order_id 
        ELSE NULL 
    END AS rto_risky,

    CASE 
        WHEN ((UPPER(rto_reimbursement.status ) = UPPER('rto'))) AND ((UPPER(CASE 
            WHEN 
                CASE 
                    WHEN UPPER(UPPER(rto_reimbursement.ml_flag)) = UPPER('green') THEN true 
                    ELSE false 
                END AND 
            CASE 
                WHEN UPPER(UPPER(rto_reimbursement.rule_flag)) = UPPER('green') THEN true 
                ELSE false 
            END
            THEN 'Safe' 
            ELSE 'Risky' 
            END) = UPPER('Safe'))) THEN rto_reimbursement.order_id 
        ELSE NULL 
    END AS rto_safe,

    CASE 
        WHEN ((UPPER(rto_reimbursement.status ) <> UPPER('rto') OR rto_reimbursement.status IS NULL)) AND ((UPPER(
            CASE 
                WHEN 
                    CASE 
                        WHEN UPPER(UPPER(rto_reimbursement.ml_flag)) = UPPER('green') THEN true 
                        ELSE false 
                    END
                AND 
                CASE 
                    WHEN UPPER(UPPER(rto_reimbursement.rule_flag)) = UPPER('green') THEN true 
                    ELSE false 
                END THEN 'Safe' 
                ELSE 'Risky' 
            END ) = UPPER('Risky'))) THEN rto_reimbursement.order_id 
            ELSE NULL 
    END AS non_rto_risky,

    CASE 
        WHEN ((UPPER(rto_reimbursement.status ) <> UPPER('rto') OR rto_reimbursement.status IS NULL)) AND ((UPPER(
            CASE 
                WHEN 
                    CASE 
                        WHEN UPPER(UPPER(rto_reimbursement.ml_flag)) = UPPER('green') THEN true 
                        ELSE false 
                    END AND 
                    CASE 
                        WHEN UPPER(UPPER(rto_reimbursement.rule_flag)) = UPPER('green') THEN true 
                        ELSE false 
                    END THEN 'Safe' 
                ELSE 'Risky' 
            END) = UPPER('Safe'))) THEN rto_reimbursement.order_id 
        ELSE NULL 
    END AS non_rto_safe
    from aggregate_pa.magic_rto_reimbursement_fact as rto_reimbursement
    WHERE rto_reimbursement.order_created_date >= '{0}' 
    and rto_reimbursement.order_created_date < '{1}'
""".format(starting_dt.strftime('%Y-%m-%d'),ending_dt))
base_df = base_db.toPandas()
base_df.head()

# COMMAND ----------

#red labelling rate

red_labelling = base_df[(base_df['merchant_id'].notna()) & (base_df['order_status'].notna()) & (base_df['result_flag'].notna())].groupby('order_created_week').agg({
    'order_id':'count',
    'risky':'count' ,
}).reset_index()

red_labelling_date = convert_week_to_date_month(red_labelling['order_created_week'])
red_labelling['orders_created_week'] = pd.DataFrame({'date':red_labelling_date})
red_labelling['risky%'] = (red_labelling['risky']/red_labelling['order_id']*100).round(2)
red_labelling =  red_labelling.drop(['order_id','risky','order_created_week'],axis=1)
red_labelling


# COMMAND ----------

red_label_img = plt.figure(figsize=(25, 8))
plt.plot(red_labelling['orders_created_week'].values, red_labelling['risky%'].values,marker='o')
plt.ylim(0, red_labelling['risky%'].max() + 5)
plt.xlabel('Magic RTO Reimbursement fact Order Created Date Week',fontsize=15 )
plt.xticks(fontsize=18,fontweight='bold')
for i, row in red_labelling.iterrows():
    plt.text(row['orders_created_week'], row['risky%'], f'{row["risky%"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')
plt.show()
#recent week

# COMMAND ----------

#red labelling model impact

total_orders = base_df[(base_df['merchant_id'].notna()) & (base_df['order_status'].notna()) & (base_df['result_flag'].notna())].groupby('order_created_week')['order_id'].count()

red_labelling_model_df = base_df[(base_df['merchant_id'].notna()) & (base_df['order_status'].notna()) & (base_df['result_flag'].notna())].groupby(['order_created_week','model_id']).agg(
    risky = ('risky','count'),
    orders = ('order_id','count'),
)
red_labelling_model_group = red_labelling_model_df
red_labelling_model_group.index = red_labelling_model_group.index.set_levels(red_labelling_model_group.index.levels[0].map(week_to_date), level='order_created_week')
total_orders.index = total_orders.index.map(week_to_date)

red_labelling_model_group['risky%'] = (red_labelling_model_group['risky'] / red_labelling_model_group['orders'] * 100).round(2)
red_labelling_model_group['order volume%'] = (red_labelling_model_group['orders'] / total_orders * 100).round(2)

red_labelling_model_group = red_labelling_model_group.drop(['orders','risky'],axis=1)

red_labelling_model_group = red_labelling_model_group.stack()
red_labelling_model_group = red_labelling_model_group.unstack(level = 0)
red_labelling_model_group = red_labelling_model_group.unstack(level = 1)
# red_labelling_model_group = red_labelling_model_group.drop(red_labelling_model_group.columns[0:8], axis=1)
red_labelling_model_group = red_labelling_model_group.drop(red_labelling_model_group.columns[0:4], axis=1)
red_labelling_model_group = red_labelling_model_group.drop(red_labelling_model_group.columns[-4:], axis=1)

last_week_date = week_to_date(week-3)
second_last_week_date = week_to_date(week-4)
second_last_week_order_total = red_labelling_model_group[(second_last_week_date,'order volume%')].sum()
last_week_order_total = red_labelling_model_group[(last_week_date, 'order volume%')].sum()
second_last_week_risky_val = red_labelling[red_labelling['orders_created_week']==second_last_week_date].values[0][1]
last_week_risky_val = red_labelling[red_labelling['orders_created_week']==last_week_date].values[0][1]
red_labelling_total = [second_last_week_risky_val, second_last_week_order_total, last_week_risky_val, last_week_order_total]
red_labelling_total = pd.DataFrame([red_labelling_total],columns=red_labelling_model_group.columns,index=['Total'])

red_labelling_model_group = pd.concat([red_labelling_model_group, red_labelling_total])
red_labelling_model_group = red_labelling_model_group.round(2).reset_index()
red_labelling_model_group = red_labelling_model_group.rename(columns={'index':'model_id'})
# for i in range(week - num_weeks + 2 , week - 2):
#     date = week_to_date(i)
#     red_labelling_model_group[(date,'order volume%')] = red_labelling_model_group[(date,'order volume%')].apply(percentage_conversion)
#     red_labelling_model_group[(date,'risky%')] = red_labelling_model_group[(date,'risky%')].apply(percentage_conversion)
    
red_labelling_model_group

# COMMAND ----------

# red_flagging_volm_and_cr_table = red_labelling_model_df
total_orders = base_df[(base_df['merchant_id'].notna()) & (base_df['order_status'].notna()) & (base_df['result_flag'].notna())].groupby('order_created_week')['order_id'].count()

red_flagging_volm_and_cr_table = base_df[(base_df['merchant_id'].notna()) & (base_df['order_status'].notna()) & (base_df['result_flag'].notna())].groupby(['order_created_week','model_id']).agg(
    risky = ('risky','count'),
    orders = ('order_id','count'),
)

red_flagging_volm_and_cr_table.index = red_flagging_volm_and_cr_table.index.set_levels(red_flagging_volm_and_cr_table.index.levels[0].map(week_to_date), level='order_created_week')
total_orders.index = total_orders.index.map(week_to_date)

red_flagging_volm_and_cr_table['risky%'] = red_flagging_volm_and_cr_table['risky']/red_flagging_volm_and_cr_table['orders']*100
red_flagging_volm_and_cr_table['order volume%'] = red_flagging_volm_and_cr_table['orders'] / total_orders * 100


red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.drop(['risky','orders'],axis=1)
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.stack()
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.unstack(level = 0)
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.unstack(level = 1)
# red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.drop(red_flagging_volm_and_cr_table.columns[0:15],axis=1)
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.drop(red_flagging_volm_and_cr_table.columns[0:4],axis=1)
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.drop(red_flagging_volm_and_cr_table.columns[-4:],axis=1)

previous_week_cr = []
previous_week_volm = []
for j in range(len(red_flagging_volm_and_cr_table[week_to_date(week-4)])):
    previous_week_cr_value = red_flagging_volm_and_cr_table[week_to_date(week-4)]['risky%'].iloc[j]
    previous_week_volume_value = red_flagging_volm_and_cr_table[week_to_date(week-4)]['order volume%'].iloc[j]
    previous_week_cr.append(previous_week_cr_value)
    previous_week_volm.append(previous_week_volume_value)
red_flagging_volm_and_cr_table[(' ','previous_week_cr')] = previous_week_cr
red_flagging_volm_and_cr_table[(' ','previous_week_volm')] = previous_week_volm

red_flagging_volm_and_cr_table[(' ','Volume Impact')] = (red_flagging_volm_and_cr_table[(' ','previous_week_cr')] * (red_flagging_volm_and_cr_table[(week_to_date(week-3)),'order volume%']-red_flagging_volm_and_cr_table[(' ','previous_week_volm')])/100).round(2)

red_flagging_volm_and_cr_table[(' ','CR Impact')] = (red_flagging_volm_and_cr_table[(week_to_date(week-3),'order volume%')] * (red_flagging_volm_and_cr_table[(week_to_date(week-3)),'risky%']-red_flagging_volm_and_cr_table[(' ','previous_week_cr')])/100)

red_flagging_volm_and_cr_table[(' ','Total Impact')] = (red_flagging_volm_and_cr_table[(' ','Volume Impact')] +  red_flagging_volm_and_cr_table[(' ','CR Impact')]).round(2)
red_flagging_volm_and_cr_table = red_flagging_volm_and_cr_table.drop(red_flagging_volm_and_cr_table.columns[4:6],axis=1)

last_week_volume_total = percentage_conversion(red_flagging_volm_and_cr_table[(' ', 'Volume Impact')].sum())
last_week_cr_total = percentage_conversion(red_flagging_volm_and_cr_table[(' ', 'CR Impact')].sum())
last_week_total_impact = percentage_conversion(red_flagging_volm_and_cr_table[(' ', 'Total Impact')].sum())
red_flag_vol_cr_total = [last_week_volume_total, last_week_cr_total, last_week_total_impact]

total = []
for weeks in range(week - num_weeks +2, week -2):
    date = week_to_date(weeks)
    risky_total = percentage_conversion(red_labelling[red_labelling['orders_created_week']==date].values[0][1])
    volumne_total = percentage_conversion(red_flagging_volm_and_cr_table[(date),'order volume%'].sum().round(2))
    red_flagging_volm_and_cr_table[(date,'order volume%')] = red_flagging_volm_and_cr_table[(date,'order volume%')].apply(percentage_conversion)
    red_flagging_volm_and_cr_table[(date,'risky%')] = red_flagging_volm_and_cr_table[(date,'risky%')].apply(percentage_conversion)
    total.append(risky_total)
    total.append(volumne_total)
red_flag_vol_cr_total
red_flagging_volm_and_cr_table[(' ', 'Total Impact')] = red_flagging_volm_and_cr_table[(' ', 'Total Impact')].apply(percentage_conversion)
red_flagging_volm_and_cr_table[(' ', 'Volume Impact')] = red_flagging_volm_and_cr_table[(' ', 'Volume Impact')].apply(percentage_conversion)
red_flagging_volm_and_cr_table[(' ', 'CR Impact')] = red_flagging_volm_and_cr_table[(' ', 'CR Impact')].apply(percentage_conversion)
red_flag_vol_cr_total = total + red_flag_vol_cr_total
red_flag_vol_cr_total = pd.DataFrame([red_flag_vol_cr_total],columns=red_flagging_volm_and_cr_table.columns,index=['Total'])


red_flagging_volm_and_cr_table = pd.concat([red_flagging_volm_and_cr_table, red_flag_vol_cr_total]).reset_index()
red_labelling_model_impact = red_flagging_volm_and_cr_table.rename(columns={'index':'model_id'})
red_labelling_model_impact


# COMMAND ----------

# waterfall chart calculation

previous_week_date = week_to_date(week-3)

volume = red_labelling_model_impact[(' ','Volume Impact')].iloc[-1]

red_flagging_chart = red_labelling_model_impact[['model_id']]

red_flagging_chart = red_flagging_chart.drop(red_labelling_model_impact.index[-1])

red_flagging_chart['Overall Impact'] = red_labelling_model_impact.loc[:,(' ','CR Impact')]

volume_impact = pd.DataFrame({('model_id', ''):['Volume Impact'],('Overall Impact', ''):[volume]})

subtotal_value = red_labelling_model_impact[(' ','Total Impact')].iloc[-1]

subtotal = pd.DataFrame({('model_id', ''):['subtotal'],('Overall Impact', ''):[subtotal_value]})

red_flag_waterfall_table = pd.concat([red_flagging_chart,volume_impact,subtotal])

second_last_row = red_flag_waterfall_table.iloc[-2].copy()  
red_flag_waterfall_table.iloc[1:-1] = red_flag_waterfall_table.iloc[:-2].values  
red_flag_waterfall_table.iloc[0] = second_last_row

pd.set_option('chained_assignment', None)
red_flag_waterfall_table

# COMMAND ----------

#waterfall_chart
red_flag_waterfall_table['Overall Impact'] = red_flag_waterfall_table['Overall Impact'].str.replace('%',' ')
red_flag_waterfall_table['Overall Impact'] = pd.to_numeric(red_flag_waterfall_table['Overall Impact'])

model_id = red_flag_waterfall_table['model_id'].values
impact = red_flag_waterfall_table['Overall Impact'].values

measures =  ['relative'] + ['relative'] * (len(model_id) - 2) + ['total']

trace = go.Waterfall(
    x=model_id,
    y=impact,
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.6,
    text=red_flag_waterfall_table['Overall Impact'].apply(float).apply(lambda x: round(x , 2)).apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    textposition='outside',
)

layout = go.Layout(
    title='Weekly Red Flagging Rate',
    title_font=dict(size=34),
    width=750,
    height=1050,
)
red_flag_fig = go.Figure(data=trace, layout=layout)
red_flag_fig = red_flag_fig.update_xaxes(tickfont=dict(size=40,family='Arial, bold'))
red_flag_fig.show()


# COMMAND ----------

#precision ,recall and fscore
precision_recall_and_fscore_df = base_df[(base_df['cod_intelligence_enabled']==True)&(base_df['experimentation']==True)&(base_df['merchant_id'].notna())&(base_df['status'] != 'created')].groupby('order_created_week').agg({
    'rto_risky':'count',
    'rto_safe':'count' ,
    'non_rto_risky':'count',
    'non_rto_safe':'count'
})
precision_recall_and_fscore_df.index = precision_recall_and_fscore_df.index.map(week_to_date)

comparison_bw_precision_recall_fscore = precision_recall_and_fscore_df
comparison_bw_precision_recall_fscore['precision'] = (comparison_bw_precision_recall_fscore['rto_risky']/(comparison_bw_precision_recall_fscore['rto_risky']+comparison_bw_precision_recall_fscore['non_rto_risky'])*100).round(2)

comparison_bw_precision_recall_fscore['recall'] = (comparison_bw_precision_recall_fscore['rto_risky']/(comparison_bw_precision_recall_fscore['rto_risky']+comparison_bw_precision_recall_fscore['rto_safe'])*100).round(2)

comparison_bw_precision_recall_fscore['fscore'] = (2*comparison_bw_precision_recall_fscore['recall']*comparison_bw_precision_recall_fscore['precision']/(comparison_bw_precision_recall_fscore['precision']+comparison_bw_precision_recall_fscore['recall'])).round(2)



comparison_bw_precision_recall_fscore =  comparison_bw_precision_recall_fscore.drop(['rto_risky','rto_safe','non_rto_risky','non_rto_safe'],axis=1)
comparison_bw_precision_recall_fscore = comparison_bw_precision_recall_fscore.reset_index()
comparison_bw_precision_recall_fscore = comparison_bw_precision_recall_fscore.fillna(0)
comparison_bw_precision_recall_fscore



# COMMAND ----------


precision_recall_fscore_comparison=plt.figure(figsize=(25, 10))
plt.plot(comparison_bw_precision_recall_fscore['order_created_week'].values, comparison_bw_precision_recall_fscore['precision'].values,marker='o',label='Precison')
plt.plot(comparison_bw_precision_recall_fscore['order_created_week'].values, comparison_bw_precision_recall_fscore['recall'].values,marker='o',label='Recall')
plt.plot(comparison_bw_precision_recall_fscore['order_created_week'].values, comparison_bw_precision_recall_fscore['fscore'].values,marker='o',label='Fscore')

plt.ylim(0, comparison_bw_precision_recall_fscore['recall'].values.max() + 5)
plt.xticks(fontsize=18,fontweight='bold')
plt.xlabel('Week Range ',fontsize=18)
plt.title('Precision, Recall and F Score Weekly Comparison')
for i, row in comparison_bw_precision_recall_fscore.iterrows():
    plt.text(row['order_created_week'], row['precision'], f'{row["precision"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

for i, row in comparison_bw_precision_recall_fscore.iterrows():
    plt.text(row['order_created_week'], row['recall'], f'{row["recall"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

for i, row in comparison_bw_precision_recall_fscore.iterrows():
    plt.text(row['order_created_week'], row['fscore'], f'{row["fscore"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

plt.legend()
plt.show()


# COMMAND ----------

#running four week avg

running_avg_df = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status'] != 'created')].groupby('order_created_week').agg({
    'rto_risky':'count',
    'rto_safe':'count' ,
    'non_rto_risky':'count',
    'non_rto_safe':'count'
})

running_avg_df.index = running_avg_df.index.map(week_to_date_running_avg)

running_avg_df['precision'] = (running_avg_df['rto_risky']/(running_avg_df['rto_risky']+running_avg_df['non_rto_risky'])*100).round(2)
running_avg_df['recall'] = (running_avg_df['rto_risky']/(running_avg_df['rto_risky']+running_avg_df['rto_safe'])*100).round(2)
running_avg_df['fscore'] = (2*running_avg_df['recall']*running_avg_df['precision']/(running_avg_df['precision']+running_avg_df['recall'])).round(2)
 
running_avg_df['running_4_month_avg_of_precision'] = running_avg_df['precision'].rolling(window =4, min_periods=1).mean()
running_avg_df['running_4_month_avg_of_recall'] = running_avg_df['recall'].rolling(window =4, min_periods=1).mean()
running_avg_df['running_4_month_avg_of_fscore'] = running_avg_df['fscore'].rolling(window =4, min_periods=1).mean()
running_avg_df = running_avg_df.drop(running_avg_df.columns[0:7],axis=1).reset_index()

moving_avg = plt.figure(figsize=(25, 10))
# running_avg_date =  convert_week_to_date_month_for_moving_avg(running_avg_df['order_created_week'])
# running_avg_df['date_month'] = pd.DataFrame({'date':running_avg_date})
plt.plot(running_avg_df['order_created_week'].values, running_avg_df['running_4_month_avg_of_precision'].values,marker='o',label='Precison')
plt.plot(running_avg_df['order_created_week'].values, running_avg_df['running_4_month_avg_of_recall'].values,marker='o',label='Recall')
plt.plot(running_avg_df['order_created_week'].values, running_avg_df['running_4_month_avg_of_fscore'].values,marker='o',label='Fscore')

plt.ylim(0, running_avg_df['running_4_month_avg_of_recall'].values.max() + 5)
plt.xlabel('Week Range ',fontsize=18)
plt.xticks(fontsize=18,fontweight='bold')
plt.title('Precision, Recall and F Score (4 week moving average)')
for i, row in running_avg_df.iterrows():
    plt.text(row['order_created_week'], row['running_4_month_avg_of_precision'], f'{row["running_4_month_avg_of_precision"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

for i, row in running_avg_df.iterrows():
    plt.text(row['order_created_week'], row['running_4_month_avg_of_recall'], f'{row["running_4_month_avg_of_recall"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

for i, row in running_avg_df.iterrows():
    plt.text(row['order_created_week'], row['running_4_month_avg_of_fscore'], f'{row["running_4_month_avg_of_fscore"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

plt.legend()
plt.show()


# COMMAND ----------

#precision

precision_df = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week').agg({
    'rto_risky':'count',
    'rto_safe':'count' ,
    'non_rto_risky':'count',
    'non_rto_safe':'count'
})
precision_df.index = precision_df.index.map(week_to_date)
precision_df['precision%'] = precision_df['rto_risky']/(precision_df['rto_risky']+precision_df['non_rto_risky'])*100
precision_df = precision_df.drop(['rto_risky','rto_safe','non_rto_risky','non_rto_safe'],axis=1).reset_index()

#precision line chart
weekly_precision_rate = plt.figure(figsize=(20, 6))
plt.plot(precision_df['order_created_week'].values, precision_df['precision%'].values,marker='o')
plt.ylim(0, precision_df['precision%'].max() + 5)
plt.xlabel('Precision Order Created Date Week',fontsize=18)
plt.xticks(fontsize=18,fontweight='bold')
for i, row in precision_df.iterrows():
    plt.text(row['order_created_week'], row['precision%'], f'{row["precision%"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')
plt.show()


# COMMAND ----------

valid_status = ['delivered','rto','cancelled','lost','partially_delivered','returned']
precision_model_df = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby(['order_created_week','model_id']).agg(
    rto_risky = ('rto_risky','count'),
    rto_safe = ('rto_safe','count') ,
    non_rto_risky = ('non_rto_risky','count'),
    non_rto_safe = ('non_rto_safe','count'),
    order_id = ('order_id','count'),
    order_count_w_delivery = ('status', lambda x: x[x.isin(valid_status)].count()),
    order_count_w_rto = ('status', lambda x: x[x.isin(['rto'])].count())
)

#shipping date available %
shipping_order = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week').agg(
    order_count_w_delivery = ('status', lambda x: x[x.isin(valid_status)].count())
)

#total shipping order
total_shipping_order = shipping_order['order_count_w_delivery'].sum()
shipping_order['shipping_order%'] = (shipping_order['order_count_w_delivery']/total_shipping_order*100).round(2)
shipping_order = shipping_order.drop(['order_count_w_delivery'],axis=1)


#total order weekly
total_order = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week')['order_id'].count()

#converting week number to date
precision_model_df.index = precision_model_df.index.set_levels(precision_model_df.index.levels[0].map(week_to_date), level='order_created_week')
total_order.index = total_order.index.map(week_to_date)
shipping_order.index = shipping_order.index.map(week_to_date)
shipping_order = shipping_order.reset_index()

#precision%
precision_model_df['precision%'] = (precision_model_df['rto_risky']/(precision_model_df['rto_risky']+precision_model_df['non_rto_risky'])*100).round(2)

#risky order volume%
precision_model_df['total_risky'] = (precision_model_df['rto_risky']+precision_model_df['non_rto_risky'])/100
precision_model_df['order%'] = precision_model_df['rto_risky']+precision_model_df['non_rto_risky']
total_risky_order = precision_model_df.groupby('order_created_week')['order%'].sum()
precision_model_df['risky order volume%'] = ((precision_model_df['rto_risky']+precision_model_df['non_rto_risky'])/total_risky_order*100).round(2)

#shipping data available%
precision_model_df['shipping data available%'] = (precision_model_df['order_count_w_delivery']/total_order*100).round(2)

precision_model_df = precision_model_df.drop(['rto_risky','rto_safe','non_rto_risky','non_rto_safe','order_id','total_risky','order%','order_count_w_delivery','order_count_w_rto'],axis=1)


precision_model_df = precision_model_df.stack()
precision_model_df = precision_model_df.unstack(level = 0)
precision_model_df = precision_model_df.unstack(level = 1)

precision_model_df = precision_model_df.drop(precision_model_df.columns[0:6], axis=1)
precision_model_df = precision_model_df.drop(precision_model_df.columns[-6:],axis=1)
precision_model_df = precision_model_df.fillna(0)
# precision_model_df

previous_week_cr = []
previous_week_volm = []
for j in range(len(precision_model_df[week_to_date(week-4)])):
    previous_week_cr_value = precision_model_df[week_to_date(week-4)]['precision%'].iloc[j]
    previous_week_volume_value = precision_model_df[week_to_date(week-4)]['risky order volume%'].iloc[j]
    previous_week_cr.append(previous_week_cr_value)
    previous_week_volm.append(previous_week_volume_value)
precision_model_df[(' ','previous_week_cr')] = previous_week_cr
precision_model_df[(' ','previous_week_volm')] = previous_week_volm
precision_model_df[(' ','Volume Impact')] = (precision_model_df[(' ','previous_week_cr')] * (precision_model_df[(week_to_date(week-3)),'risky order volume%']-precision_model_df[(' ','previous_week_volm')])/100).round(2)

precision_model_df[(' ','CR Impact')] = (precision_model_df[(week_to_date(week-3),'risky order volume%')] * (precision_model_df[(week_to_date(week-3)),'precision%']-precision_model_df[(' ','previous_week_cr')])/100).round(2)

precision_model_df[(' ','Total Impact')] = (precision_model_df[(' ','Volume Impact')] +  precision_model_df[(' ','CR Impact')]).round(2)
precision_model_df = precision_model_df.drop(precision_model_df.columns[6:8],axis=1)

last_week_volume_total = percentage_conversion(precision_model_df[(' ', 'Volume Impact')].sum())
last_week_cr_total = percentage_conversion(precision_model_df[(' ', 'CR Impact')].sum())
last_week_total_impact = percentage_conversion(precision_model_df[(' ', 'Total Impact')].sum())
precision_vol_cr_total = [last_week_volume_total, last_week_cr_total, last_week_total_impact]

precision_df_total = []
for weeks in range(week - num_weeks + 2 , week  - 2):
    date = week_to_date(weeks)
    precision_total = percentage_conversion(precision_df[precision_df['order_created_week']==date].values[0][1])
    volumne_total = percentage_conversion(precision_model_df[(date),'risky order volume%'].sum().round(2))
    shipping_total = percentage_conversion(precision_model_df[(date),'shipping data available%'].sum().round(2))
    precision_model_df[(date,'risky order volume%')] = precision_model_df[(date,'risky order volume%')].apply(percentage_conversion)
    precision_model_df[(date,'precision%')] = precision_model_df[(date,'precision%')].apply(percentage_conversion)
    precision_model_df[(date,'shipping data available%')] = precision_model_df[(date,'shipping data available%')].apply(percentage_conversion)
    precision_df_total.append(precision_total)
    precision_df_total.append(volumne_total)
    precision_df_total.append(shipping_total)

precision_model_df[(' ', 'Total Impact')] = precision_model_df[(' ', 'Total Impact')].apply(percentage_conversion)
precision_model_df[(' ', 'Volume Impact')] = precision_model_df[(' ', 'Volume Impact')].apply(percentage_conversion)
precision_model_df[(' ', 'CR Impact')] = precision_model_df[(' ', 'CR Impact')].apply(percentage_conversion)
precision_vol_cr_total = precision_df_total + precision_vol_cr_total
precision_vol_cr_total = pd.DataFrame([precision_vol_cr_total],columns=precision_model_df.columns,index=['Total'])


precision_model_df = pd.concat([precision_model_df, precision_vol_cr_total]).reset_index()
precision_model_df = precision_model_df.rename(columns={'index':'model_id'})
precision_model_df


# COMMAND ----------

# waterfall chart calculation

previous_week_date = week_to_date(week-3)

volume = precision_model_df[(' ','Volume Impact')].iloc[-1]

precision_chart = precision_model_df[['model_id']]

precision_chart = precision_chart.drop(precision_model_df.index[-1])

precision_chart['Total Impact'] = precision_model_df.loc[:,(' ','CR Impact')]

volume_impact = pd.DataFrame({('model_id', ''):['Volume Impact'],('Total Impact', ''):[volume]})

subtotal_value = precision_model_df[(' ','Total Impact')].iloc[-1]

subtotal = pd.DataFrame({('model_id', ''):['subtotal'],('Total Impact', ''):[subtotal_value]})

precision_waterfall_table = pd.concat([precision_chart,volume_impact,subtotal])

second_last_row = precision_waterfall_table.iloc[-2].copy()  
precision_waterfall_table.iloc[1:-1] = precision_waterfall_table.iloc[:-2].values  
precision_waterfall_table.iloc[0] = second_last_row

pd.set_option('chained_assignment', None)
precision_waterfall_table = precision_waterfall_table.fillna(0)
precision_waterfall_table


# COMMAND ----------

#waterfall_chart
precision_waterfall_table['Total Impact'] = precision_waterfall_table['Total Impact'].str.replace('%',' ')
precision_waterfall_table['Total Impact'] = pd.to_numeric(precision_waterfall_table['Total Impact'])
model_id = precision_waterfall_table['model_id'].values
impact = precision_waterfall_table['Total Impact'].values

measures =  ['relative'] + ['relative'] * (len(model_id) - 2) + ['total']

trace = go.Waterfall(
    x=model_id,
    y=impact,
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.6,
    text=precision_waterfall_table['Total Impact'].apply(lambda x: round(x , 2)).apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),

    textposition='outside',
)

layout = go.Layout(
    title=' Precision Rate',
    title_font=dict(size=34),
    width=750,
    height=1150,
)
precision_flag_fig = go.Figure(data=trace, layout=layout)
precision_flag_fig = precision_flag_fig.update_xaxes(tickfont=dict(size=40,family='Arial, bold'))
precision_flag_fig.show()


# COMMAND ----------

#recall

recall_df = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week').agg({
    'rto_risky':'count',
    'rto_safe':'count' ,
    'non_rto_risky':'count',
    'non_rto_safe':'count'
})
recall_df.index = recall_df.index.map(week_to_date)
recall_df['recall%'] = recall_df['rto_risky']/(recall_df['rto_risky']+recall_df['rto_safe'])*100
recall_df = recall_df.drop(['rto_risky','rto_safe','non_rto_risky','non_rto_safe'],axis=1).reset_index()
recall_df = recall_df.fillna(0)


recall_weekly_rate = plt.figure(figsize=(20, 6))
plt.plot(recall_df['order_created_week'].values, recall_df['recall%'].values,marker='o')
plt.ylim(0, recall_df['recall%'].max() + 5)
plt.xlabel('Recall Order Created Date Week',fontsize=18)
plt.xticks(fontsize=18,fontweight='bold')
for i, row in recall_df.iterrows():
    plt.text(row['order_created_week'], row['recall%'], f'{row["recall%"]:.2f}%', ha='center', va='bottom',fontsize=18, fontweight='bold')

plt.show()

# COMMAND ----------

#recall model impact
valid_status = ['delivered','rto','cancelled','lost','partially_delivered','returned']

total_order = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week')['order_id'].count()

#shipping data avialable%
shipping_order = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby('order_created_week').agg(
    order_count_w_delivery = ('status', lambda x: x[x.isin(valid_status)].count())
)
total_shipping_order = shipping_order['order_count_w_delivery'].sum()
shipping_order['shipping_order%'] = (shipping_order['order_count_w_delivery']/total_shipping_order*100).round(2)
shipping_order = shipping_order.drop(['order_count_w_delivery'],axis=1)

#base_df_recall
recall_model_impact = base_df[(base_df['cod_intelligence_enabled'])&(base_df['experimentation'])&(base_df['merchant_id'].notna())&(base_df['status']!='created')&(base_df['status'].notna())].groupby(['order_created_week','model_id']).agg(
    rto_risky = ('rto_risky','count'),
    rto_safe = ('rto_safe','count') ,
    non_rto_risky = ('non_rto_risky','count'),
    non_rto_safe = ('non_rto_safe','count'),
    order_id = ('order_id','count'),
    order_count_w_delivery = ('status', lambda x: x[x.isin(valid_status)].count()),
    order_count_w_rto = ('status', lambda x: x[x.isin(['rto'])].count())
)

#converting week number to date
recall_model_impact.index = recall_model_impact.index.set_levels(recall_model_impact.index.levels[0].map(week_to_date), level='order_created_week')
total_order.index = total_order.index.map(week_to_date)
shipping_order.index = shipping_order.index.map(week_to_date)
shipping_order = shipping_order.reset_index()

recall_model_impact['total_rto%'] =  recall_model_impact['rto_safe']+recall_model_impact['rto_risky']
total_rto_order =  recall_model_impact.groupby('order_created_week')['total_rto%'].sum()

#recall%
recall_model_impact['recall%'] = (recall_model_impact['rto_risky']/(recall_model_impact['rto_risky']+recall_model_impact['rto_safe'])*100).round(2)

#order%
recall_model_impact['rto order%'] = ((recall_model_impact['rto_safe']+recall_model_impact['rto_risky'])/total_rto_order*100).round(2)

# shipping_data
recall_model_impact['shipping data available%'] = (recall_model_impact['order_count_w_delivery']/total_order*100).round(2)

#dropping unused columns
recall_model_impact = recall_model_impact.drop(['rto_safe','non_rto_risky','non_rto_safe','order_id','rto_risky','total_rto%','order_count_w_delivery','order_count_w_rto'],axis=1)


recall_model_impact = recall_model_impact.stack()
recall_model_impact = recall_model_impact.unstack(level = 0)
recall_model_impact = recall_model_impact.unstack(level = 1)
recall_model_impact = recall_model_impact.drop(recall_model_impact.columns[0:6], axis=1)
recall_model_impact = recall_model_impact.drop(recall_model_impact.columns[-6:], axis=1)
recall_model_impact = recall_model_impact.fillna(0)

previous_week_cr = []
previous_week_volm = []
for j in range(len(recall_model_impact[week_to_date(week-4)])):
    previous_week_cr_value = recall_model_impact[week_to_date(week-4)]['recall%'].iloc[j]
    previous_week_volume_value = recall_model_impact[week_to_date(week-4)]['rto order%'].iloc[j]
    previous_week_cr.append(previous_week_cr_value)
    previous_week_volm.append(previous_week_volume_value)
recall_model_impact[(' ','previous_week_cr')] = previous_week_cr
recall_model_impact[(' ','previous_week_volm')] = previous_week_volm

recall_model_impact[(' ','Volume Impact')] = (recall_model_impact[(' ','previous_week_cr')] * (recall_model_impact[(week_to_date(week-3)),'rto order%']-recall_model_impact[(' ','previous_week_volm')])/100).round(2)

recall_model_impact[(' ','CR Impact')] = (recall_model_impact[(week_to_date(week-3),'rto order%')] * (recall_model_impact[(week_to_date(week-3)),'recall%']-recall_model_impact[(' ','previous_week_cr')])/100).round(2)

recall_model_impact[(' ','Total Impact')] = (recall_model_impact[(' ','Volume Impact')] +  recall_model_impact[(' ','CR Impact')]).round(2)
recall_model_impact = recall_model_impact.drop(recall_model_impact.columns[6:8],axis=1)
# recall_model_impact

last_week_volume_total = percentage_conversion(recall_model_impact[(' ', 'Volume Impact')].sum())
last_week_cr_total = percentage_conversion(recall_model_impact[(' ', 'CR Impact')].sum())
last_week_total_impact = percentage_conversion(recall_model_impact[(' ', 'Total Impact')].sum())
recall_vol_cr_total = [last_week_volume_total, last_week_cr_total, last_week_total_impact]


recall_df_total = []
for weeks in range(week - num_weeks + 2 , week  - 2):
    date = week_to_date(weeks)
    recall_total = percentage_conversion(recall_df[recall_df['order_created_week']==date].values[0][1])
    volumne_total = percentage_conversion(recall_model_impact[(date),'rto order%'].sum().round(2))
    shipping_total = percentage_conversion(recall_model_impact[(date),'shipping data available%'].sum().round(2))
    recall_model_impact[(date,'rto order%')] = recall_model_impact[(date,'rto order%')].apply(percentage_conversion)
    recall_model_impact[(date,'recall%')] = recall_model_impact[(date,'recall%')].apply(percentage_conversion)
    recall_model_impact[(date,'shipping data available%')] = recall_model_impact[(date,'shipping data available%')].apply(percentage_conversion)
    recall_df_total.append(recall_total)
    recall_df_total.append(volumne_total)
    recall_df_total.append(shipping_total)

recall_model_impact[(' ', 'Total Impact')] = recall_model_impact[(' ', 'Total Impact')].apply(percentage_conversion)
recall_model_impact[(' ', 'Volume Impact')] = recall_model_impact[(' ', 'Volume Impact')].apply(percentage_conversion)
recall_model_impact[(' ', 'CR Impact')] = recall_model_impact[(' ', 'CR Impact')].apply(percentage_conversion)
recall_vol_cr_total = recall_df_total + recall_vol_cr_total
recall_vol_cr_total = pd.DataFrame([recall_vol_cr_total],columns=recall_model_impact.columns,index=['Total'])


recall_model_impact = pd.concat([recall_model_impact, recall_vol_cr_total]).reset_index()
recall_model_impact = recall_model_impact.rename(columns={'index':'model_id'})
recall_model_impact


# COMMAND ----------

#waterfall chart calculation

previous_week_date = week_to_date(week-3)
volume = recall_model_impact[(' ','Volume Impact')].iloc[-1]
recall_chart = recall_model_impact[['model_id']]
recall_chart = recall_chart.drop(recall_model_impact.index[-1])
recall_chart['Total Impact'] = recall_model_impact.loc[:,(' ','CR Impact')]
volume_impact = pd.DataFrame({('model_id', ''):['volume_impact'],('Total Impact', ''):[volume]})
subtotal_value = recall_model_impact[(' ','Total Impact')].iloc[-1]
subtotal = pd.DataFrame({('model_id', ''):['subtotal'],('Total Impact', ''):[subtotal_value]})
recall_waterfall_table = pd.concat([recall_chart,volume_impact,subtotal])

second_last_row = recall_waterfall_table.iloc[-2].copy()  
recall_waterfall_table.iloc[1:-1] = recall_waterfall_table.iloc[:-2].values  
recall_waterfall_table.iloc[0] = second_last_row

pd.set_option('chained_assignment', None)
# recall_waterfall_table


#waterfall_chart
recall_waterfall_table['Total Impact'] = recall_waterfall_table['Total Impact'].str.replace('%',' ')
recall_waterfall_table['Total Impact'] = pd.to_numeric(recall_waterfall_table['Total Impact'])

model_id = recall_waterfall_table['model_id'].values
impact = recall_waterfall_table['Total Impact'].values

measures =  ['relative'] + ['relative'] * (len(model_id) - 2) + ['total']

trace = go.Waterfall(
    x=model_id,
    y=impact,
    measure=measures,
    increasing=dict(marker={'color': 'green'}),
    decreasing=dict(marker={'color': 'red'}),
    totals=dict(marker={'color': 'rgb(255,215,0)'}),
    connector=dict(line={'dash': 'solid', 'width': 1}),
    width=0.6,
    text=recall_waterfall_table['Total Impact'].apply(float).apply(lambda x: round(x , 2)).apply(lambda x:f'<b style="font-size:30pt">{x}</b>'),
    textposition='outside',
)

layout = go.Layout(
    title=' Recall Rate',
    title_font=dict(size=24),
    width=750,
    height=1050,
)
recall_flag_fig = go.Figure(data=trace, layout=layout)
recall_flag_fig = recall_flag_fig.update_xaxes(tickfont=dict(size=40,family='Arial, bold'))
recall_flag_fig.show()


# COMMAND ----------



def save_plotly_figure_as_image(fig, document,width_inches, height_inches, dpi):
    image_stream = io.BytesIO()
    fig.write_image(image_stream, format='png',width=width_inches * dpi, height=height_inches * dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def save_plotly_line_figure_as_image(fig, document,width_inches, height_inches, dpi):
    image_stream = io.BytesIO()
    fig.savefig(image_stream, format='png',dpi= dpi)
    image_stream.seek(0)
    document.add_picture(image_stream, width=Inches(width_inches), height=Inches(height_inches))

def create_document_from_multi_level_list(data):
    doc = Document()
    for section in data:
        section_title = section[0]
        section_content = section[1:]

        doc_heading = doc.add_heading(section_title, level=1)
        doc_heading.paragraph_format.space_before = Pt(8)
        doc_heading.paragraph_format.space_after = Pt(8)

        if not isinstance(section_content, list):
            section_content = [section_content]

        for item in section_content:
            if isinstance(item, list):
                subheading = item[0]
                subcontent = item[1]
                heading=doc.add_heading(subheading, level=2)
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_before = Pt(10)

                if isinstance(subcontent, pd.DataFrame):
                    paragraph = doc.add_paragraph()
                    heading.paragraph_format.space_before = Pt(8)
                    heading.paragraph_format.space_after = Pt(10)
                    table = doc.add_table(subcontent.shape[0]+2, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    last_paragraph = table.rows[-1].cells[0].paragraphs[-1]
                    last_paragraph.paragraph_format.space_after = Pt(10)  

                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(j%2 == 0 ):
                                table.cell(0,i).text = str(col)
                            else:
                                continue
                    for i,col_tuple in enumerate(subcontent.columns):
                        for j, col in enumerate(col_tuple):
                            if(j%2 != 0):
                                table.cell(1,i).text = str(col)
                            else:
                                continue
                            
                    for i, row_tuple in enumerate(subcontent.itertuples(index=False)):
                        for j, value in enumerate(row_tuple):
                            table.cell(i + 2, j).text = str(value)
                    
                    table_spacing = doc.add_paragraph()
                    table_spacing.paragraph_format.space_before = Pt(8)
                    table_spacing.paragraph_format.space_after = Pt(8)

                elif isinstance(subcontent, go.Figure):
                    subcontent.update_xaxes(tickfont=dict(size=28))
                    save_plotly_figure_as_image(subcontent, doc, width_inches=6.2, height_inches=4.8, dpi=300)
                elif isinstance(subcontent, plt.Figure):
                    # subcontent.update_xaxes(tickfont=dict(size=18))
                    save_plotly_line_figure_as_image(subcontent, doc, width_inches=6.8, height_inches=4.4, dpi=360)
                
            else:
                raise ValueError("Invalid content format. Expected list.")
    return doc


# COMMAND ----------


data_list = [
    ['COD Intelligence',
        ['Weekly Red Flagging Rate', red_label_img],
        ['Model Impact',red_labelling_model_impact],
        ['Red Flagging Waterfall Chart',red_flag_fig],
        ['Precision, Recall, F Score  (4-week Moving Average)',moving_avg],
        ['Precision, Recall, F Score - Weekly Comparison',precision_recall_fscore_comparison],
        ['Precision',weekly_precision_rate],
        ['Precision model impact',precision_model_df],
        ['Precision waterfall chart', precision_flag_fig],
        ['Recall weekly rate',recall_weekly_rate],
        ['Recall model impact',recall_model_impact],
        ['Recall waterfall chart', recall_flag_fig],
        
        
    ]
    
]
document = create_document_from_multi_level_list(data_list)
document.save('hm_cod_data_list.docx')

# COMMAND ----------

dbutils.fs.cp("file:/databricks/driver/hm_cod_data_list.docx", "dbfs:/FileStore/shared_transfer/Pallavi_Samodia/hm_cod_data_list.docx")
# https://razorpay-dev.cloud.databricks.com/files/shared_transfer/Pallavi_Samodia/hm_cod_data_list.docx

# COMMAND ----------


