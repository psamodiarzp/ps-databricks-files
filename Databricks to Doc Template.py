# Databricks notebook source
from datetime import datetime, timedelta
from datetime import date
import pandas as pd


# COMMAND ----------

import matplotlib.pyplot as plt #in case any pyplot graphs need to be added

# COMMAND ----------

!pip install python-docx

# COMMAND ----------

import docx
import os
from docx import Document
from docx.shared import Cm, Pt

# COMMAND ----------

# DBTITLE 1,Sample Table
starting_dt = '2023-07-01'
ending_dt = '2023-09-30'
data_db = sqlContext.sql(
    """
select * from batch_sheets.magic_merchant_list
where live_date >= '{0}'
and live_date <= '{1}'
""".format(
        starting_dt, ending_dt
    )
)
data_df = data_db.toPandas()
data_df.head()

# COMMAND ----------

sales_merchant_status_df = (
    (
        data_df.groupby(by=["sales_merchant_status", ])
        .agg({"mid": "count"})
        .reset_index()
    )
)
sales_merchant_status_df

# COMMAND ----------

# DBTITLE 1,Sample Table 2
segment_status_df = (
    (
        data_df.groupby(by=["segment", "sales_merchant_status"])
        .agg({"mid": "count"})
        .reset_index()
    )
    .pivot_table(index=["segment"], columns="sales_merchant_status", values="mid")
    .reset_index()
)
segment_status_df

# COMMAND ----------

# DBTITLE 1,Doc Creation Code

def create_document_from_multi_level_list(data):
    doc = docx.Document()
    for section in data:
        section_title = section[0]
        section_content = section[1:]
        
        doc.add_heading(section_title, level=1)
        
        for item in section_content:
            if isinstance(item, list):
                subheading = item[0]
                subcontent = item[1]
                doc.add_heading(subheading, level=2)
                
                if isinstance(subcontent, pd.DataFrame):
                    doc.add_paragraph()
                    table = doc.add_table(subcontent.shape[0]+1, subcontent.shape[1])
                    table.style = 'LightGrid-Accent1'
                    for i, column in enumerate(subcontent.columns):
                        table.cell(0, i).text = column
                        for j, value in enumerate(subcontent[column]):
                            table.cell(j+1, i).text = str(value)
                elif isinstance(subcontent, plt.Figure):
                    # Save the graph to a BytesIO buffer
                    buffer = BytesIO()
                    subcontent.savefig(buffer, format='png')
                    buffer.seek(0)  # Reset the buffer position to the beginning
                    
                    # Add the graph to the document
                    doc.add_picture(buffer, width=docx.shared.Inches(4), height=docx.shared.Inches(3))
                    
                    # Clear the buffer and close the figure
                    buffer.close()
                    plt.close(subcontent)
            else:
                raise ValueError("Invalid content format. Expected list.")
    return doc


# COMMAND ----------

# DBTITLE 1,Input for Doc Creation
data_list = [
    ['Acquisition',
     ['Sample Table 1',sales_merchant_status_df],
        ['Sample Table 2',segment_status_df],

    
    ],
    [
        'Section 2'
    ],
    [
        'Section 3'
    ]
]
document = create_document_from_multi_level_list(data_list)
document.save('sample.docx')

# COMMAND ----------



# COMMAND ----------

your_file_name = 'sample.docx'
dbutils.fs.cp("file:/databricks/driver/{0}".format(your_file_name), "dbfs:/FileStore/shared_transfer/{0}".format(your_file_name))
# https://razorpay-dev.cloud.databricks.com/files/shared_transfer/sample.docx
# Copy this URL to the browser and that will download your doc to the local storage

# COMMAND ----------


